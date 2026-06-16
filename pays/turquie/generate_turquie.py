"""
MAB Turquie v1 — Génération des deux documents Word
Sources internes utilisées :
  - MAB_Turquie_PREP.md (corpus PREP, juin 2026)
  - Turquie Synthèse des risques 07-2024-2.pdf (fiche pays Coface, juillet 2024)
  - Dossier Turquie PPT 2 - FINAL.pdf (recherche interne propriétaire, PESTEL/ERP/inox)
  - study_id56671_construction-industry-in-turkey (Statista, 2024)
  - TR_Bathrooms_Full_Report_Apr21.pdf (BRG Building Solutions, avril 2021, données 2020)
    → Total taps & mixers TR 2020 : 1 229,75 M TRY / 152,96 M EUR (MSP), 6 776 000 unités
    → Non-Housing (ERP) : 7,11% volume (482 010 unités)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

FONT = "Calibri"

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = FONT
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def bold_bullet(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)
    r1 = p.add_run(f"{label} : ")
    set_font(r1, bold=True)
    r2 = p.add_run(value)
    set_font(r2)
    return p

def add_para(doc, text, size=10, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, italic=italic)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {text}")
    set_font(run, size=9, italic=True, color=(0xC0, 0x50, 0x20))
    return p

def make_table(doc, rows_data, font_size=9, header_color=(0x1F, 0x49, 0x7D)):
    tbl = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    tbl.style = "Table Grid"
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.name = FONT
            run.font.size = Pt(font_size)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(*header_color)
    return tbl

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
set_margins(doc)

t = doc.add_heading("MAB TURQUIE — ÉTUDE DE MARCHÉ v2", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Robinetterie sanitaire collective / ERP — Les Robinets Presto")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(12)
sub.runs[0].font.name = FONT
date_p = doc.add_paragraph("Juin 2026")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(10)
date_p.runs[0].font.italic = True
date_p.runs[0].font.name = FONT
doc.add_paragraph()

# ─── RÉSUMÉ EXÉCUTIF ─────────────────────────────────────────────────────────
add_heading(doc, "RÉSUMÉ EXÉCUTIF", 1)
add_bullet(doc, "Marché robinetterie ERP Turquie estimé entre 20,7 et 52,9 M€ selon la méthode (extrapolation France ajustée vs BRG TR × coefficient ERP 10-15%) — fourchette cohérente avec l'estimation interne propriétaire de ~56 M€ (~4M pièces, TCAC ~4%, 2024-2030). Marché à fort volume mais très price-driven : PIB/hab 18 611 USD (0,38× la France) compensé par une population de 85,5 M habitants. (IMF WEO 10/2025 ; Dossier interne Presto)")
add_bullet(doc, "DELABIE EST ABSENT DU MARCHÉ TURC — aucune des 9 filiales internationales du groupe ne couvre la Turquie, aucune référence projet identifiée. Fenêtre d'opportunité réelle, mais signal que le marché est jugé difficile d'accès par les acteurs premium européens (sensibilité prix extrême, dominance locale Eczacıbaşı/VitrA/Artema/ECA à 60-70% volume). (Delabie.com, 2026 ; Dossier interne Presto)")
add_bullet(doc, "Pipeline d'investissement public massif et sécurisé : City Hospitals Program (105 Mds TRY/an PPP), reconstruction post-séisme (45-80 Mds$ au total, budget éducation 160 M$), programme TOKİ 500 000 logements sociaux (2025-). Construction non-résidentielle (23,5% du marché total, sous la moyenne européenne ~30,2%) tirée par la santé et l'éducation. (World Bank 2023 ; bne IntelliNews 2025 ; AGBI 2024 ; TurkStat)")
add_bullet(doc, "Top segments prioritaires : Établissements pénitentiaires (1er marché inox sanitaire turc — 68-72% du parc en inox, 200 000 pièces, 22 nouvelles prisons programmées 2025-2028), Santé (1 560 hôpitaux, programme City Hospitals), Éducation (75 500 établissements, reconstruction post-séisme). Stress hydrique structurel (Ankara, Izmir) = argument fort pour économiseurs d'eau.")
add_bullet(doc, "Risques structurels majeurs : inflation (~30,65% début 2026), dépréciation continue de la livre turque (-27% face à l'EUR sur un an, ~51 TRY/EUR en 2026), économie informelle significative, gouvernance opaque sur les marchés publics pénitentiaires (concentration sur quelques attributaires). Accès marché aligné UE (union douanière 1996) mais contrôle TSE/TAREKS systématique à l'importation au-delà du seul marquage CE.")
doc.add_paragraph()

# ─── PARTIE 1 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 1 — OVERVIEW CONTEXTE PAYS", 1)

add_heading(doc, "1.1 Analyse PESTEL", 2)
bold_bullet(doc, "POLITIQUE", "République présidentielle, R.T. Erdoğan confirmé en 2023 mais recul de son parti aux municipales 2024. Catégorie risque OCDE 5/7 (2024) ; notations S&P B+ / Moody's B3 / Fitch B. Relations tendues mais stabilisées avec l'UE/USA, renforcement des liens avec le Golfe. Climat des affaires peu prévisible pour investisseurs étrangers malgré des améliorations. (Direction Évaluation, Études et Prospective, juillet 2024)")
bold_bullet(doc, "ÉCONOMIQUE", "PIB/hab 18 200-18 611 USD (2025, fourchette selon source — 0,38× la France). Croissance attendue 3,5% (2025) → 3,7% (2026) → 4,4% (2027). Désinflation progressive 29%→18%→15% (2025-2027) mais inflation encore élevée (~30,65% en 01/2026, taux directeur 37%). Déficit public en consolidation : 4,3% (2025) → 3,6% (2026). Dette publique modérée (31% PIB) mais structure dégradée. (IMF WEO 10/2025 ; Banque mondiale ; Trading Economics)")
bold_bullet(doc, "SOCIAL", "Population 85,5 M (2025) à 86,09 M (TurkStat, fin 2025). Densité 108 hab/km². Espérance de vie 81 ans (F) / 76 ans (H). Alphabétisation 96,2%. Société à système hiérarchique ancré (Hofstede) : forte aversion à l'incertitude, culture collective, valorisation de la loyauté et de la tradition plutôt que de la réussite individuelle.")
bold_bullet(doc, "TECHNOLOGIQUE", "Marché en transition vers le sans-contact/électronique en ERP (encore minoritaire, 6-9% du marché robinetterie ERP) mais en croissance, porté par la défiance sanitaire post-pandémique et le contexte de stress hydrique. Acteurs locaux (ECA, Artema, VitrA, Creavit, Kale) couvrent déjà ce segment aux côtés des marques européennes premium.")
bold_bullet(doc, "ENVIRONNEMENTAL", "Stress hydrique structurel sévère : Ankara en crise (sécheresse la plus sèche depuis 50 ans en 2025, réserves ~1,12%), tensions résidents/tourisme à Izmir, défiance généralisée envers l'eau du robinet (recours à l'eau en bouteille). Objectif national -21% d'émissions carbone d'ici 2030. (Direction Évaluation 07/2024 ; Dossier interne Presto)")
bold_bullet(doc, "LÉGAL", "Union douanière UE-Turquie (1996) : exonération de droits de douane sur produits industriels européens, MAIS marquage CE jugé insuffisant en pratique — contrôle systématique TSE via le système TAREKS à l'importation. Normes TS EN quasi-identiques aux normes EN françaises (transpositions directes). ISO 9001/14001 attendus en complément. (ab.gov.tr ; JJRLAB 2024-2025)")
doc.add_paragraph()

add_heading(doc, "1.2 Indicateurs socio-économiques clés", 2)
make_table(doc, [
    ["Indicateur", "Valeur", "Source"],
    ["PIB/habitant", "18 200-18 611 USD (2025)", "IMF WEO 10/2025 / Worldometer"],
    ["Population", "85,5 M (2025) — 86,09 M fin 2025", "Source interne / TurkStat"],
    ["Densité", "108 hab./km²", "Source interne"],
    ["Croissance PIB", "3,5% (2025) → 3,7% (2026) → 4,4% (2027)", "Banque mondiale"],
    ["Inflation", "~30,65% (01/2026), pic 75% mi-2024", "Trading Economics"],
    ["Taux directeur banque centrale", "37%", "Trading Economics, 2026"],
    ["Déficit public", "4,3% PIB (2025) → 3,6% (2026)", "FMI / Banque mondiale"],
    ["Dette publique", "31% du PIB (structure dégradée)", "Direction Évaluation 07/2024"],
    ["Taux de change EUR/TRY", "≈51 TRY/EUR (2026) ; -27% sur 1 an", "ECB eurofxref / Capital.com 02/2026"],
    ["Catégorie risque OCDE", "5/7 (2024)", "Direction Évaluation 07/2024"],
    ["Notations", "S&P B+ / Moody's B3 / Fitch B", "Direction Évaluation 07/2024"],
])
doc.add_paragraph()

add_heading(doc, "1.3 Relations économiques et culturelles avec la France", 2)
add_bullet(doc, "Échanges bilatéraux 2024 : 23,4 Md€ (exports FR→TR 11,9 Md€ / imports 11,5 Md€ en 2023), excédent français ~1,3 Md€. France = 7e fournisseur de la Turquie. (Source interne, dossier propriétaire)")
add_bullet(doc, "Postes 2024 : exports françaises → automobile (23,7%), biens d'équipement (20,6%), aéronautique (15,3%) ; imports depuis Turquie → véhicules (35%), textile (17,7%), biens d'équipement (16,5%).")
add_bullet(doc, "Présence française structurante : ~400 entreprises, ~100 000 emplois (donnée 2012, à rafraîchir). Renault, Alstom, Thales, Schneider Electric, Air Liquide, Safran, Dassault Systèmes (industrie/tech) ; GDF Suez (énergie) ; BNP/AXA/Groupama (finance) ; Carrefour/Leroy Merlin/Décathlon (distribution). Coopération aéronautique Airbus/Turkish Airlines/Turkish Aerospace.")
add_bullet(doc, "Proximité culturelle limitée mais structurée : ambassade Ankara, consulat général Istanbul, 10 895 Français enregistrés. Hofstede : accès direct aux décideurs nécessaire, relation de confiance à construire sur la durée, formalisation forte des process attendue.")
add_bullet(doc, "Barrières d'entrée spécifiques : marquage CE insuffisant seul (contrôle TAREKS), documentation technique exigée en turc (traduction notariée), liste annuelle turque de normes d'importation obligatoires pouvant s'appliquer à certains composants de robinetterie. (ab.gov.tr ; JJRLAB 2024-2025 ; trade.gov)")
add_bullet(doc, "Opportunité : union douanière UE-Turquie = 0% droit de douane sur produits industriels européens — avantage structurel à exploiter malgré les contrôles techniques additionnels.")
doc.add_paragraph()

add_heading(doc, "1.4 Tendances d'investissement", 2)
add_para(doc, "Programmes d'investissement publics identifiés :")
make_table(doc, [
    ["Programme", "But global", "Secteur", "Budget", "Calendrier", "Opportunité Presto"],
    ["Türkiye Earthquake Recovery and Reconstruction Project (Banque mondiale)", "Restaurer logements ruraux et services publics dans les zones sinistrées (séismes février 2023)", "Logement rural, santé, infrastructures municipales", "1 Md$ prêt Banque mondiale ; coût total reconstruction 45-80 Md$", "En cours depuis 2023, pluriannuel", "Robinetterie anti-vandalisme inox, temporisée (écoles, centres de santé reconstruits)"],
    ["Plan de reconstruction post-séisme — volet logements (TOKİ)", "Reconstruire 488 000 logements + écoles/hôpitaux en zone sinistrée", "Logement, éducation, santé", "[DONNÉE NON DISPONIBLE — non isolé de l'enveloppe globale]", "Lancé 2023, livraisons 2023-2025+", "Robinetterie standard collectivité, mitigeurs thermostatiques"],
    ["City Hospitals Program (PPP santé)", "Moderniser le parc hospitalier via PPP (build-operate-transfer)", "Santé — hôpitaux universitaires/régionaux", "105 Mds TRY (2025, contractants city hospitals) ; 202 Mds TRY total PPP", "Lancé 2016 ; >20 city hospitals ouverts depuis 2022", "Robinetterie haute fréquence, PMR, thermostatique, sans-contact (plateaux techniques)"],
    ["TOKİ Urban Transformation — \"Housing Project of the Century\"", "500 000 logements sociaux à l'échelle nationale", "Logement résidentiel (effet d'entraînement parties communes)", "2 984 Mds TRY cumulés depuis 2002 ; nouveau volet 2025", "Lancé 10/2025, pluriannuel", "Robinetterie de collectivité parties communes (priorité secondaire)"],
], font_size=8)
doc.add_paragraph()
add_note(doc, "Aucun programme nommé et chiffré spécifique trouvé pour l'éducation au-delà du budget annuel 42 Mds€/an, ni pour un programme BERD/UE dédié aux ERP turcs — à approfondir en recherche complémentaire.")
doc.add_paragraph()

# ─── PARTIE 2 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 2 — MARCHÉ DE LA CONSTRUCTION", 1)

add_heading(doc, "2.1 État actuel du marché", 2)
add_bullet(doc, "Valeur du marché construction turc 2024-2025 : fourchette large selon cabinet — 53 Mds$ (Verified Market Research) à 174 Mds$ (Mordor Intelligence) ; 117,75 Mds$ (IMARC) ; 1,48 trillion TRY ≈ 35-37 Mds$ (Research and Markets). Écart méthodologique important à trancher via une source TurkStat officielle unique en phase terrain.")
add_bullet(doc, "Contribution du BTP au PIB : 1 466,47 Md TRY (2023) vs 190,62 Md TRY (2015) — croissance x7,7 en valeur nominale (effet change + inflation + volume). (TurkStat)")
add_bullet(doc, "Croissance annuelle BTP : +7,8% (2023) après -8,6% (2022) et -7,1% (2021) ; +6,5% au T2 2024. Reprise confirmée mais volatile. (Turkish Contractors Association/TurkStat)")
add_bullet(doc, "Entreprises leaders (revenus 2023) : Renaissance Construction (1 209 M$), Ant Yapi Industry & Trade (840 M$), Limak Insaat Sanayi ve Ticaret (558 M$). (Statista)")
doc.add_paragraph()

add_heading(doc, "2.2 Dynamique Neuf vs Rénovation", 2)
add_note(doc, "Estimation — à confirmer. Le ratio neuf/rénovation global (Mordor Intelligence) ne doit pas être confondu avec le ratio résidentiel/non-résidentiel (TurkStat) ci-dessous — leur proximité numérique est une coïncidence statistique.")
make_table(doc, [
    ["Segment", "Poids estimé", "Dynamique", "Source"],
    ["Neuf résidentiel", "Part majoritaire du neuf (76,35% du marché total est neuf, dont logement dominant)", "Soutenu par reconstruction post-séisme + Urban Transformation Plan", "Mordor Intelligence 2025"],
    ["Rénovation résidentielle", "23,65% du marché total (tous segments)", "RMI portée par durée de vie courte des équipements (cf. Partie 5) + rénovation parasismique", "Mordor Intelligence 2025"],
    ["Non-résidentiel (neuf + rénov.)", "23,5% du marché total construction (2023)", "Santé/éducation en croissance ; permis non-résidentiels en recul -8,7% (9 premiers mois 2024)", "TurkStat / Association of Turkish Construction Material Producers"],
    ["Génie civil / Infrastructure", "31,05% du marché total (estimation Mordor, périmètre élargi)", "270 projets PPP achevés (204 Mds$ cumulés) ; mégaprojets (Canal Istanbul)", "Mordor Intelligence 2025"],
])
doc.add_paragraph()
add_bullet(doc, "Répartition résidentiel/non-résidentiel (dépenses construction, TurkStat officiel) : 76,5% / 23,5% en 2023, vs 79,7%/20,3% en 2017 — part du non-résidentiel en hausse structurelle.")
doc.add_paragraph()

add_heading(doc, "2.3 Perspectives 2025-2030", 2)
add_bullet(doc, "Croissance attendue : CAGR 5,45% (2026-2031, Mordor Intelligence) ; +3% en 2025 (1,48 trillion TRY) puis accélération à CAGR 3,7% (2026-2029, Research and Markets) — fourchette large à conserver selon la source.")
add_bullet(doc, "Programme d'investissement public 2025 : 1,9 trillion TRY (46,2 Mds$) pour 3 783 projets multisectoriels (transport, énergie, santé, mines). (AGBI/GlobeNewswire 2025)")
add_bullet(doc, "270 projets PPP déjà achevés représentant 204 Mds$ cumulés — pipeline structurel sécurisé sur la décennie. (Mordor Intelligence 2025)")
add_bullet(doc, "Reconstruction post-séisme : facteur structurel majeur de la demande jusqu'à fin de décennie, en complément du programme de transformation urbaine (rénovation parasismique).")
add_bullet(doc, "Risques chiffrés : inflation des coûts matériaux +20-30%/an ; hausse des salaires main d'œuvre qualifiée +114% depuis 2023 ; dépréciation TRY -17 à -21% face USD (08/2024-09/2025) ; pénurie de main d'œuvre qualifiée généralisée. (Türkiye Today 2025 ; Trading Economics)")
doc.add_paragraph()

# ─── PARTIE 3 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 3 — CONSTRUCTION NON-RÉSIDENTIELLE", 1)

add_heading(doc, "3.1 État actuel", 2)
add_bullet(doc, "Part du non-résidentiel dans la construction totale : 23,5% (2023, TurkStat — donnée officielle à privilégier) — en dessous de la moyenne européenne Euroconstruct (~30,2%), signe d'un marché encore très orienté résidentiel sous pression démographique et post-séisme. (TurkStat ; Euroconstruct 2024-2025)")
add_bullet(doc, "Valeur estimée : application du ratio 23,5% à la fourchette totale 117-174 Mds$ → 27,5 à 41 Mds$. Estimation — méthode = application directe du ratio TurkStat 2023 au marché total devise forte 2024-2025, à confirmer.")
add_bullet(doc, "Croissance non-résidentielle 2024 : +11,1% en valeur de production, MAIS recul des permis de construire non-résidentiels de -8,7% sur les 9 premiers mois 2024 — signal avancé de ralentissement à surveiller pour le pipeline neuf à moyen terme.")
doc.add_paragraph()

add_heading(doc, "3.2 Dynamique Neuf vs Rénovation", 2)
add_note(doc, "Aucune source turque spécifique trouvée distinguant neuf/rénovation au sein du seul segment non-résidentiel. Estimation — à confirmer : rénovation ≈25-30% / neuf ≈70-75%, hypothèse fondée sur le ratio global ajusté à la hausse pour la rénovation du fait des obligations de mise aux normes parasismiques des bâtiments publics existants.")
make_table(doc, [
    ["Sous-segment", "Poids estimé", "Dynamique", "Source"],
    ["Santé (neuf)", "50-65% des projets santé", "City Hospitals PPP — vague de grands projets neufs", "bne IntelliNews 2025"],
    ["Santé (rénovation)", "35-50% des projets santé", "Mise aux normes parasismique, modernisation technique post-séisme", "Dossier interne Presto"],
    ["Éducation (neuf + rénov.)", "Budget 150,78 Mds TRY / 1 027 projets (2024)", "Reconstruction post-séisme (budget 160 M$, 58% salles réparées) + rénovation sanitaire/hygiène", "AGBI 01/2024 ; Dossier interne"],
])
doc.add_paragraph()

add_heading(doc, "3.3 Segments dominants", 2)
make_table(doc, [
    ["Segment", "Taille / Drivers chiffrés", "Source"],
    ["Santé (City Hospitals)", "105 Mds TRY (2025, PPP city hospitals, ~10% budget santé) ; 94,56 Mds TRY part santé budget investissement 2024 ; >20 city hospitals ouverts depuis 2022 (ex. Gaziantep, 1 875 lits, 02/2024)", "bne IntelliNews 2025 ; invest.gov.tr 2024"],
    ["Éducation", "150,78 Mds TRY pour 1 027 projets (15% budget investissement public total) ; dépense/élève en recul (4 932→4 491 USD, 2015-2022) ; part éducation budget public en baisse 12,9%→10,6%", "AGBI 01/2024 ; OCDE Education at a Glance 2025"],
    ["Sport-Jeunesse-Culture", "[DONNÉE NON DISPONIBLE en sources externes — voir Partie 4.7 pour données de parc internes]", "—"],
    ["Administratif-Sécurité (dont pénitentiaire)", "[DONNÉE NON DISPONIBLE en sources externes — voir Partie 4.8 pour données de parc internes]", "—"],
    ["Religieux", "[DONNÉE NON DISPONIBLE en sources externes — voir Partie 4.10 pour données de parc internes]", "—"],
], font_size=8)
doc.add_paragraph()

add_heading(doc, "3.4 Perspectives 2025-2030", 2)
add_bullet(doc, "Croissance 2025 attendue +3% (1,48 trillion TRY), accélération à CAGR 3,7% (2026-2029). (Research and Markets 04 et 08/2025)")
add_bullet(doc, "Segment santé : flux récurrent garanti par le modèle PPP (~10% du budget santé alloué aux contrats city hospitals) — pipeline pluriannuel visible.")
add_bullet(doc, "Segment infrastructure (31,05% du marché total) porté par 270 projets PPP achevés (204 Mds$) + mégaprojets (Canal Istanbul).")
add_bullet(doc, "Risque structurel sur l'éducation : part du budget public en baisse continue (-2,3 points 2015-2022), pouvant freiner la construction de nouveaux établissements scolaires publics — à surveiller comme frein potentiel.")
add_bullet(doc, "Freins macro transversaux : inflation matériaux, dépréciation TRY, pénurie main d'œuvre qualifiée — impact accentué sur bâtiments publics dont les équipements importés (robinetterie) renchérissent mécaniquement en devise locale.")
doc.add_paragraph()

# ─── PARTIE 4 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 4 — POIDS DES SEGMENTS ERP", 1)

add_heading(doc, "4.0 Contexte général et synthèse ERP Turquie", 2)
add_para(doc, "La Turquie dispose d'un parc ERP jeune et en forte expansion (transformation massive 2000-2020), mais marqué par une dynamique de rattrapage post-séisme et une gouvernance parfois opaque sur certains segments (pénitentiaire). Le marché est extrêmement price-driven, dominé par des acteurs locaux à 60-70% en volume.")
add_bullet(doc, "Reconstruction post-séisme structurante : budget global 45-80 Mds$, touchant éducation, santé et logement simultanément — fenêtre d'opportunité concentrée 2023-2028.")
add_bullet(doc, "Le parc inox sanitaire ERP turc est déjà cartographié par segment (dossier interne propriétaire) : pénitentiaire 1er secteur (68-72% du parc en inox), lieux de culte 2e, éducation 3e, santé 4e, sport 5e.")
add_bullet(doc, "Stress hydrique structurel (Ankara, Izmir) → argument commercial fort et récurrent pour économiseurs d'eau, transversal à tous les segments ERP.")
add_bullet(doc, "Durée de vie des robinetteries en Turquie nettement plus courte qu'en Europe (max 12 ans vs >20 ans) → cycle de remplacement plus fréquent, opportunité structurelle pour le marché du remplacement (RMI). (BRG TR Apr21)")
doc.add_paragraph()

make_table(doc, [
    ["Segment ERP", "Score Presto", "Taille / Signal", "Priorité"],
    ["4.8 Pénitentiaire", "5/5", "407 établissements, 402 000 incarcérés, 22 nouvelles prisons 2025-2028, 1er marché inox sanitaire (200 000 pièces)", "PRIORITÉ ABSOLUE"],
    ["4.2 Santé", "5/5", "1 560 hôpitaux, City Hospitals PPP 105 Mds TRY/an, 4e marché inox (105-135K pièces)", "PRIORITÉ ABSOLUE"],
    ["4.1 Éducation", "4/5", "75 500 établissements, reconstruction post-séisme (9 800 salles), 3e marché inox (120-150K pièces)", "PRIORITÉ HAUTE"],
    ["4.7 Sport & Loisirs", "3/5", "4 500 équipements, EURO 2032 co-organisation, 5e marché inox (100-130K pièces)", "SECONDAIRE"],
    ["4.10 Lieux de culte", "3/5", "82 000 mosquées, impact séisme (3 800+ endommagées), 2e marché inox (120-150K pièces)", "SECONDAIRE"],
    ["4.11 Transports (aéroports)", "3/5", "Investissements massifs, ambition hub Europe-Asie", "SECONDAIRE"],
    ["4.3 Tertiaire", "2/5", "[DONNÉE NON DISPONIBLE — non couvert par sources internes]", "OPPORTUNISTE"],
    ["4.4 Industriel", "2/5", "[DONNÉE NON DISPONIBLE — non couvert par sources internes]", "OPPORTUNISTE"],
    ["4.5 CHR", "2/5", "[DONNÉE NON DISPONIBLE — non couvert par sources internes]", "OPPORTUNISTE"],
    ["4.6 HPA", "1/5", "[DONNÉE NON DISPONIBLE]", "NON PRIORITAIRE"],
    ["4.9 Culturel", "1/5", "[DONNÉE NON DISPONIBLE]", "NON PRIORITAIRE"],
])
doc.add_paragraph()

add_heading(doc, "4.1 Établissements scolaires", 2)
add_bullet(doc, "Parc : ≈75 500 établissements (80% publics MEB, ~259 élèves/étab. ; 20% privés, ~114 élèves/étab. ; 208 universités YÖK). Par niveau : pré-primaire ~40%, primaire ~24%, collège ~21%, lycée ~14%, université 1,5%. (Dossier interne Presto)")
add_bullet(doc, "Dynamique : grandes vagues de construction 2000-2020 (doublement salles de classe depuis 2002, 734 000) ; reconstruction post-séisme (budget 160 M$, 58% salles réparées, 42% en reconstruction ~9 800 salles).")
add_bullet(doc, "Ce qui se construit (neuf) : relocalisation hors zones de faille, bâtiments modulaires/préfabriqués standardisés, périphéries urbaines en croissance (Istanbul élargi, Ankara, Izmir, Bursa, Gaziantep, Konya).")
add_bullet(doc, "Ce qui se rénove : remplacement WC-lavabos, robinetterie temporisée, ajout points d'eau/savon/séchage, accessibilité PMR — rénovation par \"paquets standard\" (kits sanitaires-cloisons-plomberie identiques).")
add_bullet(doc, "Parc sanitaire inox : 3e secteur du marché inox sanitaire turc, 13-16% du parc de pièces, 120-150K pièces inox (85% standard 22 ans / 15% anti-vandale 16 ans). Neuf+réno estimé 11-13K pièces d'ici 2030.")

add_heading(doc, "4.2 Santé / Labo", 2)
add_bullet(doc, "Parc : ≈1 560 hôpitaux (60% publics, 35% privés, 5% universitaires). Âge moyen ≈13 ans, 80% reconstruit/rénové depuis 2002. Maillage solide : 1 hôpital référence par province + réseau secondaire.")
add_bullet(doc, "Dynamique à venir : inversion de tendance — construction/grands projets ≈50-65%, rénovation/réhabilitation ≈35-50%. Budget santé total 19 Mds€/an, investissements hospitaliers 1,5-3 Mds€/an.")
add_bullet(doc, "Ce qui se construit (neuf) : \"City Hospitals\" — redéveloppement urbain à grande échelle, gigantisme et centralisation (ex. Gaziantep, 1 875 lits, ouvert 02/2024). Montée des solutions collectives (batteries lavabos, urinoirs inox, électronique-thermostatique).")
add_bullet(doc, "Ce qui se rénove : remplacement \"à l'identique\" mais plus robuste — hausse anti-vandale (entretien), temporisés mécaniques, urinoirs et accessoires inox. Bloc opératoire/réa : inox très dominant (304/316), robinetterie sans-contact ou commande coude/genou, anti-légionellose.")
add_bullet(doc, "Parc sanitaire inox : 4e secteur, 11-14% du parc de pièces, 105-135K pièces inox (85% standard 19 ans / 15% anti-vandale 25 ans). Tourisme médical en progression → modernisation/internationalisation des infrastructures.")

add_heading(doc, "4.3 Bâtiments tertiaires", 2)
add_bullet(doc, "[DONNÉE NON DISPONIBLE — segment non couvert par le dossier de recherche interne ni par les recherches externes PREP. Recherche dédiée nécessaire en phase terrain (bureaux, cantines entreprises, crèches d'entreprise).]")

add_heading(doc, "4.4 Bâtiments industriels", 2)
add_bullet(doc, "[DONNÉE NON DISPONIBLE — segment non couvert. Le secteur industriel turc est important (8e producteur mondial d'inox, 36,9 Mt en 2024) mais aucune donnée de parc sanitaire industriel n'a été identifiée. (Dossier interne Presto — marché de l'inox)]")

add_heading(doc, "4.5 CHR (cafés, hôtels, restaurants)", 2)
add_bullet(doc, "[DONNÉE NON DISPONIBLE — non couvert par les sources internes ni les recherches PREP. À rechercher via données tourisme UNWTO et fédérations hôtelières turques (TÜROB) en phase terrain. Note qualitative BRG TR : le segment non-housing inclut hôtels/terminaux/restaurants/écoles, avec forte présence de mitigeurs thermostatiques, self-closing et électroniques dans les hôtels cinq étoiles.]")

add_heading(doc, "4.6 HPA (hôtellerie plein air)", 2)
add_bullet(doc, "[DONNÉE NON DISPONIBLE — non couvert par les sources disponibles.]")

add_heading(doc, "4.7 Centres Sport & Loisirs", 2)
add_bullet(doc, "Parc : ≈4 500 équipements (terrains foot ~30%, salles/gymnases ~20%, piscines ~15%, stades ~8%, centres jeunesse ~12%, pistes athlé ~1,5%, divers ~13%).")
add_bullet(doc, "Financement : 5-6 Mds€/an (2025), public 75-85% (État GSB, mairies, organismes publics), privé 15-25% (clubs professionnels, opérateurs privés). Parc récent ~40-45% vs ancien ~55-60%.")
add_bullet(doc, "Impact séisme 2023 : budget 28,27 M€, 29 équipements déjà livrés, 123 planifiés (86 équipements sportifs + 37 centres de jeunesse).")
add_bullet(doc, "Ce qui se construit (neuf) : driver majeur EURO 2032 (co-organisation Italie-Turquie), 5 stades turcs en lice (décision 10/2026), mise à niveau multi-villes aux normes UEFA (dont sanitaires/accueil).")
add_bullet(doc, "Ce qui se rénove : marché en volume tiré à 60-70% par la rénovation/remise à niveau — gros postes techniques (chauffage, ventilation, climatisation, plomberie). Le neuf pèse davantage en valeur (grands stades).")
add_bullet(doc, "Parc sanitaire inox : 5e secteur, 8-10% du parc de pièces, 100-130K pièces inox (65% standard 19 ans / 35% anti-vandale 16 ans).")

add_heading(doc, "4.8 Établissements à sécurité renforcée — anti-vandalisme", 2)
add_bullet(doc, "Parc : 407 établissements pénitentiaires (95 fermés hommes + 9 ouverts ; 11 fermés femmes + 8 ouverts ; 13 mineurs). 402 000 incarcérés (vs 500 000 pour l'UE27 entière). Taux d'occupation 132%.")
add_bullet(doc, "Dynamique récente : +80% de détenus depuis 2002 (62 000→220 000 en 2013, +8%/an, +10% sur 2025). 78 000 agents pénitentiaires (+62% en 8 ans, ratio 1/5 détenus). Exemple méga-prison de Silivri : 1M m², 11 000 détenus, 200% taux occupation.")
add_bullet(doc, "Dynamique à venir : 22 nouvelles prisons programmées (6 fin 2025, 9 en 2026, 5 en 2027, 2 en 2028) — construction-extension privilégiée vs réhabilitation lourde, mais ne rattrapant pas durablement le besoin carcéral.")
add_bullet(doc, "Point de vigilance commerciale : gouvernance jugée opaque (conditions d'attribution peu transparentes, concentration des marchés sur quelques attributaires proches du pouvoir, achats d'urgence liés à la surcapacité).")
add_bullet(doc, "Infrastructures sanitaires : cabines de douche/lavabos/toilettes accessibles depuis dortoirs/cellules, combinaisons modulaires inox (WC+lavabo+douche), robinetterie à ouverture-fermeture automatique antiblocage, boutons piézoélectriques, distributeurs sécurisés, design compact anti-vandalisme.")
add_bullet(doc, "1er secteur du marché inox sanitaire turc : 68-72% du parc sanitaire en inox (surtout lavabo + WC), parc estimé 200 000 pièces (95% inox anti-vandale 12 ans / 5% standard 15 ans). Rénovation ~20K pièces/an, neuf 3 500-5 000 pièces/an. Neuf+réno estimé 90-100K pièces d'ici 2030.")

add_heading(doc, "4.9 Bâtiments culturels", 2)
add_bullet(doc, "[DONNÉE NON DISPONIBLE — non couvert par les sources disponibles (salles de concert, théâtres, musées, cinémas).]")

add_heading(doc, "4.10 Lieux de culte", 2)
add_bullet(doc, "Parc : ~82 000 mosquées. Type : quartier/village 70-80% du parc (4-20 points d'eau, cœur du quartier), mosquées centrales/district 15-20% (prières du vendredi), grandes mosquées 5% (rôle politique/culturel/touristique, 40-100+ points d'eau, 1 tous les 300-500m en zone urbaine dense).")
add_bullet(doc, "Population pratiquante ~99% musulmane mais érosion de la pratique assidue — logique de maillage maintenue, priorisation maintenance/sécurité plutôt que capacité. Impact séisme 2023 : 3 800+ mosquées endommagées (34% du parc des 11 provinces touchées), 43% restant à réparer fin 2025.")
add_bullet(doc, "Acteurs financeurs : Diyanet (+3 M€/an, administre toutes mosquées non-historiques), TDV (fondation parapublique, construction-réparations), municipalités (permis/foncier), Vakıflar/VGM (mosquées historiques, logique de restauration).")
add_bullet(doc, "Infrastructures sanitaires : blocs ablutions (linéaire ou circulaire), espaces homme-femme séparés (souvent moins dimensionnés pour femmes). Ablutions ~1M points d'eau (½-¼ tour ~70%, mitigeurs 10-20%, temporisé méca. 5-15%, électronique 1-3%, eau froide) ; sanitaires ~240K points d'eau (mitigeurs ~40%, temporisé méca. ~30%, électronique ~5%).")
add_bullet(doc, "2e secteur du marché inox sanitaire : 14-18% du parc (surtout lavabo + urinoirs), 120-150K pièces inox (85% standard 23 ans / 15% anti-vandale 18 ans). Neuf+réno estimé 34-40K pièces d'ici 2030.")

add_heading(doc, "4.11 Transports", 2)
add_bullet(doc, "Secteur clé du développement national : investissements massifs (aéroports, compagnie aérienne, tourisme), ambition affirmée de \"hub géographique\" Europe-Asie. Plusieurs projets de construction/extension aéroportuaires en cours.")
add_bullet(doc, "Modernisation des infrastructures existantes vers des équipements \"green\" et \"technologiques\" — budget probablement en hausse mais non chiffré dans les sources disponibles.")
add_bullet(doc, "Routes/rail : forts programmes d'investissement en cours jusqu'à 2030, corridors internationaux, soutien indirect à la construction non-résidentielle et au développement péri-urbain.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — pas de chiffrage de parc sanitaire spécifique aux aéroports/gares trouvé. À approfondir en phase terrain auprès de la Direction Générale de l'Aviation Civile turque (DHMI).]")

add_heading(doc, "4.12 Opportunités Presto par segment — Synthèse", 2)
add_note(doc, "Classement par potentiel décroissant. Score 1 (très faible) à 5 (très fort).")
make_table(doc, [
    ["Segment", "Score", "Types de produits Presto", "Arguments clés", "Canal prioritaire"],
    ["4.8 Pénitentiaire", "5/5", "Inox anti-vandalisme, robinetterie encastrée, antiblocage piézoélectrique", "1er marché inox ERP turc, 22 nouvelles prisons programmées, indestructibilité", "Ministère Justice, CTE, marchés publics (vigilance gouvernance)"],
    ["4.2 Santé", "5/5", "Mitigeurs thermostatiques, sans-contact, PMR, inox 316 bloc opératoire", "City Hospitals PPP pluriannuel, tourisme médical, hygiène anti-légionellose", "Prescription BET santé, contractants PPP"],
    ["4.1 Éducation", "4/5", "Robinetterie temporisée, PMR, kits sanitaires standardisés", "Reconstruction post-séisme, rénovation \"paquets standard\", économie eau", "MEB, municipalités, appels d'offres reconstruction"],
    ["4.7 Sport & Loisirs", "3/5", "Temporisateurs douche/piscine, inox urinoirs", "EURO 2032, normes UEFA, rénovation majoritaire", "GSB, municipalités, comités organisateurs EURO 2032"],
    ["4.10 Lieux de culte", "3/5", "Robinetterie ablutions économe en eau, inox anti-vandale", "82 000 mosquées, reconstruction post-séisme, maillage dense", "Diyanet, TDV, municipalités"],
    ["4.11 Transports", "3/5", "Robinetterie robuste haute fréquentation, sans-contact", "Ambition hub Europe-Asie, modernisation \"green\"", "DHMI, opérateurs aéroportuaires"],
    ["4.3/4.4/4.5/4.6/4.9", "1-2/5", "Non documenté", "[DONNÉE NON DISPONIBLE]", "À explorer en phase terrain"],
], font_size=8)
doc.add_paragraph()
add_para(doc, "Priorités absolues : Pénitentiaire (5/5, sous réserve de vigilance gouvernance) · Santé (5/5) · Éducation (4/5). Stress hydrique = argument commercial transversal sur tous les segments.")
doc.add_paragraph()

# ─── PARTIE 5 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 5 — TAILLE MARCHÉ : ROBINETTERIE GÉNÉRALE", 1)

add_heading(doc, "5.1 Taille et valeur du marché robinetterie", 2)
add_note(doc, "Structure v8 stricte : deux estimations côte à côte. Estimation 1 = base AFISB France extrapolée avec le coefficient brut X. Estimation 2 = données directes BRG Turquie, sans extrapolation. L'ajustement structurel n'est appliqué qu'en Partie 6.")

add_heading(doc, "Constantes de référence et coefficient d'extrapolation", 2)
make_table(doc, [
    ["Variable", "Valeur", "Source"],
    ["PIB/hab France 2025", "48 982 USD", "Worldometer (constante MAB v8)"],
    ["Population France 2025", "69,1 M", "Worldometer (constante MAB v8)"],
    ["PIB/hab Turquie 2025", "18 611 USD (fourchette 18 200-18 611)", "IMF WEO 10/2025 / Worldometer"],
    ["Population Turquie 2025", "85,5 M", "Source interne / TurkStat"],
    ["Coefficient brut X", "(18 611 / 48 982) × (85,5 / 69,1) = 0,380 × 1,237 = 0,470", "Calcul MAB v8"],
    ["Taux de change indicatif", "≈51 TRY/EUR (2026) — dépréciation continue -27%/an", "ECB eurofxref 02/2026"],
])
doc.add_paragraph()

add_heading(doc, "Estimation 1 — base AFISB France (sections A1/B1)", 2)
add_para(doc, "Source AFISB utilisée : AFISB 2021 Rapport étude de marché annuel V6 (15/05/2021) + Tendance AFISB - depuis 2020 à 2025.xlsx. Le fichier Tendance AFISB confirme les familles de marché France 2019/2020 (lavabo, douche, urinoir, WC ; onglet \"Tendance marché FRA - 2020-2025\"). Aucune donnée directe Turquie AFISB n'est disponible ; application du coefficient brut X = 0,470.")
make_table(doc, [
    ["Section", "Segment", "Valeur base France", "Extrapolation Turquie"],
    ["Section A1", "Robinetterie de collectivités", "100-125 M€ (base AFISB France après division par 2)", "100-125 M€ × 0,470 = 47,0-58,8 M€"],
    ["Section B1", "Douches & équipements connexes", "52-65 M€ (base AFISB France après division par 2)", "52-65 M€ × 0,470 = 24,4-30,6 M€"],
    ["Section B1", "Chasses d'eau & WC collectifs", "90-110 M€ (base AFISB France après division par 2)", "90-110 M€ × 0,470 = 42,3-51,7 M€"],
    ["", "TOTAL Estimation 1", "242-300 M€", "113,7-141,1 M€"],
])
doc.add_paragraph()
add_note(doc, "Estimation par extrapolation depuis France — fiabilité moyenne. Les valeurs AFISB utilisées sont des bases internes France ; le coefficient X capte population et pouvoir d'achat relatif, mais ne capte pas les spécificités sectorielles turques.")

add_heading(doc, "Estimation 2 — base BRG Turquie (sections A2/B2)", 2)
add_para(doc, "Source : BRG Building Solutions, TR_Bathrooms_Full_Report_Apr21.pdf, avril 2021, données 2020. Données directes Turquie : aucune extrapolation France n'est appliquée. Valeurs MSP retenues en TRY et EUR.")
make_table(doc, [
    ["Section", "Segment BRG", "Volume (unités)", "Valeur MSP", "Valeur EUR"],
    ["Section A2", "Bath Taps and Mixers", "750 000", "182,15 M TRY", "22,66 M EUR"],
    ["Section A2", "Shower Taps and Mixers", "1 220 000", "272,67 M TRY", "33,92 M EUR"],
    ["Total Section A2", "Bath + Shower Taps and Mixers", "1 970 000", "454,82 M TRY", "56,58 M EUR"],
    ["Section B2", "Kitchen Taps and Mixers", "1 510 000", "228,47 M TRY", "28,42 M EUR"],
    ["Section B2", "Washbasin Taps and Mixers", "3 200 000", "527,14 M TRY", "65,57 M EUR"],
    ["Total Section B2", "Kitchen + Washbasin Taps and Mixers", "4 710 000", "755,61 M TRY", "93,99 M EUR"],
    ["", "Bidet Taps and Mixers", "96 000", "19,33 M TRY", "2,40 M EUR"],
    ["", "TOTAL BRG", "6 776 000", "1 229,75 M TRY", "152,96 M EUR"],
])
doc.add_paragraph()

doc.add_page_break()
add_heading(doc, "Tableau comparatif obligatoire A1/A2 et B1/B2", 2)
make_table(doc, [
    ["Comparaison", "Estimation 1 AFISB", "Estimation 2 BRG TR", "Écart et commentaire"],
    ["Section A1 ↔ Section A2", "A1 robinetterie de collectivités : 47,0-58,8 M€", "A2 Bath + Shower : 56,58 M€", "A2 se situe dans la fourchette A1 ; écart vs point médian A1 (52,9 M€) : +3,7 M€ (+7%). Périmètre proche mais non strictement identique : BRG inclut toutes applications, AFISB cible une base collectivité extrapolée."],
    ["Section B1 ↔ Section B2", "B1 douches + WC collectifs : 66,7-82,3 M€", "B2 Kitchen + Washbasin : 93,99 M€", "B2 dépasse la fourchette B1 de +11,7 à +27,3 M€ ; écart vs point médian B1 (74,5 M€) : +19,5 M€ (+26%). Écart expliqué par le poids résidentiel fort dans Kitchen/Washbasin BRG."],
])
doc.add_paragraph()
add_note(doc, "Le tableau \"Types ERP dominants dans Non-Housing\" n'est pas inclus dans l'étude principale conformément à CLAUDE.md v8. Les données brutes Non-Housing BRG restent documentées en annexes uniquement.")

add_heading(doc, "5.2 Spécificités produit du marché turc", 2)
add_bullet(doc, "Marché à 3 niveaux de prix : premium (Hansgrohe, Grohe, Ideal Standard — import), medium (Eczacıbaşı/ECA, Elginkan — locaux), économie (jusqu'à 40 fabricants locaux + import chinois). (BRG TR Apr21)")
add_bullet(doc, "Produit dominant : washbasin mixers one-head — cœur de marché, vendu majoritairement en segment économie/lower. Tendance à la baisse des baignoires au profit des douches (one-head mixers).")
add_bullet(doc, "Durée de vie produit nettement plus courte qu'en Europe (max 12 ans vs >20 ans) — qualité perçue inférieure, marché du remplacement (RMI) structurellement plus actif. (BRG TR Apr21)")
add_bullet(doc, "Marge distributeur ~45% du prix catalogue, jusqu'à 55% sur contrats annuels signés à l'avance avec les fabricants. (BRG TR Apr21)")
add_bullet(doc, "Tendance de fond : distribution en mutation rapide vers le DIY et l'e-commerce (Trendyol, Hepsiburada, n11, Koçtaş), exerçant une pression à la baisse sur les marges des canaux traditionnels (négoce, DIY physique).")

add_heading(doc, "5.3 Canaux de distribution", 2)
add_note(doc, "Hypothèse basée sur le modèle France adaptée aux spécificités turques (BRG TR + recherche PREP) — à confirmer terrain.")
make_table(doc, [
    ["Canal", "Part estimée", "Acteurs identifiés"],
    ["Grossistes BTP généralistes", "Majoritaire, non quantifié précisément", "Dikkaya Teknik Malzeme, Borpaş Plastik, Biryapi, Feyap — aucun spécialiste ERP/collectivités dédié identifié"],
    ["Wholesale/Export spécialisé bain", "Significatif", "ELEKS Foreign Trade (Elginkan — Europe/Moyen-Orient/CEI), ELMOR (Elginkan), INTEMA (Eczacıbaşı) — couverture Turquie"],
    ["DIY Retailers", "En croissance", "KOÇTAŞ (CA 180 M€, 171 dépôts, 2020), TEKZEN (CA 190 M€, 128 dépôts, Turquie/Roumanie/Irak, 2020), BAUHAUS (10 dépôts, Istanbul/Ankara/Bursa/Antalya)"],
    ["Réseaux SAV agréés nationaux", "Atout pour ERP", "Réseau \"Yetkili Servis\" VitrA/Artema — capillaire national, critère de poids en appels d'offres publics"],
    ["E-commerce / DIY en ligne", "Croissance forte attendue", "Trendyol, Hepsiburada, n11 — concerne principalement le résidentiel grand public"],
    ["Marchés publics", "Canal clé ERP", "EKAP (Elektronik Kamu Alımları Platformu) — accès direct non exploré en PREP, à approfondir"],
], font_size=8)
doc.add_paragraph()
add_bullet(doc, "[DONNÉE NON DISPONIBLE — aucun distributeur spécialisé \"ERP/collectivités\" dédié identifié, contrairement à des marchés plus matures (France, UK). Structure dominée par des grossistes généralistes multi-produits.]")

add_heading(doc, "5.4 Dynamique et perspectives 2025-2030", 2)
add_bullet(doc, "Marché taps & mixers en reprise attendue dès 2022 après le creux 2020-2021 (COVID + crise lire), soutenu par l'Urban Transformation Plan et l'organisme public du logement TOKİ. (BRG TR Apr21)")
add_bullet(doc, "Durée de vie produit courte (12 ans) = driver structurel du marché RMI, indépendamment du cycle de construction neuve — facteur de résilience pour les ventes de remplacement.")
add_bullet(doc, "Tendance vers l'électronique et le sans-contact attendue en poursuite, portée par les segments hôtels/terminaux/restaurants/écoles non-housing — mais sensible aux investissements publics et internationaux dans le non-résidentiel.")
add_bullet(doc, "Risque majeur : volatilité de la livre turque renchérissant mécaniquement les composants et matières premières importés (énergie notamment), pression continue sur les marges fabricants malgré la hausse des prix MSP en TRY.")
doc.add_paragraph()

# ─── PARTIE 6 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 6 — TAILLE MARCHÉ : ROBINETTERIE COLLECTIVE ERP", 1)

add_heading(doc, "6.1 Taille et valeur du marché robinetterie collective ERP", 2)
add_bullet(doc, "Donnée locale interne disponible : marché robinetterie ERP Turquie ~56 M€ (~0,64 €/habitant), ~4M pièces, TCAC estimé ~4% (2024-2030). Source : dossier interne propriétaire Presto ; méthode non détaillée, donc utilisée comme validation croisée et non comme méthode officielle.")
add_bullet(doc, "Méthode 1 officielle v8 (base AFISB/France, Section A1 après division par 2, coefficient X + ajustement structurel) : 42,3-52,9 M€.")
add_bullet(doc, "Méthode 2 officielle v8 (BRG TR × coefficient ERP interne 10% et 15%, puis ajustement structurel) : 13,8-20,7 M€.")
add_bullet(doc, "Fourchette finale retenue : 20,7-52,9 M€ — entre Méthode 2 haute et Méthode 1, conformément au protocole v8. L'estimation interne propriétaire (~56 M€) se situe légèrement au-dessus et valide plutôt la borne haute.")

add_heading(doc, "6.2 Méthode d'extrapolation — deux méthodes obligatoires", 2)
add_para(doc, "Formule de base : (PIB/hab Turquie / PIB/hab France) × (Population Turquie / Population France) = Coefficient X.")
make_table(doc, [
    ["Variable", "Calcul", "Résultat"],
    ["Ratio PIB/hab", "18 611 USD / 48 982 USD", "0,380"],
    ["Ratio population", "85,5 M / 69,1 M", "1,237"],
    ["Coefficient X brut", "0,380 × 1,237", "0,470"],
    ["Ajustement structurel", "0,470 × 0,90", "0,423"],
])
doc.add_paragraph()

add_para(doc, "Méthode 1 — base AFISB/France (Section A1 robinetterie de collectivités)")
make_table(doc, [
    ["Étape", "Calcul", "Résultat"],
    ["Base France Section A1", "Robinetterie de collectivités AFISB après division par 2", "100-125 M€"],
    ["Application coefficient X brut", "100-125 M€ × 0,470", "47,0-58,8 M€"],
    ["Ajustement structurel", "47,0-58,8 M€ × 0,90", "42,3-52,9 M€"],
])
doc.add_paragraph()
add_note(doc, "Ajustement structurel -10% : économie informelle significative (réduit la part traçable du marché formel), marché extrêmement price-driven et dominé par des acteurs locaux/économie, instabilité TRY/EUR ; partiellement compensés par un cycle de remplacement court (max 12 ans vs >20 ans en Europe, BRG TR Apr21). Estimation par extrapolation avec ajustement structurel — fiabilité moyenne.")

add_para(doc, "Méthode 2 — base BRG Turquie (marché total pays × coefficient ERP interne 10% et 15%)")
make_table(doc, [
    ["Étape", "Calcul", "Résultat"],
    ["Base BRG Turquie", "Total taps & mixers TR 2020 (BRG TR Apr21, p.101)", "152,96 M€"],
    ["Coefficient ERP basse", "152,96 M€ × 10%", "15,30 M€"],
    ["Coefficient ERP haute", "152,96 M€ × 15%", "22,94 M€"],
    ["Ajustement structurel basse", "15,30 M€ × 0,90", "13,77 M€"],
    ["Ajustement structurel haute", "22,94 M€ × 0,90", "20,65 M€"],
])
doc.add_paragraph()
add_note(doc, "Conformément à v8, la part Non-Housing volume BRG n'est pas utilisée pour la Méthode 2 ; elle est remplacée par le coefficient ERP interne Presto 10-15%.")

make_table(doc, [
    ["", "Méthode 1 (base AFISB/France)", "Méthode 2 basse (BRG × 10%)", "Méthode 2 haute (BRG × 15%)"],
    ["Base utilisée", "Section A1 AFISB France : 100-125 M€", "BRG TR total : 152,96 M€", "BRG TR total : 152,96 M€"],
    ["Coefficient X", "0,470 brut", "10%", "15%"],
    ["Ajustement structurel", "-10% (coefficient ajusté 0,423)", "-10%", "-10%"],
    ["Estimation marché ERP", "42,3-52,9 M€", "13,77 M€", "20,65 M€"],
    ["Niveau de confiance", "Moyen : extrapolation France structurée", "Moyen : BRG direct mais donnée 2020 pré-crise TRY", "Moyen : BRG direct mais coefficient ERP interne à valider terrain"],
], font_size=8)
doc.add_paragraph()
add_note(doc, "Fourchette finale retenue : 20,65-52,9 M€. Limites systématiques : la méthode ne capte pas la part d'économie informelle ; elle ne reflète pas toutes les spécificités sectorielles locales ; la volatilité des taux de change peut fausser la comparaison ; l'ajustement structurel repose sur des hypothèses à confirmer terrain ; le coefficient ERP 10-15% est une estimation interne Presto — à valider terrain.")

add_heading(doc, "6.3 Évaluation du potentiel par segment ERP — Scoring", 2)
make_table(doc, [
    ["Segment", "Score", "Justification", "Hypothèses clés"],
    ["4.1 Éducation", "4/5", "75 500 établissements ; reconstruction post-séisme (9 800 salles) ; rénovation \"paquets standard\"", "Budget éducation en baisse relative (12,9%→10,6% du budget public, 2015-2022)"],
    ["4.2 Santé", "5/5", "1 560 hôpitaux ; City Hospitals PPP 105 Mds TRY/an ; tourisme médical en hausse", "Pipeline PPP pluriannuel sécurisé jusqu'en 2030+"],
    ["4.3 Tertiaire", "2/5", "[DONNÉE NON DISPONIBLE]", "Score par défaut, à valider terrain"],
    ["4.4 Industriel", "2/5", "[DONNÉE NON DISPONIBLE] — 8e producteur mondial inox (36,9 Mt 2024), segment sanitaire minime (~12-43 M$)", "Score par défaut, à valider terrain"],
    ["4.5 CHR", "2/5", "[DONNÉE NON DISPONIBLE] — present qualitativement dans le non-housing BRG (hôtels 5 étoiles, thermostatiques)", "Score par défaut, à valider terrain via UNWTO/TÜROB"],
    ["4.6 HPA", "1/5", "[DONNÉE NON DISPONIBLE]", "Score par défaut"],
    ["4.7 Sport & Loisirs", "3/5", "4 500 équipements ; EURO 2032 co-organisation ; financement 5-6 Mds€/an", "Décision finale stades 10/2026"],
    ["4.8 Pénitentiaire", "5/5", "407 établissements ; 22 nouvelles prisons 2025-2028 ; 1er marché inox sanitaire turc (200K pièces)", "Gouvernance opaque sur attribution marchés — risque commercial à intégrer"],
    ["4.9 Culturel", "1/5", "[DONNÉE NON DISPONIBLE]", "Score par défaut"],
    ["4.10 Lieux de culte", "3/5", "82 000 mosquées ; reconstruction post-séisme (3 800+ endommagées) ; financement Diyanet/TDV", "Érosion pratique religieuse assidue — maillage maintenu malgré tout"],
    ["4.11 Transports", "3/5", "Investissements aéroportuaires massifs ; ambition hub Europe-Asie", "Chiffrage parc sanitaire non disponible"],
], font_size=8)
doc.add_paragraph()
add_para(doc, "Top 3 : Pénitentiaire (5/5) · Santé (5/5) · Éducation (4/5)")

add_heading(doc, "6.4 Spécificités produit robinetterie collective turque", 2)
add_bullet(doc, "Robinetterie ERP par type de produit (estimation interne) : mitigeurs ~55-60% (ECA, Artema, VitrA + marques locales), poussoir-temporisé ~24-28% (Artema AquaTouch, ECA, GPD + Roca, Delabie, Schell), électronique-capteur ~6-9% (ECA, Artema, VitrA, Newarc, Creavit, Kale + Grohe, Hansgrohe, Ideal Standard, Geberit), autres ~5-8%.")
add_bullet(doc, "Sur le segment poussoir-temporisé spécifiquement : acteurs turcs (ECA, Artema-VitrA, GPD, Creavit, Kale) ~75-85% / acteurs européens (Roca, Grohe, Schell, Delabie) ~15-25%.")
add_bullet(doc, "Inox anti-vandalisme quasi-systématique en pénitentiaire (95% du parc, durée de vie 12 ans) ; présence croissante en santé (blocs opératoires, inox 316) et lieux de culte (urinoirs).")
add_bullet(doc, "Robinetterie ablutions (lieux de culte) à dominante mécanique simple (½-¼ tour ~70%) — opportunité de montée en gamme vers mitigeurs/temporisé pour économie d'eau.")
add_bullet(doc, "Économiseurs d'eau / limiteurs de débit : argument commercial fort et transversal compte tenu du stress hydrique structurel (Ankara, Izmir).")

add_heading(doc, "6.5 Dynamique et perspectives 2025-2030", 2)
add_bullet(doc, "Pipeline pénitentiaire sécurisé : 22 nouvelles prisons programmées 2025-2028 — visibilité de commandes sur 4 ans, sous réserve de vigilance sur la gouvernance des attributions.")
add_bullet(doc, "City Hospitals : flux récurrent garanti par le modèle PPP (~10% budget santé alloué), nouveaux établissements et maintenance des plateaux techniques en continu.")
add_bullet(doc, "Reconstruction post-séisme (éducation + santé + religieux) : fenêtre d'opportunité concentrée mais à plus court terme (objectif 2025-2028 pour la majorité des travaux).")
add_bullet(doc, "EURO 2032 : catalyseur ponctuel mais significatif pour le segment sport (5 stades turcs, décision finale 10/2026) — fenêtre de prescription 2026-2032.")
add_bullet(doc, "Risques : volatilité TRY pesant sur le pouvoir d'achat des donneurs d'ordre publics pour les équipements importés ; dominance locale très forte (60-70% volume) rendant la pénétration commerciale plus lente que dans des marchés moins price-driven ; absence de comparable Delabie sur le marché — pas de précédent direct à analyser pour calibrer une stratégie d'entrée.")
doc.add_paragraph()

# ─── PARTIE 7 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 7 — CONCURRENTS", 1)

add_heading(doc, "7.1 DELABIE (analyse prioritaire)", 2)
add_bullet(doc, "AUCUNE PRÉSENCE CONFIRMÉE EN TURQUIE. Le groupe revendique une présence dans plus de 90 pays via 9 filiales/bureaux commerciaux : Royaume-Uni, Benelux, Allemagne, Autriche, Pologne, Espagne, Portugal, Émirats Arabes Unis (Dubaï), Hong Kong — la Turquie n'apparaît pas dans cette liste. (Delabie.com, consulté juin 2026)")
add_bullet(doc, "Présence régionale la plus proche : Émirats Arabes Unis via le distributeur local MGK (Dubaï/Abu Dhabi) — aucune preuve que ce canal couvre le marché turc. (MGK.ae, consulté juin 2026)")
add_bullet(doc, "Dans le dossier de recherche interne Presto, Delabie n'apparaît qu'une seule fois, comme acteur européen marginal (~15-25% volume) sur le segment poussoir-temporisé et électronique-capteur, sans mention de présence locale, d'appel d'offres ou de référencement.")
bold_bullet(doc, "Implication stratégique", "Delabie semble absent ou très marginal en Turquie — fenêtre d'opportunité potentielle pour Presto sur le segment premium importé, mais aussi signal que le marché turc ERP est probablement jugé peu prioritaire ou difficile d'accès par les acteurs européens haut de gamme (faible volume relatif, sensibilité prix extrême, marques locales dominantes).")
add_bullet(doc, "Aucun appel d'offres public turc (hôpitaux, écoles, prisons) mentionnant Delabie identifié dans les recherches PREP — à reconfirmer via contact direct export Delabie plutôt que recherche web seule.")

add_heading(doc, "7.2 Profils des acteurs turcs dominants", 2)
make_table(doc, [
    ["Acteur", "Positionnement", "Présence ERP", "Canaux", "Forces / Faiblesses"],
    ["Eczacıbaşı / VitrA / Artema", "Leader historique turc (depuis 1958/1983), CA groupe holding 33 Md TRY (2023, toutes activités), >13 500 employés, 40 sites production", "Mitigeurs 55-60% ERP, poussoir-temporisé gamme \"AquaTouch\"", "Réseau SAV national \"Yetkili Servis\", export ~80% CA (Eczacıbaşı Yapı Ürünleri)", "F: intégration verticale céramique+robinetterie / f: CA robinetterie ERP non isolé publiquement"],
    ["ECA (Eczacıbaşı)", "Marque volume/mid-market du groupe", "Mitigeurs + poussoir-temporisé urinoir (5-30s)", "[DONNÉE NON DISPONIBLE]", "F: couverture SAV nationale, prix compétitif / f: peu de visibilité projets ERP nommés"],
    ["GPD", "Mid-market, positionnement marketing hygiène/santé", "Auto-déclaré hôpitaux/hôtels/restaurants (sans preuve projet)", "Vente directe + revendeurs (Nalburdayım.com)", "F: image hygiène / f: aucune donnée financière publique"],
    ["Creavit", "Programme transformation durable soutenu par le ministère du Commerce", "Électronique-capteur (avec ECA/Artema/VitrA/Newarc/Kale)", "[DONNÉE NON DISPONIBLE]", "[DONNÉE NON DISPONIBLE]"],
    ["Kale (Kale Musluk)", "Groupe industriel (1957, 17 sociétés, >5000 employés groupe), Kale Musluk fondée 1969", "Poussoir-temporisé + électronique-capteur", "[DONNÉE NON DISPONIBLE]", "F: ancienneté/savoir-faire forge laiton / f: vulnérabilité change/import depuis crise lire 2019-2020, perte de compétitivité confirmée"],
    ["ROCA (via NSK)", "Entrée 2016 par acquisition (46 M€, usine Eskişehir, 300 employés, capacité 8M pièces/an)", "[DONNÉE NON DISPONIBLE]", "[DONNÉE NON DISPONIBLE]", "Stratégie d'implantation par rachat industriel local plutôt qu'import"],
    ["Grohe / Hansgrohe", "Premium import, dominent le haut de gamme", "Projets prescripteurs (hôtels, aéroports, sièges)", "Sites locaux directs confirmés (groheturkiye.com, hansgrohe.com.tr)", "Présence directe structurée, contrairement à Delabie"],
    ["Geberit", "Premium", "Catalogue robinetterie photosensible référencé localement (geberit.com.tr)", "Présence directe confirmée", "—"],
], font_size=8)
doc.add_paragraph()
add_note(doc, "[DONNÉE NON DISPONIBLE — CA/effectifs spécifiques \"robinetterie ERP\" par marque non isolés du CA groupe en sources ouvertes — à creuser via rapports annuels Eczacıbaşı Yapı Ürünleri ou bases sectorielles Statista Turquie]")

add_heading(doc, "7.3 Distribution et accès marché", 2)
add_bullet(doc, "Grossistes BTP généralistes fragmentés (Dikkaya Teknik Malzeme, Borpaş Plastik, Biryapi, Feyap) — aucun spécialiste ERP/collectivités identifié.")
add_bullet(doc, "Wholesale/Export spécialisé bain : ELEKS Foreign Trade et ELMOR (Elginkan), INTEMA (Eczacıbaşı) — couverture Turquie + Europe/Moyen-Orient/CEI pour ELEKS. (BRG TR Apr21, p.25)")
add_bullet(doc, "DIY Retailers en forte croissance : KOÇTAŞ (180 M€ CA, 171 dépôts), TEKZEN (190 M€ CA, 128 dépôts, présent aussi Roumanie/Irak), BAUHAUS (10 dépôts grandes villes). (BRG TR Apr21, p.25)")
add_bullet(doc, "Réseaux SAV propriétaires des grands locaux (VitrA/Artema \"Yetkili Servis\") — atout pour appels d'offres publics (critère de continuité de service).")
add_bullet(doc, "Aucun distributeur spécialisé \"ERP/collectivités\" dédié identifié, contrairement à des marchés plus matures (France, UK, Suisse).")

add_heading(doc, "7.4 Opportunités de différenciation pour Presto", 2)
add_bullet(doc, "Robinetterie temporisée premium : différenciation possible sur fiabilité mécanique et durée de vie face aux gammes locales volume (ECA, GPD), en ciblant les segments à forte fréquentation (prisons, sport, lieux de culte — taux d'équipement inox déjà élevé).")
add_bullet(doc, "Anti-vandalisme inox haut de gamme : opportunité forte sur établissements à sécurité renforcée où le parc inox est déjà significatif (68-72% pénitentiaire) et où la durabilité prime sur le prix d'achat initial.")
add_bullet(doc, "Robinetterie PMR : axe de différenciation peu investi par les acteurs turcs identifiés (aucune mention PMR chez ECA/Artema/GPD/Creavit/Kale) — opportunité sur la conformité accessibilité, notamment bâtiments publics neufs post-séisme.")
add_bullet(doc, "Économiseurs d'eau / limiteurs de débit : argument fort dans le contexte de stress hydrique sévère documenté — levier commercial et institutionnel auprès des gestionnaires d'ERP publics.")
add_bullet(doc, "Robinetterie sans contact/infrarouge : segment encore minoritaire (~6-9%) mais en croissance, où les marques européennes (Grohe, Hansgrohe, Ideal Standard, Geberit) sont déjà présentes aux côtés des locaux — fenêtre d'entrée sur l'argument hygiène.")
add_bullet(doc, "Point de vigilance transversal : sensibilité prix très forte du marché turc, acteurs locaux couvrant tout le spectre éco→premium, absence de référence Delabie pouvant indiquer une difficulté structurelle d'accès non encore identifiée (barrières douanières/réglementaires, structure des appels d'offres favorisant le contenu local).")
doc.add_paragraph()

# ─── PARTIE 8 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 8 — NORMES & CERTIFICATIONS ROBINETTERIE", 1)

add_heading(doc, "8.1 Cadre général d'accès au marché", 2)
add_bullet(doc, "Union douanière UE-Turquie (1996) : exonération des droits de douane sur produits industriels européens — 0% sur la robinetterie (HS 8481 équivalent), MAIS contrôles techniques maintenus intégralement.")
add_bullet(doc, "Marquage CE insuffisant seul : contrôle douanier systématique via le système TAREKS (contrôle de risque) opéré par le TSE — risque de blocage/retour en cas de non-conformité documentaire.")
add_bullet(doc, "Conformité substances : RoHS (équipements électriques/électroniques, robinetterie électronique) + KKDIK/REACH (substances chimiques, obligations d'enregistrement).")
add_bullet(doc, "ISO 9001 (qualité) et ISO 14001 (environnement) attendus comme preuves de robustesse industrielle et de conformité, en complément des normes produit.")

add_heading(doc, "8.2 Normes TS EN applicables", 2)
make_table(doc, [
    ["Norme TS EN", "Équivalent EN", "Objet", "Implications techniques"],
    ["TS EN 200", "EN 200", "Robinets/mitigeurs sanitaires — exigences générales", "Débit nominal, pression de service, étanchéité, résistance mécanique, matériaux (laiton chromé/copolymère acétal)"],
    ["TS EN 817", "EN 817", "Robinetterie monocommande (mitigeurs mécaniques)", "Tests d'endurance mécanique du levier"],
    ["TS EN 1111", "EN 1111", "Robinetterie thermostatique PN10", "Stabilité thermique, anti-brûlure, résistance corrosion — certificat VitrA confirmé A-TSE-1111-03"],
    ["TS EN 15091", "EN 15091", "Robinetterie électronique (\"fotoselli\")", "[DONNÉE NON DISPONIBLE — détail technique non trouvé, texte payant TSE]"],
    ["TS EN 816", "EN 816 (≈ EN 12541 FR)", "Fermeture automatique temporisée — CŒUR DE GAMME Presto", "[DONNÉE NON DISPONIBLE — détail technique non trouvé, texte payant TSE]"],
    ["TS EN 1112/1113", "EN 1112/1113", "Douchettes et flexibles", "Débit, résistance pression, durabilité du flexible"],
    ["TS EN 274-1/2/3", "EN 274", "Siphons/raccords évacuation", "Résistance ≥80°C, résistance acides, démontable/nettoyable"],
    ["TS EN 248", "EN 248", "Essais corrosion revêtements métalliques", "Tests de résistance du chromage"],
    ["TS EN ISO 3822", "EN ISO 3822", "Mesure acoustique robinetterie", "Tests de bruit en laboratoire"],
], font_size=8)
doc.add_paragraph()
add_note(doc, "Pas de norme turque \"anti-vandalisme\" dédiée identifiée — relève du cahier des charges (matériaux/fixations) plutôt que d'une TS EN spécifique. Pas d'équivalent identifié pour EN 13904 (siphons de sol). À reconfirmer auprès du TSE directement.")

add_heading(doc, "8.3 Organismes et procédure de certification", 2)
add_bullet(doc, "TSE (Türk Standartları Enstitüsü) : organisme central. Procédure — dépôt dossier → détermination norme TS applicable → audit préliminaire → inspection site (machines, contrôle qualité, personnel) → certificat délivré par référence à la norme.")
add_bullet(doc, "Coûts et délais précis non standardisés (montants variables selon secteur/taille entreprise, non publiés). [DONNÉE NON DISPONIBLE — grille tarifaire exacte]")
add_bullet(doc, "Articulation confirmée : marquage CE + contrôle TSE via système TAREKS à l'importation. Pour les produits réglementés sous régime CE (notamment électronique RoHS), une Déclaration de Conformité CE est exigée en complément lors des contrôles TSE.")

add_heading(doc, "8.4 Écarts avec les normes françaises/européennes", 2)
add_bullet(doc, "Socle normatif quasi identique aux normes EN françaises (transpositions directes) — barrière technique théoriquement limitée pour Presto.")
add_bullet(doc, "Écart principal : procédure administrative additionnelle (TSE + marquage CE), pas de divergence technique de fond identifiée dans les normes elles-mêmes.")
add_bullet(doc, "Aucune dérogation locale identifiée. [À confirmer — recherche complémentaire nécessaire auprès du TSE directement ou d'un bureau de certification local]")

add_heading(doc, "8.5 Contraintes d'entrée et délais pratiques", 2)
add_bullet(doc, "Documentation additionnelle fréquente exigée par les douanes turques : rapports de test originaux, traductions notariées, même pour produits déjà marqués CE — source de retards récurrents signalés.")
add_bullet(doc, "Barrière linguistique administrative : documentation technique et échanges avec le TSE/douanes en turc — nécessité d'un partenaire/agent local pour la traduction notariée des dossiers.")
add_bullet(doc, "Liste annuelle turque de normes d'importation obligatoires publiée par le gouvernement — les produits non couverts par les directives UE doivent être testés/certifiés TSE avant import, pouvant concerner certains composants de robinetterie selon leur classification douanière.")
doc.add_paragraph()

# ─── PARTIE 9 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 9 — POINTS À REVÉRIFIER", 1)
add_para(doc, "Données incertaines ou manquantes à revalider par recherche complémentaire ou terrain :", italic=True)
doc.add_paragraph()

bold_bullet(doc, "Données BRG TR datées de 2020", "Le rapport BRG TR Apr21 précède la crise inflationniste et la dépréciation massive de la livre turque (2021-2023). Les valeurs EUR ont été retenues pour comparabilité mais une mise à jour BRG post-2021 est fortement recommandée. Contact : europe@brgbuildingsolutions.com")
bold_bullet(doc, "PIB/hab Turquie 2025", "Fourchette 18 200-18 611 USD selon source. Choisir une référence unique (IMF WEO octobre 2025 recommandé) pour fiabiliser le coefficient d'extrapolation en phase terrain.")
bold_bullet(doc, "Taille du marché construction total", "Écart de 53 à 174 Mds$ selon cabinet (Mordor, IMARC, Verified Market Research, Research and Markets) — à trancher via une source TurkStat officielle unique convertie au taux de change moyen annuel.")
bold_bullet(doc, "Ajustement structurel -10% (Parties 5 et 6)", "Hypothèse basée sur l'économie informelle et la sensibilité prix du marché turc, partiellement compensée par le cycle de remplacement court. À confirmer via interviews terrain (distributeurs locaux, Eczacıbaşı/VitrA).")
bold_bullet(doc, "Présence Delabie en Turquie", "Absence confirmée par recherche web croisée, mais à reconfirmer définitivement via contact direct export Delabie plutôt que recherche web seule.")
bold_bullet(doc, "CA/effectifs robinetterie ERP par marque turque (ECA, Artema, GPD, Creavit, Kale)", "Non isolés du CA groupe en sources ouvertes — à rechercher via rapports annuels Eczacıbaşı Yapı Ürünleri ou bases sectorielles Statista Turquie.")
bold_bullet(doc, "Gouvernance des marchés publics pénitentiaires", "Signalée comme potentiellement opaque (concentration sur quelques attributaires) — à vérifier avant tout engagement commercial sur ce segment malgré son potentiel élevé (5/5).")
bold_bullet(doc, "Segments 4.3 (tertiaire), 4.4 (industriel), 4.5 (CHR), 4.6 (HPA), 4.9 (culturel)", "Non couverts par le dossier interne ni les recherches PREP — à rechercher intégralement en phase terrain (ministères sectoriels, fédérations professionnelles turques).")
bold_bullet(doc, "Budget précis volet écoles/hôpitaux du plan de reconstruction post-séisme", "Non isolé de l'enveloppe globale 45-80 Mds$ — à demander directement auprès de la Banque mondiale ou du gouvernement turc (AFAD).")
bold_bullet(doc, "Appels d'offres EKAP nominatifs", "Aucun accessible en recherche PREP (plateforme nécessite navigation interactive directe sur ekap.kik.gov.tr) — recherche dédiée à mener en phase terrain.")
bold_bullet(doc, "Détail technique complet TS EN 15091 (électronique) et TS EN 816 (temporisé)", "Textes de normes payants chez TSE, non accessibles en recherche ouverte — à acquérir directement.")
bold_bullet(doc, "Présence Presto en Turquie", "Coordonnées exactes de représentants commerciaux éventuels, clients actuels, produits déjà référencés, distributeurs partenaires locaux.")
doc.add_paragraph()

# Sauvegarde
path_etude = os.path.join(OUTPUT_DIR, "MAB_Turquie_EtudeV2.docx")
doc.save(path_etude)
print(f"✓ Étude v2 sauvegardée : {path_etude}")


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT ANNEXES
# ════════════════════════════════════════════════════════════════════════════
ann = Document()
set_margins(ann)

t2 = ann.add_heading("MAB TURQUIE — ANNEXES & SOURCES v1", 0)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2 = ann.add_paragraph("Sources complètes, données brutes et compléments — Les Robinets Presto")
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.runs[0].font.size = Pt(11)
sub2.runs[0].font.name = FONT
ann.add_paragraph()

add_heading(ann, "ANNEXE 1 — LISTE DES SOURCES UTILISÉES", 1)
ann.add_paragraph("Toutes les sources consultées pour MAB_Turquie_Etude.docx v1 :").runs[0].font.italic = True
ann.add_paragraph()

sources = [
    ("Direction Évaluation, Études et Prospective — Fiche pays Turquie", "Source interne", "Français", "Juillet 2024"),
    ("Dossier Turquie PPT 2 - FINAL — recherche interne propriétaire", "Source interne", "Français", "2025-2026"),
    ("Statista — Construction industry in Turkey", "Source interne (study_id56671)", "Anglais", "2024"),
    ("BRG Building Solutions — TR Bathrooms Full Report", "Source interne (TR_Bathrooms_Full_Report_Apr21.pdf)", "Anglais", "Avril 2021 (données 2020)"),
    ("IMF World Economic Outlook — PIB/hab Turquie", "imf.org / StatisticsTimes", "Anglais", "Octobre 2025"),
    ("TurkStat — Population Turquie", "turkstat.gov.tr (via Turkish Minute)", "Anglais", "Février 2026"),
    ("World Bank — Türkiye Earthquake Recovery and Reconstruction Project", "worldbank.org", "Anglais", "Juin 2023"),
    ("World Bank/GFDRR — Estimation coût reconstruction post-séisme", "worldbank.org", "Anglais", "2023"),
    ("bne IntelliNews — Budget City Hospitals PPP 2025", "intellinews.com", "Anglais", "2025-2026"),
    ("invest.gov.tr — City Hospitals Program", "invest.gov.tr", "Anglais", "2024"),
    ("decidehealth.world — Gaziantep City Hospital", "decidehealth.world", "Anglais", "2024"),
    ("Daily Sabah/AGBI — TOKİ 500 000 logements 2025", "agbi.com", "Anglais", "Octobre 2025"),
    ("AGBI — Budget investissement public 2024 (éducation, projets)", "agbi.com", "Anglais", "Janvier 2024"),
    ("ab.gov.tr — Union douanière UE-Turquie", "ab.gov.tr", "Turc/Anglais", "Consulté 2026"),
    ("Parlement européen — Union douanière UE-Turquie", "europarl.europa.eu", "Français", "2021"),
    ("JJRLAB — Barrières import / système TAREKS", "jjrlab.com", "Anglais", "2024-2025"),
    ("trade.gov — Turkey Standards for Trade", "trade.gov", "Anglais", "Consulté 2026"),
    ("ECB eurofxref — Taux de change EUR/TRY", "ecb.europa.eu", "Anglais", "Février 2026"),
    ("Capital.com — Dépréciation livre turque", "capital.com", "Anglais", "Février 2026"),
    ("Trading Economics — Inflation, taux directeur, salaires Turquie", "tradingeconomics.com", "Anglais", "2024-2025"),
    ("Mordor Intelligence — Turkey Construction Market", "mordorintelligence.com", "Anglais", "2025"),
    ("IMARC Group — Turkey Construction Market", "imarcgroup.com", "Anglais", "2024"),
    ("Verified Market Research — Turkey Construction Market", "verifiedmarketresearch.com", "Anglais", "2024"),
    ("GlobeNewswire/Research and Markets — Turkey Construction Industry Report", "globenewswire.com", "Anglais", "Avril et août 2025"),
    ("Euroconstruct — Répartition résidentiel/non-résidentiel Europe", "euroconstruct.org", "Anglais", "2024-2025"),
    ("OCDE — Education at a Glance", "oecd.org", "Anglais", "2025"),
    ("Türkiye Today — Coûts matériaux et main d'œuvre construction", "turkiyetoday.com", "Anglais", "2025"),
    ("Delabie.com — Filiales internationales", "delabie.com", "Français/Anglais", "Consulté 2026"),
    ("MGK.ae — Distributeur Delabie UAE", "mgk.ae", "Anglais", "Consulté 2026"),
    ("Eczacıbaşı Holding / VitrA Global — Données groupe", "eczacibasi.com.tr / vitraglobal.com", "Anglais/Turc", "Consulté 2026"),
    ("GPD.com.tr — Site fabricant", "gpd.com.tr", "Turc", "Consulté 2026"),
    ("Creavit.com.tr — Communiqué Green Agreement", "creavit.com.tr", "Turc", "Consulté 2026"),
    ("Kale Group / Kalepres.com — Données entreprise", "kalepres.com", "Turc/Anglais", "Consulté 2026"),
    ("Grohe Türkiye / Hansgrohe Türkiye / Geberit Türkiye — Sites officiels", "groheturkiye.com / hansgrohe.com.tr / geberit.com.tr", "Turc", "2026"),
    ("Coşan Ticaret — Référentiel normes TS EN robinetterie", "Site fournisseur turc", "Turc", "Consulté 2026"),
    ("VitrA — Certificat conformité TS EN 1111", "vitra.com.tr", "Turc", "Consulté 2026"),
    ("webdosya.csb.gov.tr — Bordereaux de prix Ministère Environnement/Urbanisme", "csb.gov.tr", "Turc", "2020"),
    ("TSE (Türk Standartları Enstitüsü) — Procédure de certification", "tse.org.tr / isonedir.com", "Turc", "2024-2025"),
    ("EKAP — Plateforme marchés publics turcs", "ekap.kik.gov.tr", "Turc", "Consulté 2026 (accès limité)"),
]
rows = [["Source", "URL / Référence", "Langue", "Date"]] + [list(s) for s in sources]
make_table(ann, rows, font_size=8)
ann.add_paragraph()

add_heading(ann, "ANNEXE 2 — DONNÉES DE MARCHÉ ROBINETTERIE TURQUIE : ESTIMATIONS DÉTAILLÉES", 1)
add_para(ann, "Détail complet du tableau BRG TR (TR_Bathrooms_Full_Report_Apr21.pdf, p.101) — Valeurs MSP 2020 et 2019 :")
make_table(ann, [
    ["Produit", "Volumes 2020", "MSP Local 2020 (TRY)", "Valeur 2020 (M TRY)", "Valeur 2020 (M EUR)", "Volumes 2019", "Valeur 2019 (M EUR)"],
    ["Bath Taps and Mixers", "750 000", "242,87", "182,15", "22,66", "810 000", "27,13"],
    ["Bidet Taps and Mixers", "96 000", "201,31", "19,33", "2,40", "104 000", "2,89"],
    ["Kitchen Taps and Mixers", "1 510 000", "181,30", "228,47", "28,42", "1 612 000", "33,66"],
    ["Shower Taps and Mixers", "1 220 000", "223,50", "272,67", "33,92", "1 290 000", "39,78"],
    ["Washbasin Taps and Mixers", "3 200 000", "164,73", "527,14", "65,57", "3 430 000", "77,90"],
    ["GRAND TOTAL", "6 776 000", "181,49", "1 229,75", "152,96", "7 246 000", "181,36"],
], font_size=8)
ann.add_paragraph()
add_para(ann, "Détail End-Use 2020 (BRG TR, p.108) — volumes en milliers de pièces :")
make_table(ann, [
    ["Produit", "Housing RMI", "New Housing", "Non-Housing", "Total 2020"],
    ["Bath Taps and Mixers", "483,20 (64,43%)", "213,90 (28,52%)", "52,91 (7,05%)", "750,00"],
    ["Bidet Taps and Mixers", "73,37 (76,43%)", "19,06 (19,85%)", "3,57 (3,72%)", "96,00"],
    ["Kitchen Taps and Mixers", "764,48 (50,63%)", "700,34 (46,38%)", "45,18 (2,99%)", "1 510,00"],
    ["Shower Taps and Mixers", "537,49 (44,06%)", "546,00 (44,75%)", "136,51 (11,19%)", "1 220,00"],
    ["Washbasin Taps and Mixers", "1 782,59 (55,71%)", "1 173,57 (36,67%)", "243,85 (7,62%)", "3 200,00"],
    ["GRAND TOTAL", "3 641,13 (53,74%)", "2 652,87 (39,15%)", "482,01 (7,11%)", "6 776,00"],
], font_size=8)
ann.add_paragraph()
add_para(ann, "Répartition par type de produit (BRG TR, p.110) — éléments identifiés :")
add_bullet(ann, "Electronic : 27,12K unités (0,40% du marché total), intégralement sur Washbasin Taps and Mixers.")
add_bullet(ann, "Self-Closing : 195,22K unités (2,88% du marché total), intégralement sur Shower Taps and Mixers.")
add_bullet(ann, "[DONNÉE NON DISPONIBLE — détail complet One Head/Pillar/Thermostatic/Two Head non extrait avec précision suffisante depuis le graphique source]")
ann.add_paragraph()
add_note(ann, "Rappel méthodologique : la Méthode 2 officielle (Partie 6.2 du document principal) applique le coefficient ERP forfaitaire 10-15% (estimation interne Presto) au marché total BRG, conformément au protocole MAB v8 — elle ne doit pas être confondue avec le ratio Non-Housing réel de 7,11% présenté ci-dessus à titre de donnée BRG brute.")
ann.add_paragraph()

add_heading(ann, "ANNEXE 3 — DONNÉES CONSTRUCTION TURQUIE 2015-2026", 1)
make_table(ann, [
    ["Indicateur", "Valeur", "Année", "Source"],
    ["Contribution BTP au PIB", "190,62 Md TRY", "2015", "TurkStat"],
    ["Contribution BTP au PIB", "1 466,47 Md TRY", "2023", "TurkStat"],
    ["Croissance BTP", "-7,1%", "2021", "Turkish Contractors Association/TurkStat"],
    ["Croissance BTP", "-8,6%", "2022", "Turkish Contractors Association/TurkStat"],
    ["Croissance BTP", "+7,8%", "2023", "Turkish Contractors Association/TurkStat"],
    ["Croissance BTP (T2)", "+6,5%", "T2 2024", "Turkish Contractors Association/TurkStat"],
    ["Répartition résidentiel/non-résidentiel", "79,7% / 20,3%", "2017", "Association of Turkish Construction Material Producers/TurkStat"],
    ["Répartition résidentiel/non-résidentiel", "76,5% / 23,5%", "2023", "Association of Turkish Construction Material Producers/TurkStat"],
    ["Valeur marché construction total", "53 Mds$", "2024", "Verified Market Research"],
    ["Valeur marché construction total", "117,75 Mds$", "2024", "IMARC Group"],
    ["Valeur marché construction total", "173,56 Mds$", "2025", "Mordor Intelligence"],
    ["Renaissance Construction (revenus)", "1 209 M$", "2023", "Statista"],
    ["Ant Yapi Industry & Trade (revenus)", "840 M$", "2023", "Statista"],
    ["Limak Insaat Sanayi ve Ticaret (revenus)", "558 M$", "2023", "Statista"],
    ["Inflation coûts matériaux construction", "+20-30%/an", "2024-2025", "Türkiye Today"],
    ["Hausse salaires main d'œuvre qualifiée", "+114% depuis 2023", "2025", "Türkiye Today"],
    ["Dépréciation TRY/USD", "-17 à -21%", "08/2024-09/2025", "Trading Economics"],
], font_size=8)
ann.add_paragraph()

add_heading(ann, "ANNEXE 4 — PIPELINE INVESTISSEMENTS ERP TURQUIE 2023-2030", 1)
make_table(ann, [
    ["Programme", "Budget", "Période", "Source"],
    ["Türkiye Earthquake Recovery and Reconstruction Project", "1 Md$ (prêt BM) ; coût total 45-80 Mds$", "2023-en cours", "World Bank 06/2023"],
    ["Reconstruction écoles post-séisme", "160 M$ (58% salles réparées, 42% en reconstruction ~9 800)", "2023-2025+", "Dossier interne Presto"],
    ["City Hospitals Program (PPP)", "105 Mds TRY (2025, contractants) ; 202 Mds TRY total PPP", "2016-en cours, >20 ouverts depuis 2022", "bne IntelliNews 2025"],
    ["TOKİ Urban Transformation (\"Century\")", "2 984 Mds TRY cumulés depuis 2002 ; nouveau volet 500K logements", "Lancé 10/2025", "Daily Sabah/AGBI 10/2025"],
    ["Budget éducation total", "150,78 Mds TRY pour 1 027 projets (15% budget investissement public)", "2024", "AGBI 01/2024"],
    ["Programme d'investissement public global", "1,9 trillion TRY (46,2 Mds$) pour 3 783 projets", "2025", "AGBI/GlobeNewswire 2025"],
    ["Projets PPP achevés (tous secteurs)", "270 projets, 204 Mds$ cumulés", "Cumulé à 2025", "Mordor Intelligence 2025"],
    ["Réparations équipements sportifs post-séisme", "28,27 M€ (29 livrés, 123 planifiés)", "2023-2025+", "Dossier interne Presto"],
    ["Réparations mosquées post-séisme", "3 800+ endommagées (34% du parc 11 provinces), 43% restant fin 2025", "2023-2025+", "Dossier interne Presto"],
], font_size=8)
ann.add_paragraph()

add_heading(ann, "ANNEXE 5 — DONNÉES ERP TURQUIE PAR SEGMENT (PARC INOX SANITAIRE)", 1)
make_table(ann, [
    ["Segment", "Rang marché inox", "% parc inox total ERP", "Parc inox estimé", "Durée de vie (standard/anti-vandale)"],
    ["4.8 Pénitentiaire", "1er", "18-22%", "200 000 pièces", "15 ans / 12 ans (95% anti-vandale)"],
    ["4.10 Lieux de culte", "2e", "12-15%", "120-150 000 pièces", "23 ans / 18 ans (15% anti-vandale)"],
    ["4.1 Éducation", "3e", "13-16%", "120-150 000 pièces", "22 ans / 16 ans (15% anti-vandale)"],
    ["4.2 Santé", "4e", "11-14%", "105-135 000 pièces", "19 ans / 25 ans (15% anti-vandale)"],
    ["4.7 Sport & Loisirs", "5e", "11-14%", "100-130 000 pièces", "19 ans / 16 ans (35% anti-vandale)"],
], font_size=8)
ann.add_paragraph()
make_table(ann, [
    ["Segment", "Parc total (établissements)", "Source"],
    ["4.1 Éducation", "≈75 500 (80% public MEB, 20% privé, 208 universités YÖK)", "Dossier interne Presto"],
    ["4.2 Santé", "≈1 560 hôpitaux (60% publics, 35% privés, 5% universitaires)", "Dossier interne Presto"],
    ["4.7 Sport & Loisirs", "≈4 500 équipements", "Dossier interne Presto"],
    ["4.8 Pénitentiaire", "407 établissements, 402 000 incarcérés", "Dossier interne Presto"],
    ["4.10 Lieux de culte", "≈82 000 mosquées", "Dossier interne Presto"],
], font_size=8)
ann.add_paragraph()
add_note(ann, "Ces données de parc (sources internes propriétaires) sont datées de la constitution du dossier de recherche (2025-2026) sans année de référence précise systématique — recommandation de sourcer chaque chiffre via les ministères compétents (MEB éducation, Ministère Santé, Ministère Justice, Diyanet) en phase terrain.")
ann.add_paragraph()

add_heading(ann, "ANNEXE 6 — DISTRIBUTEURS ROBINETTERIE TURQUIE", 1)
make_table(ann, [
    ["Holding", "Nom commercial", "Type", "Profil produit", "CA (M€)", "Dépôts", "Couverture géographique"],
    ["ELGINKAN", "ELEKS Foreign Trade Inc.", "Wholesale/Exporter", "Chauffage et bain", "[NON DISP.]", "[NON DISP.]", "Europe, Moyen-Orient, CEI"],
    ["ELGINKAN", "ELMOR", "Wholesale", "Bain", "[NON DISP.]", "[NON DISP.]", "Turquie"],
    ["ECZACIBASI", "INTEMA", "Wholesale", "Bain", "[NON DISP.]", "[NON DISP.]", "Turquie"],
    ["BAUHAUS", "BAUHAUS", "DIY Retailer", "Assortiment DIY", "[NON DISP.]", "10", "Istanbul(5), Ankara(3), Bursa(1), Antalya(1)"],
    ["KOÇTAŞ YAPI MARKETLERI", "KOÇTAŞ", "DIY Retailer", "Assortiment DIY", "180 (2020)", "171", "Turquie"],
    ["TEKFEN", "TEKZEN", "DIY Retailer", "Assortiment DIY", "190 (2020)", "128", "Turquie, Roumanie, Irak"],
], font_size=8)
ann.add_paragraph()
add_para(ann, "Source : BRG TR Apr21, p.25 (III.2 — TR Summary of Leading Distributors).")
ann.add_paragraph()

add_heading(ann, "ANNEXE 7 — NORMES ET CERTIFICATIONS DÉTAILLÉES ROBINETTERIE TURQUIE", 1)
make_table(ann, [
    ["Norme / Certification", "Référence complète", "Objet et champ d'application", "Exigences techniques détaillées", "Lien produits Presto", "Organisme"],
    ["TS EN 200", "Adoption directe EN 200 par TSE/TSI", "Robinets et mitigeurs sanitaires — exigences générales eau froide/chaude", "Débit nominal, pression de service (PN10), étanchéité, résistance mécanique, matériaux (laiton chromé, copolymère acétal)", "Base de toute la gamme robinetterie ERP Presto", "TSE"],
    ["TS EN 817", "Adoption directe EN 817 par TSE/TSI", "Robinetterie monocommande (mitigeurs mécaniques) PN10", "Tests d'endurance mécanique du levier, mélange eau froide/chaude", "Mitigeurs thermostatiques Presto", "TSE"],
    ["TS EN 1111", "Adoption directe EN 1111 ; certificat type A-TSE-1111-03 (VitrA)", "Robinetterie thermostatique PN10", "Stabilité thermique, dispositif anti-brûlure, matériaux résistants à la corrosion", "Mitigeurs thermostatiques Presto (santé, sport)", "TSE"],
    ["TS EN 15091", "Adoption EN 15091 (texte détaillé payant TSE)", "Robinetterie à ouverture/fermeture électronique (\"fotoselli\")", "[DONNÉE NON DISPONIBLE — temps de fermeture, alimentation, IP non confirmés]", "Robinetterie sans contact/infrarouge Presto", "TSE"],
    ["TS EN 816", "Adoption EN 816, équivalent fonctionnel EN 12541 (FR)", "Robinetterie à fermeture automatique temporisée", "[DONNÉE NON DISPONIBLE — durée fermeture, débit limité non confirmés]", "CŒUR DE GAMME — robinetterie temporisée push-button Presto", "TSE"],
    ["TS EN 1112/1113", "Adoption directe EN 1112/1113", "Douchettes à main et flexibles de douche", "Débit, résistance à la pression, durabilité du flexible", "Douches ERP (sport, santé, prison)", "TSE"],
    ["TS EN 274-1/2/3", "Adoption directe EN 274", "Siphons et raccords d'évacuation", "Résistance ≥80°C, résistance aux acides, démontable/nettoyable", "Accessoires complémentaires Presto", "TSE"],
    ["TS EN 248", "Adoption directe EN 248", "Essais de corrosion des revêtements métalliques", "Tests de résistance du chromage", "Finitions robinetterie Presto", "TSE"],
    ["TS EN ISO 3822", "Adoption directe EN ISO 3822", "Mesure acoustique des robinetteries", "Tests de bruit en laboratoire", "Conformité acoustique gamme Presto", "TSE"],
    ["TSE (certification générale)", "Procédure TSE standard", "Certification produit pour mise sur marché turque", "Dépôt dossier, audit préliminaire, inspection site, certificat par référence norme", "Prérequis pour toute commercialisation ERP en Turquie", "TSE (Türk Standartları Enstitüsü)"],
    ["Marquage CE (UE 305/2011)", "Réglementation UE produits de construction", "Mise sur marché UE — reconnu via union douanière mais insuffisant seul", "Prérequis mais ne remplace pas le contrôle TSE/TAREKS à l'import", "Tous produits Presto exportés vers la Turquie", "Commission UE"],
    ["RoHS", "Directive UE 2011/65", "Restriction substances dangereuses — équipements électriques/électroniques", "Applicable à la robinetterie électronique/sans-contact Presto", "Robinetterie électronique Presto", "Commission UE / TSE (contrôle import)"],
    ["ISO 9001 / ISO 14001", "Normes internationales qualité/environnement", "Systèmes de management qualité et environnemental", "Attendus comme preuve de robustesse industrielle en appels d'offres publics", "Certification fabricant Presto (déjà détenue en France à vérifier)", "Organismes certificateurs internationaux"],
], font_size=7)
ann.add_paragraph()

add_heading(ann, "ANNEXE 8 — MATRICE DE CONFIANCE DONNÉES CLÉS", 1)
make_table(ann, [
    ["Donnée", "Valeur", "Source", "Niveau confiance", "À revérifier"],
    ["PIB/hab Turquie 2025", "18 200-18 611 USD", "IMF WEO 10/2025 / Worldometer", "Moyenne (fourchette)", "Choisir source unique IMF WEO"],
    ["Population Turquie 2025", "85,5-86,09 M", "Source interne / TurkStat", "Haute", "Non"],
    ["Marché construction total", "53-174 Mds$", "Multi-cabinets (Mordor, IMARC, VMR, R&M)", "Faible (écart large)", "Source TurkStat officielle à privilégier"],
    ["Contribution BTP au PIB 2023", "1 466,47 Md TRY", "TurkStat", "Haute", "Non"],
    ["Répartition résidentiel/non-résidentiel 2023", "76,5% / 23,5%", "TurkStat", "Haute", "Non"],
    ["Marché robinetterie ERP Turquie (interne)", "~56 M€, ~4M pièces, TCAC ~4%", "Dossier interne propriétaire", "Moyenne (méthode non détaillée)", "Confronter aux Méthodes 1 & 2 officielles"],
    ["Total taps & mixers TR 2020 (BRG)", "6 776 000 unités, 1 229,75 M TRY, 152,96 M EUR", "BRG TR Apr21, p.101", "Haute (donnée directe fabricants)", "Mise à jour post-2021 recommandée"],
    ["Non-Housing volume 2020 (BRG)", "482 010 unités (7,11%)", "BRG TR Apr21, p.108", "Haute", "Non"],
    ["Présence Delabie Turquie", "Absente/non confirmée", "Recherche web croisée + Delabie.com", "Haute (absence)", "Contact direct export Delabie recommandé"],
    ["Parc pénitentiaire", "407 établissements, 402 000 incarcérés", "Dossier interne propriétaire", "Moyenne-haute", "Sourcer via Ministère Justice turc / Prison Insider"],
    ["Parc mosquées", "≈82 000", "Dossier interne propriétaire", "Moyenne", "Sourcer via Diyanet officiel"],
    ["Parc hospitalier", "≈1 560 hôpitaux (60% publics)", "Dossier interne propriétaire", "Moyenne", "Sourcer via Ministère Santé turc"],
    ["Parc scolaire", "≈75 500 établissements", "Dossier interne propriétaire", "Moyenne", "Sourcer via MEB"],
    ["Normes TS EN robinetterie", "TS EN 200/817/1111/15091/816/1112/1113/274/248/ISO3822", "Coşan Ticaret, VitrA, CSB (bordereaux prix)", "Moyenne-haute (nomenclature confirmée, détails partiels)", "Détails TS EN 15091 et 816 à compléter"],
    ["Coefficient d'extrapolation X", "0,470 brut / 0,423 ajusté", "Calcul propre (formule MAB)", "Haute (calcul) / Moyenne (input PIB/hab)", "Lié à la fiabilisation du PIB/hab"],
    ["Programmes investissement (City Hospitals, TOKİ, reconstruction)", "Chiffrés et datés", "bne IntelliNews, invest.gov.tr, AGBI, World Bank", "Haute", "Non"],
    ["Taux de change EUR/TRY 2026", "≈51 TRY/EUR, -27%/an", "ECB eurofxref / Capital.com", "Haute", "Volatilité continue à surveiller"],
], font_size=7)
ann.add_paragraph()

# Sauvegarde annexes
path_ann = os.path.join(OUTPUT_DIR, "MAB_Turquie_Annexes.docx")
ann.save(path_ann)
print(f"✓ Annexes v1 sauvegardées : {path_ann}")
print()
print("═" * 60)
print("MAB TURQUIE v2 — Génération terminée")
print(f"  Étude   → {path_etude}")
print(f"  Annexes → {path_ann}")
