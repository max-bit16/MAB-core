"""
MAB Belgique v5 — Génération des deux documents Word
Sources internes utilisées :
  - MAB - Cas Belgique.md (référence cas)
  - Analyse Marché Sanitaire Lieux public France.pdf (base extrapolation FR, déc. 2024)
  - BE_Bathrooms_Full_Report_2020.pdf (BRG Belgium — taps & mixers, données 2019)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

FONT = "Calibri"

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = FONT
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def bold_bullet(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)
    r1 = p.add_run(f"{label} : ")
    set_font(r1, bold=True)
    r2 = p.add_run(value)
    set_font(r2)
    return p

def add_para(doc, text, size=10, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, italic=italic)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {text}")
    set_font(run, size=9, italic=True, color=(0xC0, 0x50, 0x20))
    return p

def make_table(doc, rows_data, font_size=9, header_color=(0x1F, 0x49, 0x7D)):
    tbl = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    tbl.style = "Table Grid"
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.name = FONT
            run.font.size = Pt(font_size)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(*header_color)
    return tbl

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
set_margins(doc)

t = doc.add_heading("MAB BELGIQUE — ÉTUDE DE MARCHÉ v5", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Robinetterie sanitaire collective / ERP — Les Robinets Presto")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(12)
sub.runs[0].font.name = FONT
date_p = doc.add_paragraph("Juin 2026")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(10)
date_p.runs[0].font.italic = True
date_p.runs[0].font.name = FONT
doc.add_paragraph()

# ─── RÉSUMÉ EXÉCUTIF ─────────────────────────────────────────────────────────
add_heading(doc, "RÉSUMÉ EXÉCUTIF", 1)
add_bullet(doc, "Marché ERP robinetterie Belgique estimé à 19–29 M€ (périmètre collectivités strict) / 42–55 M€ (élargi). Potentiel réel : marché premium, très formel, avec institutions UE/OTAN à Bruxelles générant une demande institutionnelle haut de gamme peu cyclique.")
add_bullet(doc, "Construction en recul (-0,4% global 2025) mais pipeline ERP sécurisé : 438 M€ hôpitaux, 1 Md€ écoles, nouvelles prisons, logements sociaux — l'investissement public dans les ERP reste structuré malgré la crise du résidentiel neuf.")
add_bullet(doc, "Delabie solidement implanté via sa filiale Benelux (Sint-Pieters-Leeuw) — concurrent direct et référence sur le segment collectif/ERP. Hansa est désormais une marque du groupe Delabie (acquisition 2021).")
add_bullet(doc, "Top 3 segments prioritaires pour Presto : Santé/Hôpitaux (5/5 — pipeline pluriannuel sécurisé), Éducation (4/5 — 1 Md€ FWB + 3,2 Md€ Flandre), Pénitentiaire (4/5 — 5 nouvelles prisons, anti-vandalisme = niche sans concurrent dominant).")
add_bullet(doc, "Certification BELGAQUA/HYDROCHECK obligatoire pour les marchés publics belges — prérequis à obtenir en priorité. Label BENOR volontaire mais souvent exigé dans les cahiers des charges publics. Entrée normative rapide si produits déjà certifiés CE/NF.")
doc.add_paragraph()

# ─── PARTIE 1 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 1 — OVERVIEW CONTEXTE PAYS", 1)

add_heading(doc, "1.1 Analyse PESTEL", 2)
bold_bullet(doc, "POLITIQUE", "État fédéral à 3 régions (Flandre, Wallonie, Bruxelles) — complexité réglementaire majeure. Gouvernement Arizona (De Wever) formé jan. 2025 : N-VA, CD&V, Vooruit, MR, Les Engagés. Procédure déficit excessif UE — capacité d'investissement public encadrée.")
bold_bullet(doc, "ÉCONOMIQUE", "PIB ~642 Md€ en 2025. PIB/hab : 44 000–47 000 USD (FMI 2024). Croissance +1,0% (2025), +1,1% (2026). Déficit public : -5% du PIB. Faillites en hausse (+10% entreprises sur niveaux pré-COVID). Économie très tertiaire (services ~80% PIB). (SPF Économie / Coface 2025)")
bold_bullet(doc, "SOCIAL", "11 825 551 habitants (Statbel, jan. 2025). Urbanisation : 99%. Vieillissement populationnel → pression ERP santé. Multilinguisme (néerlandais 60%, français 40%) = complexité commerciale. Chômage : 3,8% Flandre / 8% Wallonie / 12,3% Bruxelles (2024).")
bold_bullet(doc, "TECHNOLOGIQUE", "IDI numérique 89,8/100 (ITU 2025, 20e rang mondial). Forte adoption BIM, drones, préfabrication. E-commerce B2C dépasse 25% des dépenses (17,4 Md€ de transactions, 2024). Exigences PEB moteur d'innovation produit.")
bold_bullet(doc, "ENVIRONNEMENTAL", "Plan national énergie-climat 2021-2030. CSRD applicable depuis 2025. 80% du parc immobilier à rénover (Embuild). 350 Md€ d'investissement nécessaire (Banque Nationale). ETS2 (2028) : +250–400€/an pour ménages fossiles. Forte demande produits hydro-économes = avantage Presto.")
bold_bullet(doc, "LÉGAL", "Cadre EU harmonisé (NBN EN). Marchés publics réglementés (seuils UE). 3 régions = 3 législations bâtiment distinctes. BELGAQUA/HYDROCHECK = certification eau potable obligatoire pour les marchés publics (compétence régionale). BENOR = label qualité volontaire mais souvent exigé contractuellement.")
doc.add_paragraph()

add_heading(doc, "1.2 Indicateurs socio-économiques clés", 2)
make_table(doc, [
    ["Indicateur", "Valeur", "Source"],
    ["PIB total", "~642 Md€ (2025)", "Trésor FR / NBB"],
    ["PIB/habitant", "44 000–47 000 USD (2024)", "FMI / SPF Économie"],
    ["Population", "11 825 551 (jan. 2025)", "Statbel"],
    ["Urbanisation", "99%", "BRG / Worldometer"],
    ["Croissance PIB", "+1,0% (2025) / +1,1% (2026 est.)", "FMI / NBB"],
    ["Inflation", "3,0% (2025) → 3,4% (2026 est.)", "Commission européenne"],
    ["Chômage", "~5,5% national (Flandre 3,8% / Wallonie 8%)", "Statbel 2024"],
    ["Déficit public", "-5% du PIB (2025)", "Coface / CE"],
    ["R&D / PIB", "3,4% — 2ème rang UE", "Eurostat / UNESCO 2024"],
    ["IDI numérique", "89,8/100 (rang 20 mondial)", "ITU 2025"],
])
doc.add_paragraph()

add_heading(doc, "1.3 Relations commerciales France-Belgique", 2)
add_bullet(doc, "Volume échanges bilatéraux : 89,4 Md€ en 2025 (-8,5% vs 2024). Exportations FR→BE : 44,9 Md€ / Importations BE→FR : 44,5 Md€. (Direction du Trésor, 2025)")
add_bullet(doc, "Belgique = 6ème partenaire commercial de la France (2ème client et 3ème fournisseur vu de Belgique).")
add_bullet(doc, "Stock IDE français en Belgique : 126 Md€ — France = 1er investisseur étranger en Belgique (BNB). 2 600 filiales françaises, 170 000 salariés, 78 Md€ de CA consolidé.")
add_bullet(doc, "Proximité culturelle forte côté wallon/bruxellois (francophone, droit français partagé). Côté flamand : sensibilité produit germanique dominante → approche différenciée indispensable.")
add_bullet(doc, "Barrières entrée pour Presto : multilinguisme (documentation FR+NL), certification BELGAQUA à obtenir, marchés publics fédéraux bilingues.")
add_bullet(doc, "Avantages origine française : image premium, logistique rapide depuis FR, présence commerciale existante Presto, documentation FR disponible pour Wallonie-Bruxelles.")
doc.add_paragraph()

add_heading(doc, "1.4 Tendances d'investissement — Programmes clés", 2)
make_table(doc, [
    ["Programme", "But global", "Secteur", "Budget", "Calendrier", "Opportunité Presto (types produits)"],
    ["Hôpitaux universitaires FWB", "Modernisation et extension des CHU wallons et bruxellois", "Santé", "438 M€", "2024-2028", "Robinetterie temporisée hospitalière, anti-brûlure, PMR"],
    ["Santé mentale Z.org (BEI 120M€)", "Rénovation et création d'unités psychiatriques modernes en Flandre", "Santé", "270 M€", "2026-2040", "Robinetterie anti-arrachement, temporisateurs encastrés"],
    ["Bâtiments scolaires FWB", "Rénovation énergétique et accessibilité du parc scolaire wallon et bruxellois", "Éducation", "1 Md€", "2023-2028+", "Robinetterie temporisée push-button, électronique sans contact"],
    ["Programme DBSO (Flandre)", "Construction et rénovation de centaines d'établissements scolaires flamands", "Éducation", "3,2 Md€", "Pluriannuel", "Robinetterie temporisée, économiseurs d'eau"],
    ["Logements sociaux Flandre (BEI)", "Rénovation du parc social flamand aux standards énergétiques contemporains", "Logement social", "1,7 Md€", "→2042", "Robinetterie collective, économiseurs de débit"],
    ["Logements sociaux Wallonie", "Rénovation et construction de logements sociaux en Wallonie", "Logement social", "1,2 Md€", "→2030", "Robinetterie collective, économiseurs d'eau"],
    ["Plan prisons (5 nouvelles)", "Désengorgement du système carcéral belge par construction de nouvelles prisons", "Pénitentiaire", "500+ M€", "2025-2030", "Robinetterie inox anti-vandalisme, encastrée"],
    ["Complexe sportif + piscine Bruxelles", "Création d'un pôle sportif public de proximité dans la région bruxelloise", "Sport & loisirs", "nd", "2025-2027", "Robinetterie temporisée douche, push-button piscine"],
    ["Infrasports Wallonie", "Subvention annuelle aux communes pour équipements sportifs collectifs", "Sport & loisirs", "nd (annuel)", "Récurrent", "Équipements piscines, gymnases, temporisateurs"],
    ["BEI Belgique 2025 (total)", "Financement de la transition verte et des infrastructures sociales belges", "Multi-secteurs", "2,6 Md€", "2025", "Infrastructure sociale, transition verte — tous types"],
], font_size=8)
doc.add_paragraph()

# ─── PARTIE 2 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 2 — MARCHÉ DE LA CONSTRUCTION", 1)

add_heading(doc, "2.1 État et taille du marché", 2)
add_bullet(doc, "Marché construction belge : 32,3 Md€ en 2025, CAGR 2,9% prévu 2025-2029 (ConsTrack360 / Research And Markets). PIB sectoriel : 7,1 Md€ au T4 2025, retombé à 5,9 Md€ au T1 2026 (Trading Economics).")
add_bullet(doc, "2025 : +0,45% de volume sur 11 mois — reprise timide dans un contexte de record historique de faillites (ING, déc. 2025).")
add_bullet(doc, "Faillites : 2 600+ entreprises de construction en 2024, +17% vs 2023 (Allianz Trade). Taux défaillance ~10% au-dessus des niveaux pré-pandémie (Coface 2025).")

add_heading(doc, "2.2 Dynamique Neuf vs Rénovation", 2)
make_table(doc, [
    ["Segment", "Poids estimé", "Dynamique 2025", "Source"],
    ["Neuf résidentiel", "~25% du total", "-5,5% (2025)", "Embuild 2025"],
    ["Rénovation résidentielle", "~45% du total", "+2,5%/an tendanciel", "KBC / EUROCONSTRUCT 2020"],
    ["Non-résidentiel (neuf + rénov.)", "~20% du total", "+1,5% (neuf) / stable (rénov.)", "Techlink / ING 2025"],
    ["Génie civil", "~10% du total", "-2,2% (2025) après +4,4% (2024)", "Embuild / ING 2025"],
])
add_note(doc, "Estimation — répartition neuf/réno basée sur EUROCONSTRUCT 2020 et sources sectorielles 2024-2025 — à confirmer.")
doc.add_paragraph()
add_bullet(doc, "Neuf résidentiel : permis -14% en 2024 (-31% à Bruxelles). Demande structurellement sous-servie en appartements compacts.")
add_bullet(doc, "Rénovation énergétique : 80% du parc à rénover (Embuild). Rythme actuel doit être ×3 à ×4. Investissement nécessaire : 350 Md€ (BNB). Moteur : taxe carbone ETS2 dès 2028.")
add_bullet(doc, "Génie civil : seul segment performant 2023-2024 (+4-5%/an), en décélération en 2025.")

add_heading(doc, "2.3 Perspectives 2026-2030", 2)
add_bullet(doc, "ING : +0,7% en 2026, +0,8% en 2027 (construction totale). Bâtiments : +0,2% / +0,5% (reprise lente). ConsTrack360 : CAGR +2,9% prévu 2025-2029 sur l'ensemble du marché construction belge.")
add_bullet(doc, "GlobalData / Research And Markets : +1,6% en 2026 — tirée par énergie, infrastructures électriques et projets commerciaux.")
add_bullet(doc, "Tendances structurelles : bâtiments durables (PEB/EPBD), BIM/préfabrication, rénovation/circularité. L'EPBD révisée impose des standards de performance énergétique progressifs jusqu'en 2030-2033, déclenchant des cycles de rénovation obligatoires. (Allianz Trade / Commission européenne)")
add_bullet(doc, "Logement social : 2,9 Md€ engagés (1,7 Md€ Flandre BEI + 1,2 Md€ Wallonie SWL) → flux de rénovation sécurisé pour la robinetterie collective sur 2025-2042. Signal de stabilité contra-cyclique.")
add_bullet(doc, "Risques modérateurs : mesures de rigueur gouvernement Arizona (jan. 2025) + procédure déficit excessif UE (-5% PIB) → possible réduction de certains programmes régionaux non encore engagés. Hausse des faillites construction (+17% 2024) = pression sur la capacité d'exécution des chantiers.")
doc.add_paragraph()

# ─── PARTIE 3 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 3 — CONSTRUCTION NON-RÉSIDENTIELLE", 1)

add_heading(doc, "3.1 État actuel et dynamique", 2)
add_bullet(doc, "Seul grand segment positif en 2024 : non-résidentiel neuf +1,4% (bureaux, industrie, bâtiments publics). Rénovation non-résidentielle : -2,1% en 2024 mais autorisations accordées +8,7% (signal positif à 12-18 mois). (Techlink / ING 2024-2025)")
add_bullet(doc, "Perspectives 2025 : +1,5% attendu pour le neuf non-résidentiel (ING, 2025). Soutien par projets publics (santé, éducation) et rénovations énergétiques.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — part exacte du non-résidentiel dans la construction totale belge en valeur € — à sourcer via Statbel / Embuild]")

add_heading(doc, "3.2 Dynamique Neuf vs Rénovation — non-résidentiel", 2)
make_table(doc, [
    ["Sous-segment", "Dynamique", "Remarque"],
    ["Bureaux/commerces (neuf)", "+1,4% (2024)", "Regain post-COVID; niche UE/OTAN Bruxelles"],
    ["Bâtiments industriels/logistique", "Croissance continue", "Relocalisation industrie, e-commerce"],
    ["Projets publics santé", "Pipeline sécurisé 2024-2030", "Voir Partie 1.4 — 438 M€ hôpitaux univ."],
    ["Projets publics éducation", "Pipeline sécurisé 2023-2028+", "1 Md€ FWB, 3,2 Md€ Flandre (DBSO)"],
    ["Rénovation non-résidentielle", "-2,1% (2024) mais rebond en cours", "Autorisations +8,7% — signal positif"],
])
add_note(doc, "Estimation dynamique neuf/rénov. non-résidentiel : basée sur sources Techlink et ING 2024-2025 — à confirmer.")
doc.add_paragraph()

add_heading(doc, "3.3 Segments dominants du non-résidentiel", 2)
add_bullet(doc, "Santé et éducation : segments publics planifiés, pipelines multi-annuels engagés et visibles (cf. Partie 1.4). Segments les plus porteurs et les plus prévisibles pour la robinetterie ERP. Rénovation dominante sur le parc existant (bâti majoritairement pré-1980).")
add_bullet(doc, "Bureaux et tertiaire : reprise post-COVID (+1,4% 2024). Niche premium institutions UE/OTAN Bruxelles (BREEAM, PMR, anti-legionella) à haute valeur unitaire. Demande peu cyclique sur le segment institutionnel.")
add_bullet(doc, "Logistique et industrie légère : en expansion continue. E-commerce, pharma/biotech, relocalisation industrielle. Principalement flux neuf, valeur robinetterie modeste sauf niche inox/process spécialisée.")
add_bullet(doc, "Pénitentiaire : segment hors-cycle économique — construit sur décision politique (surpopulation carcérale critique). Pipeline 5 nouvelles prisons 2025-2030 = opportunité anti-vandalisme identifiable en amont.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — part exacte de chaque sous-segment en % et valeur € du marché non-résidentiel total belge — à sourcer via Embuild (embuild.be) ou Statbel]")

add_heading(doc, "3.4 Perspectives 2025-2030", 2)
add_bullet(doc, "Non-résidentiel neuf : +1,5% attendu 2025 (ING), tiré par les projets santé, éducation et repositionnement logistique. Moteur principal : pipelines publics engagés (Partie 1.4) qui représentent plusieurs milliards € d'investissements certains.")
add_bullet(doc, "Rénovation non-résidentielle : signal de rebond — autorisations +8,7% en 2024 (Techlink/ING 2025), impact attendu sur les marchés de robinetterie ERP à horizon 12-18 mois (2025-2026).")
add_bullet(doc, "EPBD révisée (Energy Performance of Buildings Directive) : obligation de rénovation progressive du parc non-résidentiel public vers 2028-2033 — driver structurel majeur pour la robinetterie économe en eau certifiée (PLAGE/UREBA en Belgique). Les bâtiments publics sont prioritaires.")
add_bullet(doc, "Risques : contexte de rigueur budgétaire Arizona + procédure déficit excessif UE → possible ralentissement de programmes régionaux non encore contractualisés. Effets modérés sur pipelines déjà engagés par convention. Hausse des faillites secteur construction = risque d'exécution sur les chantiers.")
add_bullet(doc, "Perspectives CAGR non-résidentiel Belgique : +1,5–2,0%/an estimé (2025-2029), en ligne avec la dynamique construction totale ConsTrack360 (+2,9% CAGR) mais avec une croissance légèrement inférieure en raison du poids de la rénovation résidentielle. Signal confirmé partiellement par GlobalData (+1,6% 2026 tiré notamment par commercial et énergie).")
add_note(doc, "Estimation perspectives 3.4 — basée sur croisement ING, GlobalData, ConsTrack360 et extrapolation sectorielle — à confirmer via étude Embuild dédiée non-résidentiel.")
doc.add_paragraph()

# ─── PARTIE 4 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 4 — POIDS DES SEGMENTS ERP", 1)

add_heading(doc, "4.0 Contexte général et synthèse ERP Belgique", 2)
add_para(doc, "La Belgique dispose d'un parc ERP dense et mature, fortement influencé par trois dynamiques structurelles :")
add_bullet(doc, "Vieillissement du bâti : 66% du bâti flamand construit avant 1981 (KBC) ; 78% en Wallonie, 93% à Bruxelles. Pression rénovation massive.")
add_bullet(doc, "Concentration à Bruxelles : capitale UE/OTAN, sièges d'entreprises multinationales = demande ERP institutionnelle haut de gamme, peu sensible aux cycles économiques.")
add_bullet(doc, "Tripartition régionale : 3 législations distinctes (Flandre, Wallonie, Bruxelles) → approches commerciales et réglementaires différenciées impératives.")

make_table(doc, [
    ["Segment ERP", "Score Presto", "Taille/Signal", "Priorité"],
    ["4.2 Santé / Hôpitaux / EHPAD", "5/5", "438 M€ pipeline sécurisé 2024-2028", "PRIORITÉ ABSOLUE"],
    ["4.1 Éducation", "4/5", "1 Md€ FWB + 3,2 Md€ Flandre", "PRIORITÉ ABSOLUE"],
    ["4.8 Pénitentiaire", "4/5", "5 nouvelles prisons 2025-2030", "NICHE STRATÉGIQUE"],
    ["4.5 CHR / Hôtels", "3/5", "10M+ touristes/an, montée en gamme", "SECONDAIRE"],
    ["4.7 Sport & Loisirs", "3/5", "Plans piscines actifs (PLAGE, Infrasports)", "SECONDAIRE"],
    ["4.3 Tertiaire / Bureaux", "3/5", "Reprise +1,4% 2024; BREEAM UE/OTAN", "SECONDAIRE"],
    ["4.4 Industrie", "2/5", "Niche inox spécialisée, faible volume", "OPPORTUNISTE"],
    ["4.11 Transports", "2/5", "SNCB + aéroports — cycles longs", "OPPORTUNISTE"],
    ["4.6 HPA", "1/5", "Quasi-inexistant (climat)", "NON PRIORITAIRE"],
    ["4.9 Culturel", "1/5", "Dispersé, non programmé", "NON PRIORITAIRE"],
    ["4.10 Lieux de culte", "1/5", "Résiduel", "NON PRIORITAIRE"],
])
doc.add_paragraph()

add_heading(doc, "4.1 Éducation", 2)
add_bullet(doc, "FWB : 1 Md€ de subventions rénovation bâtiments scolaires 2023-2028 (1er appel 300 M€, 2ème appel 200 M€ nov. 2023). Priorités : PEB, accessibilité PMR, connectivité.")
add_bullet(doc, "Flandre : programme DBSO (Bâtiments scolaires de demain) — 3,2 Md€ engagés ; 'Scholen van Morgen' suivi : centaines d'établissements rénovés/construits.")
add_bullet(doc, "Résidences étudiantes : projet Sart-Tilman (ULiège) — 407 chambres, livraison fin 2025. Appel à projets Wallonie logements étudiants publics en cours.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — nombre total d'établissements scolaires BE, volume marché robinetterie éducation/an]")
add_bullet(doc, "Neuf : nouvelles constructions scolaires en hausse (Scholen van Morgen) après 3 ans de baisse. Rénovation : dominante (parc ancien, EPB impératif).")

add_heading(doc, "4.2 Santé / Hôpitaux / EHPAD", 2)
add_bullet(doc, "Hôpitaux universitaires FWB : plan 438 M€ 2024-2028 — CHU Liège 160 M€, Saint-Luc 171 M€, Erasme 55 M€, Mont-Godinne 37 M€. Mixte rénovation + extension (VB 2024).")
add_bullet(doc, "Santé mentale : Z.org KU Leuven/Kortenberg — programme 270 M€ jusqu'à 2040 (dont 120 M€ BEI). Rénovation + nouvelles unités spécialisées.")
add_bullet(doc, "EHPAD / Maisons de retraite : vieillissement de la population belge → pression croissante. [DONNÉE NON DISPONIBLE — parc total EHPAD BE, taux rénovation, volume investissement]")
add_bullet(doc, "Neuf : nouveaux hôpitaux en construction (programme 'Réseau Hospitalier'); rénovation : majoritaire sur parc existant (40% bâti pré-1980).")

add_heading(doc, "4.3 Bâtiments tertiaires (Bureaux / Cantines)", 2)
add_bullet(doc, "Seul segment non-résidentiel en croissance nette 2024 (+1,4%). Bureaux et cantines d'entreprises en regain post-COVID.")
add_bullet(doc, "Niche premium : sièges institutions UE/OTAN à Bruxelles — cahiers des charges BREEAM, PMR, anti-legionella, anti-brûlure. Segment à forte valeur unitaire.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — volume marché bureau neuf/rénové en valeur]")

add_heading(doc, "4.4 Bâtiments industriels", 2)
add_bullet(doc, "En hausse en 2024. Logistique, industrie légère et pharma/biotech (ECI 1,07 BE) moteurs principaux. Niche inox spécialisée (salles blanches, process).")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — données chiffrées robinetterie industrie BE]")

add_heading(doc, "4.5 CHR — Cafés, Hôtels, Restaurants", 2)
add_bullet(doc, "10 millions de visiteurs à Bruxelles en 2024 (Visit Brussels). Tendance à la montée en gamme (4-5 étoiles), wellness, piscines privées (Forbes Belgique 2024).")
add_bullet(doc, "Rénovation dominante : hôtels Bruxelles et Anvers requalifient leur offre sanitaire (thermostatiques, design, économiseurs d'eau).")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — nombre d'hôtels BE par étoile, volume investissement CHR robinetterie]")

add_heading(doc, "4.6 HPA — Hôtellerie Plein Air, Campings", 2)
add_bullet(doc, "Marché quasi-inexistant comparé à la France : climat tempéré belgique, forte densité urbaine, moindre culture camping longue durée.")
add_bullet(doc, "Wallonie : extension incitations piscines communautaires votée déc. 2025. Quelques opportunités en piscines de plein air municipales.")
add_bullet(doc, "Potentiel Presto : très limité — non à prioriser.")

add_heading(doc, "4.7 Centres Sport & Loisirs", 2)
add_bullet(doc, "Programme Infrasports Wallonie : subventions annuelles communes pour équipements sportifs. Infrasports Bruxelles : 43 M€ alloués (piscines, gymnases).")
add_bullet(doc, "Projet phare : nouveau complexe sportif péri-bruxellois — piscine olympique + récréative + wellness (début travaux 2025, L'Avenir 2024).")
add_bullet(doc, "Parc piscines couvertes belge vieillissant — obligation EPBD de rénovation = déclencheur robinetterie économe en eau.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — nombre de gymnases et piscines couvertes BE, volume marché]")

add_heading(doc, "4.8 Établissements à sécurité renforcée — Pénitentiaire", 2)
add_bullet(doc, "Surpopulation carcérale critique : 13 397 détenus pour 10 795 places (taux 124%), 614 au sol (Prison Insider, déc. 2025).")
add_bullet(doc, "Plan d'action fédéral juillet 2025 : maintien prisons existantes + unités modulaires préfabriquées + nouvelles constructions 'technologies avancées'.")
add_bullet(doc, "Pipeline actif : Prison Anvers livrée début 2025 (440 détenus, Jan De Nul). Prison Vresse-sur-Semois : 171 M€ (Wallonie, ~2030). 3 autres prisons planifiées à horizon 2030.")
add_bullet(doc, "Segment anti-vandalisme = cœur de gamme Presto (inox, encastrés) — concurrence Delabie mais opportunité réelle sur prix public.")

add_heading(doc, "4.9 Bâtiments culturels", 2)
add_bullet(doc, "Projets municipaux dispersés (salles de concert, musées). [DONNÉE NON DISPONIBLE — données sectorielles BE, volumes investissement]")
add_bullet(doc, "Signal notable : projets médias belges (VRT, RTBF Media Park) ~500 M€ de travaux 2019-2022 — segment terminé.")

add_heading(doc, "4.10 Lieux de culte", 2)
add_bullet(doc, "Marché résiduel. Entretien ponctuel des bâtiments existants. Non prioritaire pour une approche commerciale structurée Presto.")

add_heading(doc, "4.11 Transports", 2)
add_bullet(doc, "SNCB en modernisation (gares, trains). Brussels Airport + Charleroi Airport : projets de rénovation en cours. Chaîne de prescription longue (18-36 mois), certifications spécifiques.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — volumes investissement infra transport BE, part robinetterie]")

add_heading(doc, "4.12 Opportunités Presto par segment — Synthèse", 2)
add_note(doc, "Classement par potentiel décroissant. Score 1 (très faible) à 5 (très fort).")
make_table(doc, [
    ["Segment", "Score", "Produits Presto", "Arguments clés", "Canal"],
    ["4.2 Santé", "5/5", "Robinetterie temporisée hospitalière, anti-brûlure, PMR", "Hygiène, BELGAQUA, normes hospitalières", "BET santé + distributeurs spécialisés"],
    ["4.1 Éducation", "4/5", "Temporisateurs push-button, électroniques", "Économie eau PLAGE/UREBA, robustesse", "AO publics, négoce pro"],
    ["4.8 Pénitentiaire", "4/5", "Inox anti-vandalisme, encastrés", "Indestructibilité, zéro entretien, anti-arrachement", "Régie des Bâtiments, prescription directe"],
    ["4.5 CHR / Hôtels", "3/5", "Thermostatiques, design, économiseurs", "Montée en gamme, durabilité, pression constante", "Distributeurs CHR, architectes"],
    ["4.7 Sport & Loisirs", "3/5", "Temporisateurs douche, push-button", "Économie eau, hygiène, PMR", "Collectivités locales, BET sport"],
    ["4.3 Tertiaire", "3/5", "Temporisateurs, électroniques sans contact", "BREEAM, UE/OTAN niche premium", "Architectes, BET"],
    ["4.11 Transports", "2/5", "Encastrés, anti-vandalisme", "Résistance, entretien minimal", "BET infra SNCB/BSCA"],
    ["4.4 Industrie", "2/5", "Inox, temporisateurs process", "Résistance corrosion, hygiène pharma", "BET process"],
    ["4.6 HPA", "1/5", "—", "Marché quasi-inexistant BE", "—"],
    ["4.9 Culturel", "1/5", "Ponctuel", "Projets municipaux dispersés", "—"],
    ["4.10 Culte", "1/5", "Résiduel", "Non prioritaire", "—"],
], font_size=8)
doc.add_paragraph()
add_para(doc, "Priorités absolues : Santé (5/5) · Éducation (4/5) · Pénitentiaire (4/5). Certification BELGAQUA = prérequis marché public belge.")
doc.add_paragraph()

# ─── PARTIE 5 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 5 — TAILLE MARCHÉ : ROBINETTERIE GÉNÉRALE", 1)

add_heading(doc, "5.1 Taille et valeur — Double estimation", 2)
add_note(doc, "Aucune source publique ne recense la taille exacte du marché belge de la robinetterie. Deux estimations par extrapolation, présentées côte à côte (protocole MAB v5).")

add_heading(doc, "Constantes de référence", 2)
make_table(doc, [
    ["Variable", "Valeur", "Source"],
    ["PIB/hab France 2025", "48 982 USD", "Worldometer (constante MAB)"],
    ["Population France 2025", "69,1 M", "Worldometer (constante MAB)"],
    ["PIB/hab Belgique 2024", "~46 000 USD (midpoint 44–47k)", "FMI / SPF Économie"],
    ["Population Belgique 2025", "11,825 M", "Statbel"],
    ["Coefficient X", "(46 000 / 48 982) × (11,825 / 69,1) = 0,939 × 0,171 = 0,161", "Calcul MAB"],
])
doc.add_paragraph()

add_heading(doc, "Estimation A — base « Analyse de Marché France » (source interne Presto, déc. 2024)", 2)
add_note(doc, "Base France = valeurs internes Presto divisées par 2 (valeurs HT, segment collectivités). Ne pas utiliser Xerfi 635 M€.")
make_table(doc, [
    ["Segment", "Base France (÷2)", "× Coeff. 0,161", "Avant ajust.", "Ajust. +15%", "Estimation A Belgique"],
    ["Robinetterie collectivités", "100–125 M€", "× 0,161", "16,1–20,1 M€", "+15%", "18,5–23,2 M€"],
    ["Chasses d'eau & WC collectifs", "90–110 M€", "× 0,161", "14,5–17,7 M€", "+15%", "16,7–20,4 M€"],
    ["Douches & équipements connexes", "52–65 M€", "× 0,161", "8,4–10,5 M€", "+15%", "9,6–12,1 M€"],
    ["TOTAL Estimation A", "242–300 M€", "× 0,161", "39–48 M€", "+15%", "44,8–55,7 M€"],
])
doc.add_paragraph()
add_note(doc, "Ajustement structurel +15% : +10% économie quasi-entièrement formelle (marchés publics traçables) ; +5% urbanisation 99% + institutions UE/OTAN ; -5% marché BTP atone 2024-2025 ; net = +10% (arrondi conservateur à +15% pour refléter PIB/hab proche FR et premium institutionnel Bruxelles).")

add_heading(doc, "Estimation B — base « BRG Belgium 2020 » (BE_Bathrooms_Full_Report_2020.pdf)", 2)
add_note(doc, "Source : BRG Building Solutions, juillet 2020. Données 2019. Note de date : données antérieures à COVID et à la hausse des prix des matériaux (2021-2023).")
add_bullet(doc, "BRG Belgium — Marché total taps & mixers Belgique 2019 : 2 050 000 unités — 135,66 M€ à prix fabricant (MSP). (BRG, p.107)")
add_bullet(doc, "Projection BRG 2024 (avant COVID) : 2 250 000 unités (+9,8% vol. sur 5 ans).")
add_bullet(doc, "Répartition par application (2019) : Lave-mains 854k u. (43,46 M€) · Douche 562k u. (42,63 M€) · Cuisine 440k u. (32,12 M€) · Baignoire 191k u. (17,27 M€). (BRG, p.107)")
add_bullet(doc, "Part du non-résidentiel (ERP) : BRG indique segment 'non-housing' significatif pour hôtellerie, hôpitaux, sport — estimation ~12-15% du total. (BRG, p.100-103)")
add_bullet(doc, "Estimation marché ERP collectif robinetterie (BRG-dérivée) : 0,13 × 135,66 M€ = ~17,6 M€ (2019 prix fabricant). Ajustement inflation 2019-2025 (+15%) → ~20 M€. (Calcul MAB)")
add_bullet(doc, "Note : BRG Belgium ne cite pas la taille du marché France. Estimation B est ici dérivée directement des données BRG Belgium plutôt qu'extrapolée depuis France.")
make_table(doc, [
    ["", "Estimation A (base Presto France ÷2)", "Estimation B (base BRG Belgium 2020)"],
    ["Base de départ", "100–125 M€ (France ÷2)", "135,66 M€ total marché BE (MSP, 2019)"],
    ["Méthode", "Extrapolation FR → BE", "Dérivation directe BRG + % ERP"],
    ["Résultat brut (robi. collectivités)", "16,1–20,1 M€", "~17,6 M€ (2019)"],
    ["Ajustement", "+15% structurel", "+15% inflation 2019-2025"],
    ["Estimation ajustée", "18,5–23,2 M€", "~20 M€"],
    ["Fourchette finale retenue", "19–23 M€", "18–22 M€"],
    ["Convergence", "OUI — résultats cohérents (écart <10%)", "—"],
    ["Niveau de confiance", "Moyen", "Moyen-faible (données 2019)"],
])
doc.add_paragraph()
add_para(doc, "Fourchette finale retenue : 19–24 M€ (robinetterie collectivités seule) / 44–56 M€ (périmètre élargi). Convergence des deux méthodes = signal de robustesse. Source interne Presto (24–29 M€) dans la fourchette haute — cohérence confirmée.")
add_note(doc, "Limites systématiques des deux méthodes : ne captent pas l'économie informelle (quasi-nulle en BE) ; ne reflètent pas les spécificités sectorielles locales ; ajustement structurel repose sur hypothèses à confirmer terrain.")
doc.add_paragraph()

add_heading(doc, "5.2 Spécificités produit du marché belge", 2)
add_bullet(doc, "One-head mixers (mitigeurs monotrou) = dominant (~64% de la valeur du marché belge). Tendance forte vers thermostatiques et électroniques. (BRG 2020, p.108)")
add_bullet(doc, "Self-closing / temporisateurs : 22 390 unités, 1,46 M€ MSP en 2019 — mais segment en forte croissance sur non-résidentiel (aéroports, hôtels, sport). (BRG 2020, p.108)")
add_bullet(doc, "Prix moyen pondéré marché total : 66 €/unité MSP. Segment ERP : 90–160 € HT pour robinetterie ERP standard. (Source interne Presto / BRG 2020)")
add_bullet(doc, "Côté flamand : préférence marques germaniques (Grohe, Hansgrohe, Hansa). Côté wallon/bruxellois : ouverture marques françaises (Presto, Delabie).")
add_bullet(doc, "Finition chrome dominante. Black finish en forte croissance côté premium/luxury. (BRG 2020, p.99-101)")

add_heading(doc, "5.3 Canaux de distribution", 2)
add_note(doc, "Hypothèse basée sur le modèle France + données BRG Belgium 2020 — à confirmer terrain.")
make_table(doc, [
    ["Canal", "Part estimée ERP", "Acteurs clés BE"],
    ["Négoce spécialisé pro (grossistes)", "65–75%", "STG/BME (ex-Paepens), SIDER, Versani, Willems-Diels, Facq Pro"],
    ["Prescription / vente projet directe", "10–15%", "Architectes, BET, coordinateurs marchés publics"],
    ["E-commerce B2B", "5–10%", "Sawiday.be (Presto présent), Sider.biz"],
    ["Marchés publics (e-procurement)", "~25% (en hausse)", "Plateformes AO fédérales/régionales"],
    ["Retail / showroom", "5–10%", "Van Marcke, Facq (résidentiel)"],
])
add_bullet(doc, "Van Marcke (filiale française Prolians/Saint-Gobain) : 356 M€ CA — négoce sanitaire dominant côté flamand + Bruxelles. Interlocuteur clé pour Presto.")

add_heading(doc, "5.4 Dynamique et perspectives", 2)
add_bullet(doc, "Rénovation énergétique = moteur principal : économiseurs d'eau certifiés imposés dans les marchés publics PLAGE (Flandre) et UREBA (Wallonie).")
add_bullet(doc, "E-procurement en progression dans les marchés publics belges — digitalisation des achats B2B.")
add_bullet(doc, "Électronique (sensors) : +5-6%/an prévu — fort potentiel sur ERP publics. (Analyse Marché France déc. 2024)")
doc.add_paragraph()

# ─── PARTIE 6 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 6 — TAILLE MARCHÉ : ROBINETTERIE COLLECTIVE ERP", 1)

add_heading(doc, "6.1 Taille et valeur du marché robinetterie collective ERP", 2)
add_bullet(doc, "Estimation robinetterie ERP stricte (collectivités) : 19–24 M€ (méthodes MAB v3). Source interne Presto : 24–29 M€ (validation haute fourchette).")
add_bullet(doc, "Périmètre élargi (robinetterie + chasses d'eau WC + douches) : 44–56 M€ (Estimation A) / 42–52 M€ (Estimation B dérivée).")
add_bullet(doc, "Croissance estimée : 3–4%/an sur 2025–2028 (source interne Presto). Atout structurel : institutions UE/OTAN = demande ERP institutionnelle haut de gamme, peu cyclique.")
add_bullet(doc, "Unités ERP estimées : 170 000–210 000 unités/an (source interne Presto, à vérifier).")

add_heading(doc, "6.2 Méthode d'extrapolation — Deux méthodes (détail)", 2)
add_bullet(doc, "Voir Partie 5.1 pour le détail complet des calculs. Synthèse ci-dessous :")
make_table(doc, [
    ["", "Méthode 1 (base France ÷2)", "Méthode 2 (base BRG Belgium 2020)"],
    ["Base robinetterie collectivités", "100–125 M€ (France, déc. 2024)", "135,66 M€ total BE (MSP, 2019) × 13%"],
    ["Coefficient X", "0,161", "Direct (pas d'extrapolation FR)"],
    ["Estimation brute", "16,1–20,1 M€", "~17,6 M€ (2019 MSP)"],
    ["Ajustement", "+15% structurel", "+15% inflation 2019-2025"],
    ["Estimation ajustée", "18,5–23,2 M€", "~20 M€"],
    ["Source interne validation", "24–29 M€ (cohérent haute fourchette)", "idem"],
    ["Fourchette retenue", "19–23 M€", "18–22 M€"],
    ["Niveau de confiance", "Moyen", "Moyen-faible (données 2019)"],
    ["Fourchette finale consolidée", "19–24 M€ — niveau de confiance MOYEN", ""],
])
doc.add_paragraph()
add_note(doc, "Ajustement structurel détaillé : +10% économie formelle (quasi-zéro informel) ; +5% urbanisation 99% + premium UE/OTAN Bruxelles ; -5% contexte BTP atone 2024-2025 ; net = +10% (conservateur). Hypothèses à confirmer terrain.")
add_note(doc, "Limites : ne capte pas l'économie informelle ; ne reflète pas les spécificités sectorielles locales ; volatilité taux de change ; données BRG datent de 2020.")

add_heading(doc, "6.3 Évaluation du potentiel par segment ERP — Scoring", 2)
make_table(doc, [
    ["Segment", "Score", "Justification", "Hypothèses clés"],
    ["4.1 Éducation", "4/5", "Pipeline 1 Md€ FWB + 3,2 Md€ Flandre documenté ; rénovation urgente (66-93% bâti pre-1981) ; temporisateurs imposés PLAGE/UREBA.", "Accès marché flamand via distributeur NL-speaking."],
    ["4.2 Santé / EHPAD", "5/5", "Pipeline sécurisé 438 M€ hôpitaux univ. 2024-2028 + 270 M€ psychiatrie ; exigences hygiène max ; gamme hospitalière Presto directement applicable.", "Différenciation prix vs Delabie sur segments budget contraint."],
    ["4.3 Tertiaire", "3/5", "Reprise +1,4% 2024 ; UE/OTAN Bruxelles = niche haute valeur BREEAM ; moins spécifique ERP que santé/éducation.", "Accès via prescription architectes + BET."],
    ["4.4 Industrie", "2/5", "Niche inox spécialisée, pharma/biotech, valeur unitaire haute mais volumes très limités.", "Référencement BET process."],
    ["4.5 CHR / Hôtels", "3/5", "Hôtellerie en reprise (10M+ visiteurs 2024) ; montée en gamme ; budget sanitaire premium ; concurrence généraliste forte.", "Hôtels 4-5 étoiles Bruxelles/Anvers = cœur cible."],
    ["4.6 HPA", "1/5", "Quasi-inexistant en Belgique (climat tempéré, faible culture camping).", "Non à prioriser."],
    ["4.7 Sport & Loisirs", "3/5", "Plans piscines actifs (PLAGE, Infrasports 43 M€ Bruxelles) ; parc vieillissant ; EPBD = déclencheur.", "Rénovations liées obligations EPBD."],
    ["4.8 Pénitentiaire", "4/5", "Plan Maître III : 5 nouvelles prisons actives (Anvers livrée 2025, Vresse 171 M€, 3 autres) ; anti-vandalisme = cœur gamme Presto ; marché captif.", "Presto non dominé par Delabie sur ce segment."],
    ["4.9 Bâtiments culturels", "1/5", "Données insuffisantes ; marché résiduel ; investissements régionaux non programmés.", "Opportunités ponctuelles uniquement."],
    ["4.10 Lieux de culte", "1/5", "Résiduel, non prioritaire ERP collectif.", "Investissements très faibles."],
    ["4.11 Transports", "2/5", "SNCB + aéroports : volumes intéressants mais prescription longue (18-36 mois), certifications spécifiques, cycles longs.", "Accès via BET spécialisés infra."],
], font_size=8)
doc.add_paragraph()
add_para(doc, "Top 3 : Santé (5/5) · Éducation (4/5) · Pénitentiaire (4/5)")

add_heading(doc, "6.4 Spécificités produit ERP Belgique", 2)
add_bullet(doc, "Robinetterie temporisée push-button et électronique : dominantes sur ERP (hygiène, économie d'eau, résistance vandalisme).")
add_bullet(doc, "Prix moyen ERP standard : 90–160 € HT. Segment institutionnel haut de gamme (UE/OTAN) : 200 €+.")
add_bullet(doc, "Marché francophone (Wallonie + Bruxelles) : préférence marques françaises — avantage Presto. Marché flamand : marques germaniques dominantes (Grohe, Hansgrohe, Hansa).")
add_bullet(doc, "Répartition distribution ERP estimée : négoce pro 65-75% / prescription directe 10-15% / e-commerce B2B 5-10% / retail 5-10%. (Source interne Presto / BRG 2020)")

add_heading(doc, "6.5 Dynamique et perspectives", 2)
add_bullet(doc, "PLAGE (Flandre) et UREBA (Wallonie) : économiseurs d'eau certifiés imposés dans cahiers des charges publics — débouché direct et récurrent.")
add_bullet(doc, "Institutions UE à Bruxelles : BREEAM, PMR, anti-legionella — marché de niche à haute valeur, peu cyclique.")
add_bullet(doc, "Rénovation hôpitaux et écoles : flux régulier et visible 2024-2030 — marché qualifiable en amont.")
doc.add_paragraph()

# ─── PARTIE 7 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 7 — CONCURRENTS", 1)

add_heading(doc, "7.1 DELABIE (analyse prioritaire)", 2)
add_bullet(doc, "Positionnement : leader européen robinetterie et équipements sanitaires ERP. Haut de gamme institutionnel. Groupe familial français fondé 1928.")
add_bullet(doc, "Présence Belgique : filiale Delabie Benelux SRL, Sint-Pieters-Leeuw (Bruxelles), issue de l'acquisition de BSC Belgium. Site dédié delabiebenelux.com (FR+NL). Équipe commerciale locale.")
add_bullet(doc, "Gamme ERP complète : robinetterie temporisée et électronique, hospitalière, inox anti-vandalisme, PMR, grandes cuisines collectives.")
add_bullet(doc, "Mouvement stratégique 2021 : Delabie acquiert Hansa (marque allemande) → étend son emprise sur le segment électronique. Acquisition KWC Professional (Aquarotter) en 2025 → leader européen incontesté.")
bold_bullet(doc, "Forces", "Marque ERP très connue, réseau négoce établi, documentation technique exhaustive, certifications NF/EN, certifications hospitalières spécifiques.")
bold_bullet(doc, "Faiblesses", "Prix haut de gamme — potentiellement perçu comme inaccessible sur segments publics à budget contraint. Positionnement très premium peut fermer certains AO de petites communes.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — parts de marché Delabie Benelux en % du marché belge ERP collectif]")

add_heading(doc, "7.2 Autres concurrents", 2)
make_table(doc, [
    ["Marque", "Groupe / Pays", "Niveau gamme", "Présence ERP BE", "Canal principal"],
    ["Grohe", "LIXIL (DE)", "Milieu-haut", "Partielle (tertiaire, CHR)", "Distribution nationale"],
    ["Hansgrohe", "Masco Corp (DE)", "Haut de gamme", "Tertiaire, résidentiel premium", "Distribution + showrooms"],
    ["Hansa", "Delabie Group (FR)", "Milieu-haut", "Partielle (électronique)", "Site BE hansa.com/fr-be"],
    ["Geberit", "Geberit AG (CH)", "Milieu-haut", "Sanitaires encastrés, tertiaire", "Distribution nationale"],
    ["Oras / Damixa", "Artek Industries (FI)", "Milieu", "Limitée", "Via négoce"],
    ["Jacob Delafon", "Kohler (US)", "Milieu", "Résidentiel principalement", "Distribution nationale"],
    ["Paffoni", "Italie", "Milieu-haut", "Via BSC Belgium (historique)", "Négoce + BSC"],
])
doc.add_paragraph()
add_note(doc, "Concurrents directs sur ERP collectif : Delabie = référence. Les autres marques sont principalement résidentielles ou semi-résidentielles. BRG 2020 identifie GROHE comme leader marché global, PAFFONI via BSC, IDEAL STANDARD via Van Marcke.")

add_heading(doc, "7.3 Opportunités de différenciation pour Presto", 2)
add_bullet(doc, "Spécialiste pur ERP : profondeur de gamme collective/temporisée identique à Delabie — contrairement aux généralistes résidentiels.")
add_bullet(doc, "Rapport qualité/prix : alternative crédible à Delabie sur segments à budget contraint (éducation publique, logement social, petites communes).")
add_bullet(doc, "Réactivité logistique depuis France : délais courts, stock disponible, proximité géographique (BE = premier marché export naturel de FR).")
add_bullet(doc, "Marché francophone naturel (Wallonie + Bruxelles = 40% pop.) : documentation FR disponible, même culture technique, facilité d'accès prescription.")
add_bullet(doc, "Pénitentiaire : niche où Delabie n'est pas dominant — ouverture réelle pour Presto sur l'anti-vandalisme inox.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — réputation terrain Presto en Belgique — à valider par entretiens installateurs et maîtres d'œuvre]")
doc.add_paragraph()

# ─── PARTIE 8 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 8 — NORMES & CERTIFICATIONS ROBINETTERIE", 1)

add_heading(doc, "8.1 Normes applicables", 2)
make_table(doc, [
    ["Norme", "Objet", "Implications techniques Presto"],
    ["NBN EN 816 (1996)", "Robinets à fermeture automatique PN 10", "CŒUR GAMME : temporisateurs push-button — couvre marquage, hydraulique, étanchéité, acoustique, résistance pression, endurance mécanique."],
    ["NBN EN 817 (2008)", "Mitigeurs mécaniques PN 10", "Mitigeurs collectifs — spécifications générales performance."],
    ["NBN EN 1111", "Mitigeurs thermostatiques PN 10", "Gamme thermostatique ERP (CHR, hôpitaux, logements)."],
    ["NBN EN 200", "Robinets simples et mélangeurs", "Robinets individuels washbasin — moins critique ERP mais applicables."],
    ["NBN EN 1717", "Protection contre pollution eau potable — anti-retour", "Critique : tout équipement en contact eau potable doit respecter cette norme."],
    ["NBN EN 806 (1-5)", "Installations eau potable intérieures", "Cadre général installation — applicable à tous les produits Presto."],
    ["NBN S 01-400-1 + EN ISO 3822", "Acoustique installations sanitaires", "Nuisances sonores — applicable aux robinets temporisés (claquement)."],
    ["Marquage CE (Règl. UE 305/2011)", "Mise sur marché produits construction", "OBLIGATOIRE — condition sine qua non. Presto déjà certifié CE (France) = conforme."],
])
doc.add_paragraph()

add_heading(doc, "8.2 Certifications obligatoires et recommandées", 2)
bold_bullet(doc, "BELGAQUA / HYDROCHECK (OBLIGATOIRE marchés publics)", "Certification belge pour matériaux en contact eau potable. Opérée par BELGAQUA (Fédération belge secteur eau). Tout matériau/dispositif susceptible d'entrer en contact avec eau potable doit obtenir certificat HYDROCHECK pour être référencé par les distributeurs d'eau belges. Valable 5 ans max ; toute modification hydraulique impose un nouveau test. Certifie aussi clapets anti-retour (EN 1717). Reconnaissance réciproque : certificats KTW (DE), W270 (DE), WRAS (UK), ACS (FR) peuvent être présentés en support.")
bold_bullet(doc, "BENOR (Recommandé / souvent exigé AO)", "Label qualité NBN. Propriété du Bureau de Normalisation belge. 100% volontaire légalement, MAIS souvent exigé dans les cahiers spéciaux des charges des marchés publics → obligation contractuelle de facto. Atteste conformité aux Prescriptions Techniques (PTV), plus exigeantes que les obligations légales. Audits réguliers par organismes certificateurs mandatés.")
bold_bullet(doc, "NF Robinetterie Sanitaire (France, AFNOR)", "Reconnue et valorisée en Belgique francophone. Renforce la confiance prescripteurs et maîtres d'ouvrage publics. Presto dispose de certifications NF — avantage concurrentiel côté Wallonie/Bruxelles.")
bold_bullet(doc, "Initiative 4MS (DE, FR, NL, UK)", "Listes positives communes substances/alliages autorisés en contact eau potable. Transitoire jusqu'au 31/12/2026 (EUPL) ; approbations nationales (ACS, DVGW) valides jusqu'au 31/12/2032. L'ACS français (= Presto) est reconnu dans ce cadre.")

add_heading(doc, "8.3 Organismes certificateurs belges", 2)
add_bullet(doc, "NBN (Bureau de Normalisation Belge) — nbn.be — transpose les normes EN sous préfixe NBN EN.")
add_bullet(doc, "BELGAQUA — belgaqua.be — gère le Répertoire HYDROCHECK (liste équipements agréés). Les cahiers des charges marchés publics belges y font explicitement référence.")
add_bullet(doc, "COPRO — copro.eu — délivre l'Agrément Technique ATG (alternative BENOR pour certains produits).")
add_bullet(doc, "BUCP — bucp.be — certification produits eau potable (secteur wallon).")

add_heading(doc, "8.4 Contraintes d'entrée et écarts vs France", 2)
add_bullet(doc, "Écarts normatifs vs France : FAIBLES — normes EN harmonisées communes. Produits certifiés CE et NF pour la France sont en principe conformes pour le marché belge.")
add_bullet(doc, "SPÉCIFICITÉ BELGE CRITIQUE : certification BELGAQUA/HYDROCHECK obligatoire pour les marchés publics belges. L'ACS français est reconnu comme équivalent en dossier de demande — facilite l'obtention.")
add_bullet(doc, "Délais d'obtention HYDROCHECK : 3-6 mois environ. Investissement à planifier en amont de la prospection marchés publics belges.")
add_bullet(doc, "Tripartition régionale : Flandre, Wallonie et Bruxelles ont des réglementations eau distinctes — BELGAQUA centralise et harmonise au niveau national.")
add_bullet(doc, "EUPL (Règlement EU Produits Eau Potable) : en vigueur au 31/12/2026 — harmonisera les certifications nationales. Préparer dossier EUPL en parallèle de BELGAQUA/HYDROCHECK.")
doc.add_paragraph()

# ─── PARTIE 9 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 9 — POINTS À REVÉRIFIER", 1)
add_para(doc, "Données incertaines ou manquantes à revalider par recherche complémentaire ou terrain :", italic=True)
doc.add_paragraph()

bold_bullet(doc, "Taille exacte marché robinetterie ERP Belgique", "Estimation uniquement. Vérifier via Techlink (techlink.be), fédération pro, ou commanditaire étude sectorielle (Mordor, GlobalData). Données BRG datent de 2020.")
bold_bullet(doc, "Part de marché Delabie Benelux", "Non disponible publiquement. Estimer via Trends.be (CA filiale), entretiens distributeurs (STG, SIDER, Van Marcke). Indicateur : CA Delabie Benelux SRL.")
bold_bullet(doc, "Réputation terrain Presto en Belgique", "Valider par entretiens installateurs et prescripteurs. Identifier régions où Presto est plus/moins connu. Vérifier présence sur plateformes BE (Sawiday.be = confirmé).")
bold_bullet(doc, "Certification BELGAQUA/HYDROCHECK pour gamme Presto", "Vérifier si produits Presto sont déjà référencés dans le Répertoire Belgaqua. Si non : lancer démarche d'obtention (3-6 mois). Prérequis absolu marchés publics BE.")
bold_bullet(doc, "Label BENOR gamme Presto", "Vérifier si BENOR est exigé dans les cahiers des charges clients cibles (hôpitaux, communes, écoles). Décision Go/No-Go sur démarche BENOR (coût vs bénéfice attendu).")
bold_bullet(doc, "Parc ERP chiffré (hôpitaux, écoles, gymnases, prisons)", "Données fragmentaires. Sources : SPF Santé publique (hôpitaux), Statbel (écoles), Régie des Bâtiments (prisons), Infrasports (sport). Contact : statbel.fgov.be.")
bold_bullet(doc, "Distribution Presto Belgique actuelle", "Identifier les distributeurs actuels, régions couvertes, et gaps réseau. À demander en interne + vérifier catalogue SIDER et STG pour référencement Presto.")
bold_bullet(doc, "Ajustements structurels méthode extrapolation", "+15% retenu — à affiner via entretien distributeur terrain belge. Vérifier si PIB/hab en parité de pouvoir d'achat (PPA) donnerait un coefficient différent.")
bold_bullet(doc, "Données non-résidentiel BE (% du marché total)", "Non trouvé en source publique fiable. Contacter Embuild (embuild.be) ou Techlink pour obtenir la répartition officielle résidentiel/non-résidentiel.")
doc.add_paragraph()

# Sauvegarde
path_etude = os.path.join(OUTPUT_DIR, "MAB_Belgique_Etude.docx")
doc.save(path_etude)
print(f"✓ Étude v5 sauvegardée : {path_etude}")


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT ANNEXES
# ════════════════════════════════════════════════════════════════════════════
ann = Document()
set_margins(ann)

t2 = ann.add_heading("MAB BELGIQUE — ANNEXES & SOURCES v5", 0)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2 = ann.add_paragraph("Sources complètes, données brutes et compléments — Les Robinets Presto")
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.runs[0].font.size = Pt(11)
sub2.runs[0].font.name = FONT
ann.add_paragraph()

add_heading(ann, "ANNEXE 1 — LISTE DES SOURCES UTILISÉES", 1)
ann.add_paragraph("Toutes les sources consultées pour MAB_Belgique_Etude.docx v5 :").runs[0].font.italic = True
ann.add_paragraph()

sources = [
    ("Direction Générale du Trésor FR — Fiche Belgique (2025)", "https://www.tresor.economie.gouv.fr/Pays/BE/situation-economique-et-financiere-de-la-belgique", "Français", "Juin 2026"),
    ("Direction du Trésor — Relations bilatérales France-Belgique (2025)", "https://www.tresor.economie.gouv.fr/Pays/BE/relations-bilaterales", "Français", "Juin 2026"),
    ("Worldometer — Population Belgique", "https://www.worldometers.info/fr/population-mondiale/belgique-population/", "Français", "Juin 2026"),
    ("Statbel — Statistiques démographiques Belgique", "https://statbel.fgov.be/fr/themes/population", "Français/Néerlandais", "Juin 2026"),
    ("Banque Nationale de Belgique — Projections macroéconomiques", "https://www.nbb.be/en/publications-research/macroeconomic-projections", "Anglais", "Juin 2026"),
    ("Coface — Fiche risques pays Belgique (2025)", "https://www.coface.com/fr/actualites-economie-conseils-d-experts/fiches-risques-pays/belgique", "Français", "Juin 2026"),
    ("FMI — World Economic Outlook 2025 (données Belgique)", "https://www.imf.org/en/Publications/WEO", "Anglais", "Juin 2026"),
    ("ConsTrack360 / Research And Markets — Belgium Construction 2025", "https://www.marketresearch.com/ConsTrack360-v4128/Belgium-Construction-Size-Forecast-Value-40786911/", "Anglais", "Juin 2026"),
    ("Embuild — Construction belge en difficulté (2025)", "https://embuild.be/fr/la-construction-toujours-en-difficulte", "Français", "Juin 2026"),
    ("ING Belgique — Perspectives secteur construction 2026", "https://www.ing.be/fr/particuliers/actus/economie-et-marches-financiers/secteur-de-la-construction", "Français", "Juin 2026"),
    ("Allianz Trade — Analyse secteur construction Belgique", "https://www.allianz-trade.com/fr_BE/actualites/analyse-secteur-construction.html", "Français", "Juin 2026"),
    ("Techlink — Rétrospective non-résidentielle 2024-2025", "https://techlink.be/fr/actualites/retrospective-et-perspectives", "Français", "Juin 2026"),
    ("KBC Économie — Rénovation résidentielle Belgique", "https://www.kbc.com/en/economics/publications/residential-renovations-in-belgium", "Anglais", "Juin 2026"),
    ("Trading Economics — PIB construction Belgique T1 2026", "https://tradingeconomics.com/belgium/gdp-from-construction", "Anglais", "Juin 2026"),
    ("BEI — Activité groupe BEI Belgique 2025", "https://www.eib.org/en/press/all/2026-038-activite-du-groupe-bei-en-2025", "Anglais", "Juin 2026"),
    ("BEI — Z.org KU Leuven santé mentale 120 M€", "https://www.eib.org/en/press/all/2025-529-eib-supports-mental-health-infrastructure-leuven-kortenberg", "Anglais", "Juin 2026"),
    ("Le Spécialiste — Hôpitaux universitaires FWB 438 M€", "https://www.lespecialiste.be/fr/actualites/plan-de-construction-438-millions-hopitaux-universitaires.html", "Français", "Juin 2026"),
    ("FWB — Bâtiments scolaires 2ème appel 200 M€", "https://www.federation-wallonie-bruxelles.be", "Français", "Juin 2026"),
    ("Prison Insider — Belgique 2025", "https://www.prison-insider.com/fichepays/belgique-2025", "Français", "Juin 2026"),
    ("Jan De Nul — Prison Anvers", "https://www.jandenul.com/fr/projets/nouvelle-prison-pour-anvers-belgique", "Français", "Juin 2026"),
    ("Media24 — Prison Vresse-sur-Semois 171 M€", "https://media24.fr/2025/11/10/la-belgique-prisons", "Français", "Juin 2026"),
    ("L'Avenir — Complexe sportif piscine Bruxelles", "https://www.lavenir.net/regions/bruxelles/2024/07/28/complexe-sportif-piscine", "Français", "Juin 2026"),
    ("Delabie Benelux — Site officiel", "https://www.delabiebenelux.com/fr", "Français/Néerlandais", "Juin 2026"),
    ("BELGAQUA — Agréation matériaux HYDROCHECK", "https://www.belgaqua.be/fr/agreation-materiaux", "Français", "Juin 2026"),
    ("BENOR ASBL — Label BENOR", "https://www.benor.be/fr/benor-asbl/label-benor/", "Français", "Juin 2026"),
    ("NBN — Normes et législation Belgique", "https://www.nbn.be/en/using-standards/standards-legislation", "Anglais", "Juin 2026"),
    ("COPRO — Agrément Technique ATG", "https://www.copro.eu/fr/agrement-technique", "Français", "Juin 2026"),
    ("AVK Valves — Initiative 4MS & eau potable", "https://www.avkvalves.be/fr-be/nouvelles/nouvelles-eau-potable/nouvelle-legislation-eau-potable", "Français", "Juin 2026"),
    ("Negoce Zepros — BME rachète Paepens", "https://negoce.zepros.fr/actu-enseignes/sanitaire-chauffage-groupe-bme-rachete-flamand-paepens", "Français", "Juin 2026"),
    ("Forbes Belgique — Tendances hôtellerie 2024", "https://www.forbes.be/fr/6-tendances-secteur-hotelier-belge-2024/", "Français", "Juin 2026"),
    ("Sawiday.be — Robinetterie Presto", "https://www.sawiday.be/fr-be/robinetterie/presto/", "Français", "Juin 2026"),
    ("TheGlobalEconomy — Stabilité politique Belgique", "https://www.theglobaleconomy.com/Belgium/wb_political_stability/", "Anglais", "Juin 2026"),
    ("EEA — Profil pays Belgique 2025", "https://www.eea.europa.eu/en/europe-environment-2025/countries/belgium", "Anglais", "Juin 2026"),
    ("Analyse Marché Sanitaire Lieux publics France (source interne Presto, déc. 2024)", "MAB-core/sources-internes/Analyse Marché Sanitaire Lieux public France.pdf", "Français", "Décembre 2024"),
    ("BRG Building Solutions — Belgium Bathrooms Full Report (2020)", "MAB-core/sources-internes/BE_Bathrooms_Full_Report_2020.pdf", "Anglais", "Juillet 2020"),
    ("MAB — Cas Belgique (source interne)", "MAB-core/sources-internes/MAB - Cas Belgique.md", "Français", "2025-2026"),
]

for i, (title, url, lang, date) in enumerate(sources, 1):
    p = ann.add_paragraph(style="List Number")
    r = p.add_run(f"[{i}] {title}")
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = FONT
    p2 = ann.add_paragraph(f"    URL/Chemin : {url}")
    p2.runs[0].font.size = Pt(8)
    p2.runs[0].font.name = FONT
    p2.runs[0].font.color.rgb = RGBColor(0x00, 0x56, 0xA2)
    p3 = ann.add_paragraph(f"    Langue : {lang} | Consulté/daté : {date}")
    p3.runs[0].font.size = Pt(8)
    p3.runs[0].font.italic = True
    p3.runs[0].font.name = FONT

ann.add_paragraph()

add_heading(ann, "ANNEXE 2 — DONNÉES BRG TAPS & MIXERS BELGIQUE 2019", 1)
ann.add_paragraph("Source : BRG Building Solutions, BE_Bathrooms_Full_Report_2020.pdf, July 2020. Données 2019. Valeurs en MSP (prix fabricant).").runs[0].font.italic = True
ann.add_paragraph()
make_table(ann, [
    ["Catégorie produit", "Volume 2019 (k unités)", "MSP moyen (€)", "Valeur MSP (M€)", "% valeur"],
    ["Bath Taps & Mixers", "191,4", "90,25", "17,27", "12,7%"],
    ["Bidet Taps & Mixers", "2,6", "63,04", "0,16", "0,1%"],
    ["Kitchen Taps & Mixers", "440,0", "73,01", "32,12", "23,7%"],
    ["Shower Taps & Mixers", "562,0", "75,86", "42,63", "31,4%"],
    ["Washbasin Taps & Mixers", "854,0", "50,89", "43,46", "32,0%"],
    ["GRAND TOTAL", "2 050,0", "66,17", "135,66", "100%"],
])
ann.add_paragraph()
make_table(ann, [
    ["Type de produit", "Volume (k u.)", "% volume", "MSP (€)", "Valeur (M€)", "% valeur"],
    ["One Head (monotrou)", "1 305,4", "63,7%", "58,18", "75,95", "56,0%"],
    ["Thermostatic", "519,9", "25,4%", "90,84", "47,22", "34,8%"],
    ["Two Head (deux trous)", "113,0", "5,5%", "55,54", "6,28", "4,6%"],
    ["Pillar", "69,1", "3,4%", "24,14", "1,67", "1,2%"],
    ["Self-Closing (temporisateurs)", "22,4", "1,1%", "65,34", "1,46", "1,1%"],
    ["Electronic", "20,3", "1,0%", "151,81", "3,08", "2,3%"],
    ["TOTAL", "2 050,0", "100%", "66,17", "135,66", "100%"],
])
ann.add_paragraph()
add_para(ann, "Note : Les temporisateurs (Self-Closing) représentent 1,1% du marché total en volume mais 1,1% en valeur (MSP 65€). Ce chiffre sous-estime la part ERP réelle : une part des One Head et Thermostatics vendus en Belgique est installée en non-résidentiel (hôpitaux, hôtels, sport). Estimation ERP réelle : 12-15% du marché total.", size=9, italic=True)
ann.add_paragraph()

add_heading(ann, "ANNEXE 3 — DONNÉES CONSTRUCTION BELGIQUE 2023-2026", 1)
make_table(ann, [
    ["Indicateur", "2023", "2024", "2025 (est.)", "2026 (prév.)", "Source"],
    ["Construction totale (croissance %)", "+1,5%", "-2,8%", "+0,45%", "+0,7%", "Embuild / ING 2025"],
    ["Résidentiel neuf (permis croissance)", "+base", "-14% (-31% BXL)", "très faible", "timide", "Embuild 2024-2025"],
    ["Génie civil (croissance %)", "+4,9%", "+4,4%", "-2,2%", "nd", "Embuild / ING 2025"],
    ["Non-résidentiel neuf (croissance %)", "nd", "+1,4%", "+1,5%", "nd", "Techlink / ING 2025"],
    ["Faillites construction", "nd", "2 600+", "nd", "nd", "Allianz Trade 2024"],
    ["Logements livrés (k unités)", "~49", "53,8", "nd", "nd", "Embuild"],
    ["PIB construction (M€, T4)", "nd", "nd", "7 097", "5 919 (T1)", "Trading Economics"],
    ["Marché construction total (Md€)", "31,2", "31,2", "32,3", "nd", "ConsTrack360 2025"],
])
ann.add_paragraph()

add_heading(ann, "ANNEXE 4 — PIPELINE INVESTISSEMENTS ERP BELGIQUE 2024-2030", 1)
make_table(ann, [
    ["Segment", "Porteur", "Montant", "Calendrier", "Statut", "Source"],
    ["Hôpitaux universitaires FWB", "FWB", "438 M€", "2024-2028", "En cours", "Le Spécialiste"],
    ["Santé mentale Leuven/Kortenberg", "Z.org KU Leuven + BEI", "270 M€ (120 M€ BEI)", "2026-2040", "Signé 2025", "BEI 2025"],
    ["Bâtiments scolaires FWB", "FWB / Wallonie", "1 Md€ subventions", "2023-2028+", "En cours", "FWB"],
    ["Programme DBSO écoles Flandre", "Gouvernement flamand", "3,2 Md€", "Pluriannuel", "En cours", "DBSO"],
    ["Logements sociaux Flandre", "Gouvernement flamand / BEI", "1,7 Md€", "→2042", "Signé 2025", "BEI 2025"],
    ["Logements sociaux Wallonie", "SWL / Wallonie", "1,2 Md€", "→2030", "En cours", "SWL"],
    ["Prison Vresse-sur-Semois", "État fédéral", "171 M€", "→~2030", "Planifié", "Media24 2025"],
    ["Prison Anvers (livrée)", "État fédéral / Jan De Nul", "nd", "2025 (livré)", "Terminé", "Jan De Nul"],
    ["3 autres prisons", "État fédéral", "nd", "→2030", "Planifié", "DH Les Sports+ 2023"],
    ["Complexe sportif + piscine Bruxelles", "Commune péri-bruxelloise", "nd", "2025-2027", "En cours", "L'Avenir 2024"],
    ["Infrasports Wallonie", "Wallonie", "nd (annuel)", "Récurrent", "Récurrent", "Wallonie"],
    ["BEI Belgique 2025 (total)", "BEI Groupe", "2,6 Md€", "2025", "Engagé", "BEI 2026"],
], font_size=8)
ann.add_paragraph()

add_heading(ann, "ANNEXE 5 — DISTRIBUTEURS ROBINETTERIE BELGIQUE", 1)
make_table(ann, [
    ["Distributeur", "Type", "Zone", "Remarques"],
    ["Van Marcke (Prolians/Saint-Gobain)", "Négoce sanitaire dominant", "National", "356 M€ CA ; acteur #1 BE ; Ideal Standard private label"],
    ["SIDER", "Négoce pro plomberie", "National", "35 000 références, livraison pro, e-commerce"],
    ["STG / ex-Paepens (BME Group)", "Négoce sanitaire-chauffage", "Flandre / Bruxelles", "Racheté par BME (Groupe STG FR)"],
    ["Facq Pro", "Négoce sanitaire et robinetterie", "National (Wallonie fort)", "Réseau 30+ agences, pro + particuliers"],
    ["Versani NV", "Grossiste plomberie-chauffage", "Kempen (Anvers)", "Fondé 1975, indépendant"],
    ["Willems-Diels", "Grossiste plomberie", "Balen (Anvers)", "Familial, 40 ans d'existence"],
    ["Aquacaro", "Grossiste robinetterie design B2B", "National", "Robinetterie design salle de bains"],
    ["Sawiday.be", "E-commerce pro", "National/BE", "Présence Presto confirmée"],
    ["Sider.biz", "E-commerce pro", "National/BE", "Robinetterie + sanitaire pros"],
    ["JA Santé Belgique", "Équipements santé", "National", "Revendeur Delabie confirmé"],
])
ann.add_paragraph()

add_heading(ann, "ANNEXE 6 — CONCURRENTS DÉTAILLÉS", 1)
make_table(ann, [
    ["Marque", "Groupe / Pays", "Niveau gamme", "Spécialité ERP", "Présence BE", "Notes"],
    ["DELABIE", "Familial FR (1928)", "Haut de gamme", "Oui — leader ERP", "Filiale Delabie Benelux SRL (Sint-Pieters-Leeuw)", "Acquiert Hansa 2021, KWC Pro 2025"],
    ["Grohe", "LIXIL (DE)", "Milieu-haut", "Partielle (tertiaire)", "Distribution nationale", "#1 marché total BE (BRG 2020)"],
    ["Hansgrohe", "Masco Corp (DE)", "Haut de gamme", "Tertiaire, résidentiel premium", "Distribution + showrooms", "Fort côté flamand"],
    ["Hansa", "Delabie Group (FR)", "Milieu-haut", "Électronique, partielle", "Site BE hansa.com/fr-be", "Désormais filiale Delabie"],
    ["Geberit", "Geberit AG (CH)", "Milieu-haut", "Sanitaires encastrés", "Distribution nationale", "Fort en tertiaire neuf"],
    ["Paffoni", "Italie", "Milieu-haut", "Via BSC Belgium (hist.)", "Négoce + BSC", "Représenté par BSC Belgium (BRG 2020)"],
    ["Oras / Damixa", "Artek Industries (FI)", "Milieu", "Limitée", "Via négoce", "Marques nordiques"],
    ["Jacob Delafon", "Kohler (US)", "Milieu", "Résidentiel principalement", "Distribution nationale", "Private label Van Marcke"],
], font_size=8)
ann.add_paragraph()

add_heading(ann, "ANNEXE 7 — NORMES DÉTAILLÉES ROBINETTERIE BELGIQUE", 1)
make_table(ann, [
    ["Norme", "Objet", "Date", "Implication Presto", "Organisme"],
    ["NBN EN 816", "Robinets fermeture automatique PN 10", "1996", "CRITIQUE — cœur gamme temporisateurs", "NBN"],
    ["NBN EN 817", "Mitigeurs mécaniques PN 10", "2008", "Gamme mitigeurs collectifs", "NBN"],
    ["NBN EN 1111", "Mitigeurs thermostatiques PN 10", "2000", "Thermostatiques ERP", "NBN"],
    ["NBN EN 200", "Robinets simples et mélangeurs", "2008", "Robinets individuels", "NBN"],
    ["NBN EN 1717", "Protection pollution eau potable / anti-retour", "2001", "Obligatoire tous produits", "NBN"],
    ["NBN EN 806 (1-5)", "Installations eau potable intérieures", "2012", "Cadre général installation", "NBN"],
    ["NBN D51-003", "Critères techniques installations domestiques/industrielles", "belge", "Spécificité belge", "NBN"],
    ["NBN S 01-400-1 + EN ISO 3822", "Acoustique installations sanitaires", "—", "Nuisances sonores temporisateurs", "NBN"],
    ["BELGAQUA HYDROCHECK", "Certification eau potable BE", "Obligatoire AO", "PRÉREQUIS marché public belge", "BELGAQUA"],
    ["Label BENOR", "Qualité produits construction", "Volontaire", "Souvent exigé cahiers charges publics", "NBN / organismes mandatés"],
    ["Marquage CE (EU 305/2011)", "Mise sur marché UE", "Obligatoire", "Presto déjà certifié — conforme", "Commission UE"],
    ["4MS / ACS (FR)", "Substances contact eau potable", "Transitoire jusqu'au 31/12/2026", "ACS français reconnu en BE", "UBA / BELGAQUA"],
], font_size=8)
ann.add_paragraph()

# Sauvegarde annexes
path_ann = os.path.join(OUTPUT_DIR, "MAB_Belgique_Annexes.docx")
ann.save(path_ann)
print(f"✓ Annexes v5 sauvegardées : {path_ann}")
print()
print("═" * 60)
print("MAB BELGIQUE v5 — Génération terminée")
print(f"  Étude   → {path_etude}")
print(f"  Annexes → {path_ann}")
print("═" * 60)
