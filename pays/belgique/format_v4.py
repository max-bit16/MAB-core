"""
Post-traitement mise en page v4 — MAB Belgique
Source  : pays/belgique/outputs/MAB_Belgique_Etude.docx  (= v3 contenu)
Cible   : pays/belgique/outputs/MAB_Belgique_EtudeV4_mise_en_page.docx

Changements appliqués (contenu inchangé) :
  1. Police Calibri partout
  2. Orange → gris #595959 ; bleus titres → gris #595959
  3. Espacement augmenté avant/après paragraphes et titres
  4. Marges 2,5 cm sur les 4 côtés
  5. Fond #EBEBEB sur les lignes d'en-tête de tous les tableaux
  6. Saut de page obligatoire avant chaque PARTIE X (Heading 1 "PARTIE …")
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), "outputs")

JOBS = [
    {
        "src": os.path.join(OUT_DIR, "MAB_Belgique_Etude.docx"),
        "dst": os.path.join(OUT_DIR, "MAB_Belgique_EtudeV4_mise_en_page.docx"),
        "break_prefix": "PARTIE",
    },
    {
        "src": os.path.join(OUT_DIR, "MAB_Belgique_Annexes.docx"),
        "dst": os.path.join(OUT_DIR, "MAB_Belgique_AnnexesV4.docx"),
        "break_prefix": "ANNEXE",
    },
]

GRAY    = RGBColor(0x59, 0x59, 0x59)
FONT    = "Calibri"

# Couleurs à remplacer → gris
REPLACE_COLORS = {
    (0xC0, 0x50, 0x20),   # orange des notes
    (0x1F, 0x49, 0x7D),   # bleu H1
    (0x2E, 0x75, 0xB6),   # bleu H2
    (0x00, 0x56, 0xA2),   # bleu URLs annexes
}

def set_cell_bg(cell, hex_fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tcPr.append(shd)

def fix_run(run):
    run.font.name = FONT
    try:
        if run.font.color.type is not None:
            rgb = run.font.color.rgb
            if tuple(rgb) in REPLACE_COLORS:
                run.font.color.rgb = GRAY
    except Exception:
        pass

def fix_para(para, is_heading=False):
    for run in para.runs:
        fix_run(run)
    pf = para.paragraph_format
    if is_heading:
        pf.space_before = Pt(14)
        pf.space_after  = Pt(6)
    else:
        if pf.space_before is None or pf.space_before < Pt(4):
            pf.space_before = Pt(4)
        if pf.space_after is None or pf.space_after < Pt(4):
            pf.space_after  = Pt(4)

def process(src, dst, break_prefix):
    doc = Document(src)

    # 1. Marges 2,5 cm
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # 2. Paragraphes + sauts de page avant sections principales
    for para in doc.paragraphs:
        is_h = para.style.name.startswith("Heading")
        fix_para(para, is_heading=is_h)
        if is_h and para.text.strip().upper().startswith(break_prefix.upper()):
            para.paragraph_format.page_break_before = True

    # 3. Tableaux
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT
                        try:
                            if run.font.color.type is not None:
                                rgb = run.font.color.rgb
                                if tuple(rgb) in REPLACE_COLORS:
                                    run.font.color.rgb = GRAY
                        except Exception:
                            pass
                if i == 0:
                    set_cell_bg(cell, "EBEBEB")

    # 4. Styles du document
    for style in doc.styles:
        try:
            if style.font:
                style.font.name = FONT
        except Exception:
            pass
        try:
            if style.name in ("Heading 1", "Heading 2", "Heading 3"):
                pf = style.paragraph_format
                pf.space_before = Pt(14)
                pf.space_after  = Pt(6)
                if style.font.color.type is not None:
                    if tuple(style.font.color.rgb) in REPLACE_COLORS:
                        style.font.color.rgb = GRAY
        except Exception:
            pass

    doc.save(dst)
    print(f"✓ Sauvegardé : {dst}")

# ── Traitement des deux fichiers ─────────────────────────────────────────────
for job in JOBS:
    process(job["src"], job["dst"], job["break_prefix"])
