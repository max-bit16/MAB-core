"""
MAB Suisse v1 — Génération des deux documents Word
Sources internes utilisées :
  - MAB_Suisse_PREP.md (corpus PREP, juin 2026)
  - Analyse Marché Sanitaire Lieux public France.pdf (base extrapolation FR, déc. 2024)
  - EMAE - Extrapolation, notes de recherche.docx (données construction CH)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

FONT = "Calibri"

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = FONT
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def bold_bullet(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.5)
    r1 = p.add_run(f"{label} : ")
    set_font(r1, bold=True)
    r2 = p.add_run(value)
    set_font(r2)
    return p

def add_para(doc, text, size=10, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, italic=italic)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {text}")
    set_font(run, size=9, italic=True, color=(0xC0, 0x50, 0x20))
    return p

def make_table(doc, rows_data, font_size=9, header_color=(0x1F, 0x49, 0x7D)):
    tbl = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    tbl.style = "Table Grid"
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.name = FONT
            run.font.size = Pt(font_size)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(*header_color)
    return tbl

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
set_margins(doc)

t = doc.add_heading("MAB SUISSE — ÉTUDE DE MARCHÉ v1", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Robinetterie sanitaire collective / ERP — Les Robinets Presto")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(12)
sub.runs[0].font.name = FONT
date_p = doc.add_paragraph("Juin 2026")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(10)
date_p.runs[0].font.italic = True
date_p.runs[0].font.name = FONT
doc.add_paragraph()

# ─── RÉSUMÉ EXÉCUTIF ─────────────────────────────────────────────────────────
add_heading(doc, "RÉSUMÉ EXÉCUTIF", 1)
add_bullet(doc, "Marché robinetterie ERP Suisse estimé à 27–33 M€ (robinetterie collective seule, estimation par extrapolation ajustée — Méthode A). Marché premium à très fort pouvoir d'achat (PIB/hab 89 783 USD — 1,8× la France) : exigences qualité élevées, prix acceptés, marges potentiellement supérieures à celles réalisées en France. Droits de douane HS 8481 supprimés depuis janvier 2024 — accès sans taxe pour produits français. (Worldometers 2024 / kmu.admin.ch)")
add_bullet(doc, "Construction résiliente (CHF 68,9 Mds, +2,1% en 2025) avec la rénovation désormais majoritaire (57% du total). Pipelines ERP publics sécurisés sur 5-10 ans : 528 M CHF/an rénovation énergétique (Programme Bâtiments), 8,5 Mds CHF plan ferroviaire CFF 2035, >1,5 Mds CHF programmes hospitaliers HUG+CHUV, 560 M CHF CAP2030 aéroport Genève. (SSE / Eminence.ch / news.admin.ch / CFF / HUG)")
add_bullet(doc, "MENACE CONCURRENTIELLE MAJEURE — Delabie a finalisé l'acquisition de KWC Professional (Unterkulm, Argovie) en Q3 2025 : KWC était la marque suisse historique de robinetterie ERP (hôpitaux, écoles, prisons, transports). Delabie dispose désormais d'une production locale en Suisse et d'un réseau installé (marque Aquarotter, leader germanophone ERP). Geberit (siège Rapperswil-Jona, CH) domine le marché sanitaire global. Presto = 3 commerciaux terrain, pas de filiale — à renforcer impérativement face à ces acteurs. (delabie.com / geberit.ch, 2025)")
add_bullet(doc, "Top 3 segments prioritaires : Santé/EMS (5/5 — 278 hôpitaux + 1 465 EMS + vieillissement démographique fort + pipeline HUG+CHUV >1,5 Mds CHF), Transports (5/5 — 764 gares CFF + programme accessibilité PMR >2,5 Mds CHF + aéroport Genève CAP2030 560 M CHF), CHR (4/5 — 42,8 M nuitées 2024 record absolu, ~5 000 hôtels, marché premium). (OFS 2023-2024 / CFF / GVA newsroom)")
add_bullet(doc, "Atout Suisse romande : 26% de la population est francophone (Genève, Vaud, Valais, Neuchâtel) — marché d'entrée naturel pour Presto, documentation déjà disponible, proximité culturelle forte. La France est le 4e fournisseur de la Suisse (7,9% des imports hors or) et 1 706 filiales françaises sont implantées. (DG Trésor 2026)")
doc.add_paragraph()

# ─── PARTIE 1 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 1 — OVERVIEW CONTEXTE PAYS", 1)

add_heading(doc, "1.1 Analyse PESTEL", 2)
bold_bullet(doc, "POLITIQUE", "État fédéral à 26 cantons — stabilité politique exemplaire, démocratie directe (référendums fréquents). Gouvernement fédéral collégial (Conseil fédéral, 7 membres). Hors UE mais accords bilatéraux CH-UE en vigueur. Neutralité internationale. Décisions d'investissement public stables et prévisibles sur le long terme. (admin.ch, 2025)")
bold_bullet(doc, "ÉCONOMIQUE", "PIB ~900 Mds USD (2024). PIB/hab : 89 783 USD — 1,8× France, 3e rang mondial (Worldometers/Trading Economics, 2024). Croissance PIB : +1,4% (2024 et 2025 est.). Inflation maîtrisée : ~1,1% (OFS 2024). Chômage : 2,3% — quasi plein emploi (SECO, 2024). Franc suisse très fort : 1 CHF = 1,05 EUR (moy. 2024). Économie formelle quasi totale.")
bold_bullet(doc, "SOCIAL", "Population 8,9 M habitants (2024) — 9,0 M en 2026 (OFS / Worldometers). Urbanisation : 74%, 52 agglomérations, 6,6 M habitants urbains (Union des villes suisses 2024). 26% francophones (Suisse romande). Vieillissement accéléré : +1,2%/an de personnes âgées → pression EMS. 4 langues nationales (DE 63%, FR 23%, IT 8%, romanche 1%).")
bold_bullet(doc, "TECHNOLOGIQUE", "Forte adoption BIM : 62% (Eminence.ch 2025). Smart building : CHF 1,35 Md de marché actuel, croissance. Robinetterie infrarouge/connectée favorisée dans ERP neufs. Rang 2 mondial innovation (GII 2024). ETH Zurich et EPFL : centres de recherche de premier plan.")
bold_bullet(doc, "ENVIRONNEMENTAL", "Stratégie énergie 2050 (sortie nucléaire progressive, renouvelables). Neutralité carbone 2050 inscrite dans la Constitution depuis 2023. Programme Bâtiments : 528 M CHF versés 2024 pour rénovation énergétique. Taux rénovation actuel 1,4%/an vs objectif 3%/an — fort retard = fort potentiel. OPBD : réglementation eau potable très stricte, interdiction du plomb depuis 1904. (news.admin.ch)")
bold_bullet(doc, "LÉGAL", "Hors UE → marquage CE insuffisant pour eau potable. Certification SVGW (ZertW) obligatoire de facto pour robinetterie en contact avec l'eau potable. Norme SIA 500 (accessibilité PMR) en révision (prSIA 500:2025). Directive W3 SVGW pour installations eau potable. Droits de douane HS 8481 : 0% depuis janvier 2024. (svgw.ch / kmu.admin.ch)")
doc.add_paragraph()

add_heading(doc, "1.2 Indicateurs socio-économiques clés", 2)
make_table(doc, [
    ["Indicateur", "Valeur", "Source"],
    ["PIB total", "~900 Mds USD (2024)", "Worldometers / OFS 2024"],
    ["PIB/habitant", "89 783 USD (2024)", "Trading Economics / Worldometers"],
    ["Population", "8,9 M (2024) — 9,007 M (2026)", "OFS / Worldometers"],
    ["Urbanisation", "74% (52 agglomérations)", "Union villes suisses / OFS 2024"],
    ["Croissance PIB", "+1,4% (2024 et 2025 est.)", "SECO 2024-2025"],
    ["Inflation", "~1,1% (2024)", "OFS 2024"],
    ["Chômage", "2,3%", "SECO 2024"],
    ["Taux de change", "1 CHF = 1,05 EUR (moy. 2024)", "CDTF / Banque de France"],
    ["Population >65 ans", "~19,5% (en hausse +1,2%/an)", "OFS 2024"],
    ["Rang innovation", "2e mondial (GII 2024)", "Global Innovation Index 2024"],
])
doc.add_paragraph()

add_heading(doc, "1.3 Relations commerciales France-Suisse", 2)
add_bullet(doc, "Exports France → Suisse (biens hors or) : 19,8 Mds EUR (2024) → 22,6 Mds EUR (2025, +14%). Solde commercial France : +2,3 Mds EUR (2024) → +4,7 Mds EUR (2025). (Direction du Trésor, 2026)")
add_bullet(doc, "Suisse = 9e partenaire commercial de la France. France = 4e fournisseur de la Suisse (7,9% des imports hors or). (DG Trésor)")
add_bullet(doc, "Présence française en Suisse : 964 groupes via 1 706 filiales, ~76 755 emplois directs, IDE stock 48,2 Mds CHF. Grands employeurs : Bouygues (6 300 emp.), Vinci Energies (3 260 emp.), AXA, Saint-Gobain. (DG Trésor, données 2021)")
add_bullet(doc, "Atouts origine française : Suisse romande (26% population) = marché d'entrée francophone naturel. Image qualité France bien perçue. Accords bilatéraux CH-UE facilitent les échanges. Droits de douane HS 8481 supprimés depuis 2024.")
add_bullet(doc, "Presto en Suisse : quelques commerciaux terrain (pas de filiale constituée). Présence à structurer pour répondre à la demande ERP sur un marché très compétitif.")
add_bullet(doc, "Freins pour Presto : 3 langues de marché (FR + DE + IT). Concurrents locaux très forts (Geberit siège suisse ; Delabie-KWC fabrication locale post-acquisition 2025). Prix en CHF : tarification à adapter (marché premium). Certification SVGW à sécuriser si non encore obtenue.")
doc.add_paragraph()

add_heading(doc, "1.4 Tendances d'investissement — Programmes clés", 2)
make_table(doc, [
    ["Programme", "But global", "Secteur", "Budget", "Calendrier", "Opportunité Presto (types produits)"],
    ["Programme Bâtiments (rénovation énergétique)", "Réduction des émissions CO2 du parc bâti résidentiel et tertiaire suisse", "Tous bâtiments", "528 M CHF/an (247 M fédéral + 275 M cantons, 2024)", "Continu 2024-2030+", "Économiseurs d'eau, temporisateurs lors rénovations sanitaires"],
    ["Plan ferroviaire CFF 2035", "Extension de l'offre ferroviaire nationale ; +20% places assises sur 60 lignes", "Transports / gares", "8,5 Mds CHF", "2024-2035", "Robinetterie temporisée gares ; infrarouge ; PMR"],
    ["Programme accessibilité CFF (PMR)", "Adaptation des 764 gares aux normes d'accessibilité pour personnes à mobilité réduite", "Transports", ">2,5 Mds CHF d'ici 2028", "2024-2028", "Robinetterie PMR, temporisée, infrarouge sans contact"],
    ["HUG Genève — rénovation et extension hospitalière", "Modernisation plateau technique, maternité, psychiatrie (EviPsy), hôpital des enfants", "Santé", ">58,5 M CHF/an ; 1,5 Mds CHF HUG+CHUV", "2024-2031", "Temporisée hospitalière, infrarouge, anti-brûlure, PMR, anti-arrachement psychiatrie"],
    ["CAP2030 — Aéroport Genève", "Rénovation du terminal principal et construction du Satellite 10", "Transports", "560 M CHF 2023-2032", "2025-2032", "Temporisée, infrarouge, économiseurs eau, PMR aéroport"],
    ["FORTA (fonds routes nationales)", "Entretien et développement du réseau routier national et des agglomérations", "Infrastructures / aires de repos", "~3 Mds CHF/an ; 11,6 Mds CHF à 2030", "Continu", "Temporisée aires de repos et stations-service"],
    ["Budget FRI 2025-2028 (formation-recherche)", "Financement des EPF, universités et HES suisses", "Éducation", "29,2 Mds CHF (coupes 460 M CHF/an)", "2025-2028", "Robinetterie établissements académiques, résidences étudiantes, infrarouge"],
], font_size=8)
doc.add_paragraph()

# ─── PARTIE 2 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 2 — MARCHÉ DE LA CONSTRUCTION", 1)

add_heading(doc, "2.1 État et taille du marché", 2)
add_bullet(doc, "Investissement total construction Suisse 2024 : CHF 68,9 Mds (+1,8% YoY), soit ~9,5% du PIB. (Eminence.ch, Swiss Construction Insights 2025)")
add_bullet(doc, "CA secteur principal (entrepreneurs) : CHF 23,4 Mds en 2024. Prévisions : CHF 23,9 Mds (+2,1% en 2025) et CHF 24,4 Mds (+1,9% en 2026). (SSE / Baumeister.swiss)")
add_bullet(doc, "Carnets de commandes : 7,4 mois de travail assuré fin 2024 — niveau le plus élevé depuis 6 ans. (SIA Q4 2024)")
add_bullet(doc, "Emplois secteur : 352 000 personnes (8,1% du travail national). (Eminence.ch)")
add_bullet(doc, "Marché robuste, bonne résilience face aux turbulences européennes. Croissance progressive sans surchauffe. (Source interne EMAE, 2025)")

add_heading(doc, "2.2 Dynamique Neuf vs Rénovation", 2)
make_table(doc, [
    ["Segment", "Poids estimé", "Dynamique 2025", "Source"],
    ["Neuf résidentiel", "~38% du total", "+4,8% (2025)", "EMAE / Eminence.ch 2025"],
    ["Rénovation résidentielle (énergie)", "~19% du total", "Forte hausse (soutien réglementaire)", "news.admin.ch / Programme Bâtiments"],
    ["Non-résidentiel (neuf + rénov.)", "~25-30% du total", "Stable — pas de croissance attendue", "EMAE / Eminence.ch 2025"],
    ["Génie civil / Infrastructures", "~15-20% du total", "Stable (investissements fédéraux CFF/FORTA)", "EMAE 2025"],
])
add_note(doc, "Estimation — répartition neuf/réno basée sur EMAE (source interne, 2025) et Eminence.ch. La rénovation atteint 57% du total (CHF 39,2 Mds vs CHF 29,3 Mds neuf). Décomposition sous-segments à confirmer via OFS (bfs.admin.ch).")
doc.add_paragraph()
add_bullet(doc, "Rénovation = 57% de l'investissement total (CHF 39,2 Mds) vs neuf 43% (CHF 29,3 Mds) — la rénovation dépasse le neuf pour la première fois. 59% du parc résidentiel date d'avant 1980. (Eminence.ch 2024)")
add_bullet(doc, "Résidentiel neuf : +4,8% en 2025, moteur principal. 50 000 logements/an nécessaires vs 42 000-45 000 construits réellement. 82% des projets = logements collectifs. Pénurie structurelle persistante. (EMAE 2025)")
add_bullet(doc, "Rénovation énergétique : forte hausse, 2e pilier de croissance. Taux actuel 1,4%/an, objectif net zéro 2050 = 3%/an. Programme Bâtiments : 528 M CHF versés en 2024. (news.admin.ch)")
add_bullet(doc, "Régions dynamiques : Zurich, Suisse centrale, Lac Léman. Bâle : accent rénovation et modernisation énergétique. (EMAE 2025)")

add_heading(doc, "2.3 Perspectives 2025-2030", 2)
add_bullet(doc, "Croissance construction attendue : +1,9% (2026) puis ~+2%/an selon SSE. Marché robuste, résilient, sans surchauffe malgré turbulences européennes. (SSE 2025)")
add_bullet(doc, "Résidentiel : pénurie structurelle (50 000 logements/an nécessaires, ~42 000 construits) → moteur de croissance durable à horizon 2030. 82% logements collectifs = marchés robinetterie collective par projet.")
add_bullet(doc, "Rénovation énergétique : tendance lourde irréversible. Stratégie énergie 2050 génère un cycle de rénovation du parc bâti sur 25 ans. Objectif : 3%/an vs 1,4% actuel → potentiel de multiplication ×2 du rythme. (news.admin.ch)")
add_bullet(doc, "BIM : 62% d'adoption (Eminence.ch 2025) → prescripteurs (architectes, ingénieurs SIA) très influents dans le choix des équipements sanitaires. Référencement BIM de Presto à développer.")
add_bullet(doc, "Smart building et robinetterie connectée : marché existant CHF 1,35 Md, en croissance → robinetterie infrarouge et économiseurs eau avec comptage connecté favorisés dans les ERP neufs.")
add_bullet(doc, "Non-résidentiel : prudence maintenue, rénovation prime sur neuf. Bureaux en légère surcapacité (5% vacance). Pipelines publics (santé, éducation, transports) sécurisent la demande ERP.")
add_bullet(doc, "Risques : pression sur les coûts (main-d'œuvre, matériaux), CHF fort pouvant pénaliser la compétitivité, investissements non-résidentiels reportés (EMAE 2025).")
doc.add_paragraph()

# ─── PARTIE 3 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 3 — CONSTRUCTION NON-RÉSIDENTIELLE", 1)

add_heading(doc, "3.1 État actuel", 2)
add_bullet(doc, "Non-résidentiel Suisse : estimé à 25-30% de l'investissement total = CHF 17-21 Mds (2024). (Estimation — décomposition exacte à confirmer via OFS bfs.admin.ch)")
add_bullet(doc, "Pas de croissance attendue en 2025 ; investissements non-résidentiels reportés dus à la pression des coûts. (EMAE / Eminence.ch 2025)")
add_bullet(doc, "Bureaux (5 grandes villes) : 995 500 m² disponibles fin 2024, taux de vacance 5,0% (+0,9 pt YoY). Légère surcapacité → rénovation prioritaire sur neuf. (CBRE / immoday.ch 2024)")
add_bullet(doc, "Santé : CHF 58,5 M investis aux HUG seuls en 2024 ; programme HUG+CHUV >1,5 Mds CHF engagés sur 2024-2031. (hug.ch)")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — part exacte non-résidentiel en valeur CHF/EUR 2024 — à sourcer via OFS bfs.admin.ch/statistiques/construction-logement]")

add_heading(doc, "3.2 Dynamique Neuf vs Rénovation — non-résidentiel", 2)
make_table(doc, [
    ["Sous-segment", "Dynamique", "Remarque"],
    ["Bureaux / tertiaire neuf", "Stable-décroissant (vacance 5%)", "Légère surcapacité → rénovation prime"],
    ["Santé (hôpitaux, EMS)", "Croissance soutenue", "Vieillissement démographique + pipelines HUG/CHUV"],
    ["Éducation (universités, EPF, HES)", "Stable (coupes FRI 460 M CHF/an)", "Rénovation primaire ; neuf résidences étudiantes"],
    ["Industrie / logistique", "Faible croissance", "Pharma (Bâle), horlogerie (Jura/Neuchâtel)"],
    ["Transports (gares, aéroports)", "Forte croissance (pipelines engagés)", "CFF 8,5 Mds CHF ; CAP2030 Genève 560 M CHF"],
])
add_note(doc, "Estimation dynamique neuf/rénov. non-résidentiel basée sur sources EMAE, CBRE, hug.ch — à confirmer via OFS.")
doc.add_paragraph()

add_heading(doc, "3.3 Segments dominants du non-résidentiel", 2)
add_bullet(doc, "Santé et EMS : 1er segment non-résidentiel par dynamique. 278 hôpitaux (OFS 2022), 1 465 EMS/100 540 places (OFS SOMED 2023). Vieillissement démographique → +1,2%/an de personnes âgées. Pipeline HUG+CHUV >1,5 Mds CHF. Rénovation très majoritaire (parc hospitalier ancien).")
add_bullet(doc, "Transports : 2e segment par budget engagé. CFF : 764 gares, 8,5 Mds CHF plan 2035, programme accessibilité PMR >2,5 Mds CHF. Aéroport Genève : CAP2030 560 M CHF 2025-2032. (CFF / newsroom.gva.ch)")
add_bullet(doc, "Tertiaire / Bureaux : 3e segment. 5 grandes villes actives mais légère surcapacité (5% vacance). Demande rénovation dominante. Niche premium : sièges de multinationales (Nestlé, Novartis, ABB, Roche).")
add_bullet(doc, "Éducation : segment stable, rénovation principale. ETH Zurich, EPFL, 37 HES + universités cantonales. ~11 700 établissements scolaires (données 2017/18 — à actualiser OFS).")
add_bullet(doc, "Industrie : pharma (Bâle), horlogerie (Jura, Neuchâtel). Sanitaires vestiaires/cantines — segment secondaire pour Presto.")

add_heading(doc, "3.4 Perspectives 2025-2030", 2)
add_bullet(doc, "Santé et EMS : croissance structurelle portée par vieillissement démographique. EFAS 2028 (nouveau mode financement soins) à surveiller. Pressions sur capacité : 97% des prisons et saturation EMS prévisible horizon 2030-2035.")
add_bullet(doc, "Transports : pipelines CFF et aéroports = visibilité 2025-2035. Lots sanitaires gares CFF incluent systématiquement robinetterie temporisée/PMR/infrarouge lors rénovations.")
add_bullet(doc, "Éducation : budget FRI 2025-2028 (29,2 Mds CHF) mais coupes 460 M CHF/an → croissance limitée sur constructions neuves. Rénovation reste active.")
add_bullet(doc, "Rénovation énergétique non-résidentiel : Programme Bâtiments inclut le tertiaire → vague de rénovation sanitaire lors des travaux énergétiques à horizon 2025-2030.")
add_bullet(doc, "Smart building : 62% d'adoption BIM — les maîtres d'ouvrage intègrent des critères de performance eau dans les ERP. Robinetterie infrarouge et connectée favorisée.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — CAGR non-résidentiel Suisse 2025-2030 — aucune source directe identifiée]")
doc.add_paragraph()

# ─── PARTIE 4 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 4 — POIDS DES SEGMENTS ERP", 1)

add_heading(doc, "4.0 Contexte général et synthèse ERP Suisse", 2)
add_para(doc, "La Suisse dispose d'un parc ERP de haute qualité, dont la majorité a été construite avant 1980 et nécessite une rénovation sanitaire. Le marché est caractérisé par des exigences techniques très élevées, des budgets publics stables et une culture de l'entretien rigoureux.")
add_bullet(doc, "Parc ERP ancien : 59% du bâti construit avant 1980 (Eminence.ch) → cycle de rénovation sanitaire important et structurel.")
add_bullet(doc, "Exigences qualité premium : le cahier des charges ERP suisse impose en règle générale des produits haut de gamme. Durabilité, économie eau et accessibilité PMR sont des critères standard.")
add_bullet(doc, "Vieillissement démographique : +1,2%/an de personnes âgées → pression sur EMS et hôpitaux à horizon 2030-2040.")
add_bullet(doc, "Pipelines publics sécurisés : CFF, HUG/CHUV, Programme Bâtiments — visibilité d'investissement sur 5-10 ans.")
doc.add_paragraph()

make_table(doc, [
    ["Segment ERP", "Score Presto", "Taille / Signal", "Priorité"],
    ["4.2 Santé / Hôpitaux / EMS", "5/5", "278 hôpitaux + 1 465 EMS + pipeline HUG+CHUV >1,5 Mds CHF", "PRIORITÉ ABSOLUE"],
    ["4.11 Transports", "5/5", "764 gares CFF + CAP2030 Genève 560 M CHF + PMR 2,5 Mds CHF", "PRIORITÉ ABSOLUE"],
    ["4.5 CHR / Hôtels", "4/5", "42,8 M nuitées 2024 (record) ; ~5 000 établissements", "PRIORITÉ HAUTE"],
    ["4.1 Éducation", "4/5", "ETH + EPFL + 37 HES + ~11 700 établissements", "PRIORITÉ HAUTE"],
    ["4.7 Sport & Loisirs", "3/5", "32 000 installations ; ~1 000 centres fitness", "SECONDAIRE"],
    ["4.3 Tertiaire / Bureaux", "3/5", "Stable ; niche premium multinationales", "SECONDAIRE"],
    ["4.8 Pénitentiaire", "3/5", "90 prisons ; 97% occupation ; besoin rénov./construction", "SECONDAIRE"],
    ["4.4 Industrie", "2/5", "Pharma/horlogerie ; vestiaires/cantines", "OPPORTUNISTE"],
    ["4.6 HPA (campings)", "2/5", "Données insuffisantes ; tourisme montagne", "OPPORTUNISTE"],
    ["4.9 Culturel", "2/5", "Musées réputés mais dispersés", "OPPORTUNISTE"],
    ["4.10 Lieux de culte", "1/5", "Sécularisation accélérée ; -7% groupes depuis 2008", "NON PRIORITAIRE"],
])
doc.add_paragraph()

add_heading(doc, "4.1 Éducation", 2)
add_bullet(doc, "~11 700 établissements scolaires (primaire + secondaire I, données OFS 2017/18 — à actualiser). 10 universités cantonales + ETH Zurich (23 500 étudiants) + EPFL (12 000 étudiants) + 37 HES. (OFS 2024 / ETH / EPFL)")
add_bullet(doc, "Taux scolarisation ~99% → parc établissements stable en volume mais rénovation active. Budget FRI 2025-2028 : 29,2 Mds CHF mais coupes de 460 M CHF/an (Conseil fédéral).")
add_bullet(doc, "Résidences étudiantes : développement actif (Lausanne/EPFL, projets privés) → segment croissant. Dynamique : rénovation > neuf pour primaire/secondaire.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — volume marché robinetterie éducation/an en Suisse]")

add_heading(doc, "4.2 Santé / Hôpitaux / EMS", 2)
add_bullet(doc, "Hôpitaux : 278 établissements (264 hôpitaux + 14 maternités). Lits totaux : ~37 925-38 000 (OFS 2023). Lits/1 000 hab : 4,4-4,6 (OCDE) vs France 5,9. Facteur santé = 4,5/5,9 = 0,76. (OFS 2022-2023 / OCDE)")
add_bullet(doc, "EMS (établissements médico-sociaux) : 1 465 établissements ; 100 540 places ; 170 211 résidents ; coûts 11,65 Mds CHF. (OFS SOMED 2023 / Senesuisse)")
add_bullet(doc, "Pipeline HUG Genève : rénovation maternité 2024-2030, nouvel hôpital psychiatrique (EviPsy), hôpital des enfants → CHF 58,5 M investis en 2024. (hug.ch)")
add_bullet(doc, "Vieillissement démographique → fort accroissement demande EMS à horizon 2030-2040. EFAS 2028 = nouveau système financement hospitalier-ambulatoire à surveiller.")
add_bullet(doc, "Neuf : nouveaux hôpitaux rares (HUG Psychiatrie EviPsy). Rénovation : très majoritaire sur parc existant.")

add_heading(doc, "4.3 Bâtiments tertiaires (Bureaux)", 2)
add_bullet(doc, "995 500 m² disponibles (5 grandes villes) fin 2024, taux vacance 5,0% (CBRE / immoday.ch 2024). Légère surcapacité → rénovation prioritaire sur neuf.")
add_bullet(doc, "Niche premium : sièges de multinationales (Nestlé, Novartis, ABB, Roche) = cahiers des charges exigeants. Canaux de prescription : architectes + BET.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — volume marché robinetterie tertiaire/an]")

add_heading(doc, "4.4 Bâtiments industriels", 2)
add_bullet(doc, "Secteurs clés : pharma (Bâle — Novartis, Roche), horlogerie (Jura, Neuchâtel, Vallée de Joux), machines-outils, medtech.")
add_bullet(doc, "Sanitaires vestiaires, cantines, salles blanches pharma (exigences hygiéniques très élevées). Segment secondaire pour Presto en volume, mais valeur unitaire élevée.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — nombre établissements industriels, volume robinetterie/an]")

add_heading(doc, "4.5 CHR — Cafés, Hôtels, Restaurants", 2)
add_bullet(doc, "~5 000 hôtels et établissements (OFS / HotellerieSuisse 2024). 42,8 millions nuitées hôtelières 2024 (+2,6% vs 2023) — record absolu. 22,0 millions nuitées étrangères (+5,1%). (OFS 2024)")
add_bullet(doc, "Recettes touristiques : 20,33 Mds EUR (Fédération suisse du tourisme 2024). Exigences qualité premium → robinetterie design, durabilité, anti-calcaire. Marché très favorable aux produits haut de gamme.")
add_bullet(doc, "Facteur CHR = 22 M nuitées étrangères CH / 102 M arrivées France ≈ 0,216. Estimation marché CHR corrigée de ce facteur par rapport à la France.")

add_heading(doc, "4.6 HPA — Hôtellerie Plein Air, Campings", 2)
add_bullet(doc, "[DONNÉE NON DISPONIBLE — nombre campings / HPA Suisse non identifié. TCS gère un réseau de campings suisses avec sanitaires modernisés.]")
add_bullet(doc, "Segment de taille limitée mais qualité sanitaires attendue élevée (marché suisse premium). Tourisme montagne = clientèle internationale exigeante.")

add_heading(doc, "4.7 Centres Sport & Loisirs", 2)
add_bullet(doc, "32 000 installations sportives en Suisse (EDA / aboutswitzerland.eda.admin.ch). ~980-1 000 centres de fitness ; CA CHF 1,3 Md en 2024 (+7,6%) ; 1,37 million de membres. (LFM / Watson.ch 2024)")
add_bullet(doc, "75% de la population pratique le sport régulièrement (EDA). Stades, piscines cantonales, gymnases communaux.")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — dépenses publiques sport % PIB Suisse (Eurostat ne couvre pas CH)]")

add_heading(doc, "4.8 Établissements à sécurité renforcée — Pénitentiaire", 2)
add_bullet(doc, "90 établissements pénitentiaires ; 7 373 places déclarées. 6 881 détenus au 31/01/2024 ; 7 119 début 2026 — record depuis 1988. Taux d'occupation : 97% (2026). (Prison Insider / RTS 2024)")
add_bullet(doc, "Taux incarcération : 78/100 000 hab (vs France 111/100 000). Facteur pénitentiaire = 78/111 = 0,70. 72% de détenus étrangers → spécificité suisse.")
add_bullet(doc, "Saturation → besoin de construction/rénovation mais décisions politiques complexes (vote populaire possible). Hôpitaux psychiatriques (EviPsy HUG) = segment additionnel anti-arrachement.")

add_heading(doc, "4.9 Bâtiments culturels", 2)
add_bullet(doc, "[DONNÉE NON DISPONIBLE — nombre cinémas, théâtres, musées Suisse non disponible via recherches]")
add_bullet(doc, "Musées réputés : Kunsthaus Zurich (rénové 2021), fondations Lausanne, musées cantonaux. Segment dispersé, commandes ponctuelles.")

add_heading(doc, "4.10 Lieux de culte", 2)
add_bullet(doc, "1 599 paroisses catholiques (OFS / ciao.ch 2022). ~200 mosquées (OFS 2022). ~40 synagogues. 5 883 groupes religieux locaux en 2022 (-7,2% vs 2008). (Kirchenstatistik)")
add_bullet(doc, "Sécularisation accélérée : 36,8% sans religion en 2024 → reconversions lieux de culte croissantes (RTS 2024). Segment résiduel pour Presto.")

add_heading(doc, "4.11 Transports", 2)
add_bullet(doc, "Aéroport Genève : 17,8 M passagers 2024. Programme CAP2030 : 560 M CHF (rénovation terminal + Satellite 10, 2025-2032). (newsroom.gva.ch)")
add_bullet(doc, "CFF : 764 gares. Plan ferroviaire 2035 : 8,5 Mds CHF. Programme accessibilité PMR : >2,5 Mds CHF d'ici 2028. Gare Lausanne en travaux 2024-2030 (550 M CHF réseau romand). (CFF / news.sbb.ch)")
add_bullet(doc, "Aéroports ZRH (Zurich) et BSL (Bâle-Mulhouse) : [DONNÉE NON DISPONIBLE — projets spécifiques]. Autoroutes / aires de repos : FORTA (CHF 3 Mds/an).")

add_heading(doc, "4.12 Opportunités Presto par segment — Synthèse", 2)
add_note(doc, "Classement par potentiel décroissant. Score 1 (très faible) à 5 (très fort).")
make_table(doc, [
    ["Segment", "Score", "Types de produits Presto", "Arguments clés", "Canal prioritaire"],
    ["4.2 Santé / EMS", "5/5", "Temporisée hospitalière, infrarouge, anti-brûlure, PMR, anti-arrachement psychiatrie", "Hygiène anti-legionella, exigences techniques élevées, pipelines pluriannuels", "Prescription BET santé, marchés cantonaux (simap.ch)"],
    ["4.11 Transports", "5/5", "Temporisée haute durabilité, infrarouge, PMR, économiseurs eau", "Fréquentation 24h/24, accessibilité obligatoire CFF, robustesse", "CFF marchés publics, GVA CAP2030 (simap.ch)"],
    ["4.5 CHR", "4/5", "Mitigeurs thermostatiques, design, économiseurs", "Record nuitées 2024, marché premium, durabilité", "Grossistes (Meier Tobler), architectes intérieur hôtels"],
    ["4.1 Éducation", "4/5", "Temporisateurs push-button, infrarouge, PMR (SIA 500)", "Économie eau, hygiène post-COVID, robustesse, SIA 500", "Communes, cantons, simap.ch"],
    ["4.7 Sport & Loisirs", "3/5", "Temporisateurs douche, push-button piscines", "Économie eau, hygiène, PMR, 32 000 installations", "Collectivités locales, BET sport"],
    ["4.3 Tertiaire", "3/5", "Temporisateurs, infrarouge, économiseurs", "Smart building, niche multinationales premium", "Architectes, BET, prescripteurs BIM"],
    ["4.8 Pénitentiaire", "3/5", "Inox anti-vandalisme, encastrée, anti-arrachement", "Indestructibilité, saturation prisons = rénov./constructions", "Cantons, administrations pénitentiaires, simap.ch"],
    ["4.4 Industrie", "2/5", "Inox, temporisateurs process", "Hygiène pharma, résistance corrosion", "BET process, installateurs spécialisés"],
    ["4.6 HPA", "2/5", "Économiseurs, temporisateurs", "Tourisme montagne premium", "Groupes camping (TCS)"],
    ["4.9 Culturel", "2/5", "Ponctuel", "Musées réputés, qualité premium", "Ponctuel"],
    ["4.10 Culte", "1/5", "Résiduel", "Sécularisation accélérée", "Non prioritaire"],
], font_size=8)
doc.add_paragraph()
add_para(doc, "Priorités absolues : Santé/EMS (5/5) · Transports (5/5) · CHR (4/5) · Éducation (4/5). Suisse romande = porte d'entrée francophone naturelle.")
doc.add_paragraph()

# ─── PARTIE 5 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 5 — TAILLE MARCHÉ : ROBINETTERIE GÉNÉRALE", 1)

add_heading(doc, "5.1 Taille et valeur — Estimation A (base Analyse de Marché France)", 2)
add_note(doc, "Aucune source publique ne recense la taille exacte du marché suisse de la robinetterie. Estimation par extrapolation depuis la France (protocole MAB). Pas d'étude BRG Suisse disponible : Estimation B non calculable.")
add_heading(doc, "Constantes de référence", 2)
make_table(doc, [
    ["Variable", "Valeur", "Source"],
    ["PIB/hab France 2025", "48 982 USD", "Worldometer (constante MAB)"],
    ["Population France 2025", "69,1 M", "Worldometer (constante MAB)"],
    ["PIB/hab Suisse 2024", "89 783 USD", "Trading Economics / Worldometers"],
    ["Population Suisse 2024", "8,9 M", "OFS / Worldometers"],
    ["Coefficient brut X", "(89 783 / 48 982) × (8,9 / 69,1) = 1,833 × 0,1288 = 0,236", "Calcul MAB"],
    ["Ajustement structurel", "+12% (marché premium, PIB/hab très élevé, culture entretien rigoureuse)", "Hypothèse MAB — à confirmer terrain"],
    ["Coefficient ajusté", "0,236 × 1,12 = 0,265", "Calcul MAB"],
    ["Taux de change", "1 CHF = 1,05 EUR (2024)", "CDTF / Banque de France"],
])
doc.add_paragraph()

add_heading(doc, "Estimation A — Marché robinetterie Suisse (base France ÷ 2, coefficient 0,265)", 2)
make_table(doc, [
    ["Segment", "Base France (÷2)", "× Coeff. 0,265", "Estimation Suisse (M€)", "En M CHF (~×1,05)"],
    ["Robinetterie de collectivités", "100–125 M€", "× 0,265", "26,5 — 33,1 M€", "~28 — 35 M CHF"],
    ["Chasses d'eau & WC collectifs", "90–110 M€", "× 0,265", "23,9 — 29,2 M€", "~25 — 31 M CHF"],
    ["Douches & équipements connexes", "52–65 M€", "× 0,265", "13,8 — 17,2 M€", "~14 — 18 M CHF"],
    ["TOTAL Estimation A", "242–300 M€", "× 0,265", "64,1 — 79,4 M€", "~67 — 83 M CHF"],
])
doc.add_paragraph()
add_note(doc, "Ajustement +12% justifié : PIB/hab très élevé (1,8× France), marché premium avec exigences qualité élevées, prix CHF supérieurs aux prix EUR, culture d'entretien et de renouvellement rigoureuse en Suisse. Estimation par extrapolation — fiabilité moyenne — à confirmer terrain (Meier Tobler, Suissetec, SVGW).")

add_heading(doc, "Estimation B — Base Études BRG", 2)
add_para(doc, "[DONNÉE NON DISPONIBLE — Aucune étude BRG spécifique à la Suisse n'a été identifiée dans les sources internes ou externes. L'unique étude BRG disponible en sources internes couvre la Belgique (BE_Bathrooms_Full_Report_2020.pdf). L'Estimation B ne peut pas être calculée. À demander : europe@brgbuildingsolutions.com]")
doc.add_paragraph()

add_heading(doc, "5.2 Spécificités produit du marché suisse", 2)
add_bullet(doc, "Exigences qualité premium : le marché suisse ERP tolère des prix supérieurs à ceux pratiqués en France, en contrepartie de durabilité et de performance élevées.")
add_bullet(doc, "Laiton sans plomb obligatoire depuis 1904 (OPBD) — standard le plus ancien d'Europe. Alliages certifiés SVGW (ex : CC246E, CuSi4Zn9MnP). Matières plastiques : doivent figurer sur listes positives SVGW.")
add_bullet(doc, "Économie d'eau : fortement valorisée dans les ERP. Limiteurs de débit et aérateurs économiseurs = standard dans les cahiers des charges publics. Smart metering et comptage connecté en croissance.")
add_bullet(doc, "Accessibilité PMR : norme SIA 500 (révision prSIA 500:2025) = standard obligatoire en ERP neuf. Robinetterie à levier ou infrarouge préférée. Hauteur 90-130 cm.")
add_bullet(doc, "4 langues de marché : documentation produit attendue en DE + FR au minimum, IT en option.")
add_bullet(doc, "Robinetterie infrarouge/électronique : croissance forte (smart building CHF 1,35 Md). Préférence pour les produits sans contact dans les ERP post-COVID.")

add_heading(doc, "5.3 Canaux de distribution", 2)
add_note(doc, "Hypothèse basée sur le modèle France (sources internes Presto) adaptée aux spécificités suisses — à confirmer terrain.")
make_table(doc, [
    ["Canal", "Part estimée ERP", "Acteurs clés CH"],
    ["Négoce sanitaire-chauffage (grossistes)", "60–70%", "Meier Tobler SA (CA CHF 496 M 2024 ; 47 points de vente ; 10 000 clients)"],
    ["Grossistes technique bâtiment", "10–15%", "Debrunner Acifer (>160 000 articles, e-shop B2B national)"],
    ["Grossistes Suisse romande", "5–10%", "Bringhen Group / Saneo SA (Bulle) — partenaire Geberit"],
    ["Prescripteurs / vente projet directe", "10–15%", "Architectes SIA, ingénieurs, BIM managers (très influents)"],
    ["Marchés publics (simap.ch)", "~30% en hausse", "simap.ch (obligatoire depuis juil. 2024 au-dessus des seuils)"],
    ["Installateurs (formation Suissetec)", "Canal indirect", "Suissetec : 3 600 entreprises, 24 sections régionales"],
])
add_bullet(doc, "Meier Tobler SA = porte d'entrée distributeur prioritaire pour Presto en Suisse. Canal à développer en priorité. CA 2024 : CHF 496 M (sanitaire + chauffage combinés).")
add_bullet(doc, "[DONNÉE NON DISPONIBLE — répartition exacte des ventes grossistes / prescripteurs / e-commerce en Suisse]")

add_heading(doc, "5.4 Dynamique et perspectives 2025-2030", 2)
add_bullet(doc, "Marché robinetterie CH corrélé à la construction : +2%/an prévu 2025-2026, résilience démontrée. (SSE 2025)")
add_bullet(doc, "Rénovation énergétique = driver structurel → robinetterie économiseurs eau intégrée aux lots sanitaires lors des rénovations (Programme Bâtiments 528 M CHF/an).")
add_bullet(doc, "Robinetterie électronique et infrarouge : croissance forte en Suisse (smart building CHF 1,35 Md). Préférence marquée dans ERP publics post-2022.")
add_bullet(doc, "Vieillissement démographique → augmentation EMS et rénovation hospitalière → segment santé en croissance structurelle 2025-2040.")
add_bullet(doc, "Droits de douane supprimés (HS 8481, 0% depuis jan. 2024) → amélioration compétitivité produits français importés. (kmu.admin.ch)")
doc.add_paragraph()

# ─── PARTIE 6 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 6 — TAILLE MARCHÉ : ROBINETTERIE COLLECTIVE ERP", 1)

add_heading(doc, "6.1 Taille et valeur du marché robinetterie collective ERP", 2)
add_bullet(doc, "[DONNÉE NON DISPONIBLE — Aucune donnée directe sur la taille du marché robinetterie collective ERP en Suisse. Estimations ci-dessous par extrapolation depuis France (méthode MAB). À confirmer via Meier Tobler, Suissetec, SVGW ou études sectorielles.]")
add_bullet(doc, "Estimation robinetterie ERP Suisse (collectivités strictes) : 26,5–33 M€ (~28–35 M CHF) selon Méthode 1. Périmètre élargi (collectivités + WC + douches) : 64–79 M€ (~67–83 M CHF).")

add_heading(doc, "6.2 Méthode d'extrapolation — Détail du calcul", 2)
add_para(doc, "Méthode 1 — base Analyse de Marché France (source interne Presto, déc. 2024)")
make_table(doc, [
    ["Étape", "Calcul", "Résultat"],
    ["PIB/hab Suisse 2024", "Worldometers / Trading Economics", "89 783 USD"],
    ["PIB/hab France 2025 (constante MAB)", "Worldometers", "48 982 USD"],
    ["Ratio PIB/hab", "89 783 / 48 982", "1,833"],
    ["Population Suisse 2024", "OFS / Worldometers", "8,9 M"],
    ["Population France 2025 (constante MAB)", "Worldometers", "69,1 M"],
    ["Ratio population", "8,9 / 69,1", "0,1288"],
    ["Coefficient brut X", "1,833 × 0,1288", "0,236"],
    ["Ajustement structurel", "+12% (marché premium, CHF fort, culture entretien)", "× 1,12"],
    ["Coefficient ajusté", "0,236 × 1,12", "0,265"],
    ["Base France robinetterie collectivités (÷2)", "100–125 M€", "—"],
    ["Estimation Suisse ajustée", "100–125 M€ × 0,265", "26,5–33,1 M€"],
])
doc.add_paragraph()

add_para(doc, "Méthode 2 — base Études BRG")
add_para(doc, "[DONNÉE NON DISPONIBLE — Aucune étude BRG Suisse disponible. Méthode 2 non calculable. Fourchette finale retenue : 26,5–33 M€ (Méthode 1 uniquement). Niveau de confiance : MOYEN. À confirmer terrain.]")
doc.add_paragraph()

make_table(doc, [
    ["", "Méthode 1 (base France ÷ 2)", "Méthode 2 (base BRG Suisse)"],
    ["Base robinetterie collectivités", "100–125 M€ (France, déc. 2024)", "NON DISPONIBLE"],
    ["Coefficient", "0,265 (ajusté +12%)", "N/A"],
    ["Estimation", "26,5–33,1 M€ (~28–35 M CHF)", "N/A"],
    ["Fourchette retenue", "26,5–33 M€", "N/A"],
    ["Niveau de confiance", "Moyen", "N/A"],
    ["Fourchette finale consolidée", "26,5–33 M€ — niveau de confiance MOYEN", ""],
])
doc.add_paragraph()
add_note(doc, "Limites : méthode ne capte pas l'économie informelle (quasi-nulle en Suisse → biais limité). Ne reflète pas les spécificités sectorielles locales. Volatilité CHF/EUR peut fausser la comparaison. Ajustement +12% repose sur hypothèses à confirmer terrain (Meier Tobler, installateurs).")

add_heading(doc, "6.3 Évaluation du potentiel par segment ERP — Scoring", 2)
make_table(doc, [
    ["Segment", "Score", "Justification", "Hypothèses clés"],
    ["4.1 Éducation", "4/5", "Parc dense ~11 700 établissements ; rénovation active ; budget FRI 29,2 Mds CHF", "Données parc à actualiser (OFS 2017/18)"],
    ["4.2 Santé / EMS", "5/5", "278 hôpitaux + 1 465 EMS ; vieillissement ; pipeline HUG+CHUV >1,5 Mds CHF", "Lits/1000 hab 4,5 vs France 5,9 (facteur 0,76)"],
    ["4.3 Tertiaire", "3/5", "Stable ; niche premium multinationales ; surcapacité légère bureaux", "Rénovation > neuf"],
    ["4.4 Industrie", "2/5", "Niche pharma/horlogerie ; valeur unitaire haute mais volumes limités", "Référencement BET process"],
    ["4.5 CHR", "4/5", "42,8 M nuitées record 2024 ; 5 000 établissements ; marché premium", "Facteur CHR 0,216 ; robinetterie design haut de gamme"],
    ["4.6 HPA", "2/5", "Données insuffisantes ; tourisme montagne ; TCS campings", "Qualité attendue élevée malgré segment limité"],
    ["4.7 Sport & Loisirs", "3/5", "32 000 installations ; 1 000 fitness ; 75% pop. sportive", "Dép. sport % PIB non disponible"],
    ["4.8 Pénitentiaire", "3/5", "90 prisons ; 97% occupation ; besoins rénov./constructions", "Facteur incarcération 0,70 ; décisions politiques"],
    ["4.9 Culturel", "2/5", "Musées réputés mais dispersés ; commandes ponctuelles", "Segment opportuniste"],
    ["4.10 Lieux de culte", "1/5", "Sécularisation accélérée ; fermetures > ouvertures", "Résiduel"],
    ["4.11 Transports", "5/5", "764 gares CFF ; CAP2030 560 M CHF ; PMR 2,5 Mds CHF ; lots sanitaires systématiques", "Pipelines publics engagés et visibles sur 10 ans"],
], font_size=8)
doc.add_paragraph()
add_para(doc, "Top 3 : Santé/EMS (5/5) · Transports (5/5) · CHR (4/5) · Éducation (4/5)")

add_heading(doc, "6.4 Spécificités produit robinetterie collective suisse", 2)
add_bullet(doc, "Robinetterie sans plomb obligatoire (OPBD) → laiton sans plomb, alliages certifiés SVGW. Exigence en vigueur depuis 1904.")
add_bullet(doc, "Temporisateurs et infrarouge : standard dans ERP, en forte hausse dans les nouvelles rénovations. Robinetterie connectée (IoT) en croissance.")
add_bullet(doc, "PMR (SIA 500) : obligatoire dans tout ERP neuf ou rénové. Robinetterie à levier ou infrarouge préférée. Hauteur 90-130 cm imposée.")
add_bullet(doc, "Anti-vandalisme inox : requis dans pénitentiaire, psychiatrie, certaines gares. Spécification technique fréquente dans cahiers des charges.")
add_bullet(doc, "Économiseurs d'eau : limiteurs de débit, aérateurs, systèmes anti-gaspillage — valorisés par les donneurs d'ordre publics (objectifs environnement 2050).")

add_heading(doc, "6.5 Dynamique et perspectives 2025-2030", 2)
add_bullet(doc, "Croissance estimée du marché robinetterie collective CH : corrélée à la construction (+2%/an) avec prime pour la rénovation (57% du marché total).")
add_bullet(doc, "Vieillissement démographique → explosion demande EMS sur 2025-2040 : nouveau cycle de construction/rénovation EMS = segment à très fort potentiel de long terme.")
add_bullet(doc, "Transports : pipelines CFF et CAP2030 visibles sur 10 ans = commandes robinetterie PMR/infrarouge/temporisée prévisibles et récurrentes.")
add_bullet(doc, "Rénovation énergétique (Programme Bâtiments 528 M CHF/an) : chaque chantier de rénovation = opportunité de remplacement robinetterie lors des travaux de second œuvre.")
add_bullet(doc, "Risques : Delabie-KWC (production locale post-acquisition Q3 2025) = concurrent dominant durablement implanté. Geberit = force de distribution incomparable (siège suisse). CHF fort = prix compétitifs des produits locaux vs importations. Certification SVGW = prérequis à sécuriser.")
doc.add_paragraph()

# ─── PARTIE 7 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 7 — CONCURRENTS", 1)

add_heading(doc, "7.1 DELABIE (analyse prioritaire)", 2)
add_bullet(doc, "Positionnement : leader européen robinetterie et équipements sanitaires ERP. Haut de gamme institutionnel. Groupe familial français fondé 1928.")
add_bullet(doc, "ACQUISITION STRATÉGIQUE MAJEURE — KWC Professional (Unterkulm, Argovie, CH) : finalisée Q3 2025. Portefeuille acquis : KWC (robinetterie ERP suisse historique), Aquarotter (leader robinetterie collective germanophone), DVS (UK). Production locale maintenue à Unterkulm. (delabie.com / architecture-hospitalière.fr, 2025)")
add_bullet(doc, "Post-acquisition : Delabie = acteur dominant ERP en Suisse, avec usine locale, marques reconnues historiquement, réseau installateurs anciens clients KWC. 70% du CA Delabie sera hors France.")
bold_bullet(doc, "Forces", "Production locale (Unterkulm AG), marques très reconnues (KWC/Aquarotter), réseau installateurs capillaire existant, documentation DE + FR, gamme ERP très complète.")
bold_bullet(doc, "Faiblesses", "Intégration post-acquisition en cours — possibles tensions commerciales transitoires. Prix haut de gamme potentiellement perçus comme inaccessibles sur segments à budget contraint.")
add_bullet(doc, "Impact pour Presto : menace directe et immédiate sur tous les segments ERP suisses. Réponse stratégique urgente nécessaire — différenciation par rapport qualité/prix, réactivité logistique, et Suisse romande.")

add_heading(doc, "7.2 Geberit", 2)
add_bullet(doc, "Siège social : Rapperswil-Jona (Saint-Gall), Suisse. Coté SIX Swiss Exchange. CA mondial 2024 : CHF 3,08 Mds (+2,5% hors change). ~11 000 employés, 26 sites de production, 50+ pays. (geberit.ch / zonebourse.com)")
add_bullet(doc, "Spécialités : bâti-supports, systèmes canalisations, douches, produits salle de bains. ERP : aéroports, stades, bibliothèques, santé.")
add_bullet(doc, "Distribution Suisse : Geberit Distribution SA (Lausanne), réseau installateurs capillaire national, Meier Tobler.")
bold_bullet(doc, "Forces", "'Home market' suisse — image, confiance totale des installateurs, support technique imbattable localement.")
bold_bullet(doc, "Faiblesses", "Gamme moins spécialisée sur robinetterie temporisée ERP stricto sensu vs Delabie/Presto.")

add_heading(doc, "7.3 Autres concurrents", 2)
make_table(doc, [
    ["Concurrent", "Groupe / Pays", "Positionnement", "Gamme ERP", "Canal CH", "Forces / Faiblesses"],
    ["Delabie + KWC", "France / Suisse (post-acq.)", "Haut de gamme, spécialiste ERP", "Complet (temporisée, anti-vand., PMR, IR)", "Distributeurs + installateurs anciens KWC", "F: production locale, marques reconnues ; f: intégration en cours"],
    ["Geberit", "Suisse", "Leader sanitaire global", "Bâti-supports, douches (moins focalisé rob. temporisée)", "Réseau capillaire CH, Meier Tobler", "F: home market suisse ; f: moins spécialisé temporisateurs ERP"],
    ["Grohe", "Allemagne (LIXIL Japon)", "Milieu-haut de gamme", "Thermostatique, infrarouge", "Meier Tobler, revendeurs", "F: notoriété ; f: CA ERP CH non disponible"],
    ["Hansgrohe", "Allemagne", "Premium", "HansaPublic (infrarouge, électronique)", "hansgrohe.ch, revendeurs", "F: design premium ; f: réseau limité ERP"],
    ["Franke", "Suisse (Arlesheim)", "A vendu KWC Pro à Delabie", "Retrait ERP collectif (depuis 2024)", "—", "Retrait du segment — opportunité pour Presto"],
    ["Presto", "France", "Spécialiste ERP", "Temporisée, IR, anti-vand., PMR, encastrée", "3 commerciaux terrain", "F: expertise ERP, rapport Q/P ; f: pas de filiale, notoriété à construire"],
])
doc.add_paragraph()
add_note(doc, "[DONNÉE NON DISPONIBLE — parts de marché précises CHF des concurrents en robinetterie ERP Suisse — à estimer via entretiens Meier Tobler ou installateurs]")

add_heading(doc, "7.4 Opportunités de différenciation pour Presto", 2)
add_bullet(doc, "Niche anti-vandalisme inox : KWC/Aquarotter désormais chez Delabie — les anciens clients KWC cherchant à diversifier leurs achats peuvent se reporter sur Presto. Argument prix/qualité alternatif.")
add_bullet(doc, "Suisse romande : marché naturel francophone (26% population, Genève, Vaud, Valais, Neuchâtel) — avantage commercial direct, documentation FR disponible, proximité culturelle forte.")
add_bullet(doc, "Rapport qualité/prix : alternative crédible à Delabie sur segments à budget contraint (communes, cantons, petits ERP). Profondeur de gamme ERP comparable.")
add_bullet(doc, "Robinetterie infrarouge connectée (IoT) : segment en forte croissance (smart building CHF 1,35 Md). Presto peut se différencier sur ce créneau technologique.")
add_bullet(doc, "PMR (SIA 500) : obligatoire dans tout ERP neuf → levier de prescription systématique via architectes et BET.")
add_bullet(doc, "Arguments écologiques : sans plomb (conforme OPBD), économiseurs eau, durabilité longue → parfaitement alignés avec exigences helvétiques. Stratégie environnement 2050.")
doc.add_paragraph()

# ─── PARTIE 8 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 8 — NORMES & CERTIFICATIONS ROBINETTERIE", 1)

add_heading(doc, "8.1 Certification SVGW (ZertW) — certification eau potable suisse", 2)
bold_bullet(doc, "Organisme", "SVGW (Association suisse du gaz et des eaux / Schweizerischer Verein des Gas- und Wasserfaches). Accrédité SAS selon SN EN ISO/IEC 17065 (n° SCESp 0028). Sièges : Zurich, Lausanne, Bellinzone.")
bold_bullet(doc, "Principe", "Obligatoire de facto pour tout produit en contact avec l'eau potable en Suisse. Le marquage CE européen seul est insuffisant (Suisse hors UE).")
bold_bullet(doc, "Durée et renouvellement", "5 ans, renouvelable. Toute modification hydraulique impose une nouvelle certification.")
bold_bullet(doc, "Périmètre testé", "Hygiène, hydraulique (débits, pressions), mécanique (endurance, étanchéité), physique.")
bold_bullet(doc, "Contact", "info@svgw.ch | +41 44 288 33 33 | svgw.ch/fr/certification/certification-des-produits/")
add_bullet(doc, "Coûts et délais : [NON DISPONIBLES — à demander directement à SVGW. Estimer 3-6 mois selon volume et complexité de la gamme.]")
add_bullet(doc, "Lien avec certifications françaises : les produits conformes NF/ACS français sont proches des exigences SVGW mais ne sont pas automatiquement acceptés. L'ACS peut être présenté comme document de support.")

add_heading(doc, "8.2 Directive W3 SVGW (installations eau potable en bâtiment)", 2)
add_bullet(doc, "Vitesse max. conduites distribution : 1,0 — 2,0 m/s. Pertes de charge compteurs eau : 20-40 kPa. Pertes de charge groupes appareils : 100 kPa standard.")
add_bullet(doc, "Compléments : W3/C3 (hygiène eau potable en bâtiment), W3/C4 (autocontrôle qualité eau).")
add_bullet(doc, "Implications pour Presto : les robinets temporisés doivent respecter ces valeurs de débit et pression. À intégrer dans les fiches techniques suisses.")

add_heading(doc, "8.3 Normes EN applicables (adoptées en Suisse via SNV)", 2)
make_table(doc, [
    ["Norme", "Objet", "Pertinence Presto", "Implications techniques"],
    ["EN 200", "Robinetterie sanitaire générale (eau froide et chaude)", "Base pour toute robinetterie ERP", "Pression nominale PN10, tests endurance, étanchéité, débit, acoustique"],
    ["EN 816", "Robinets à fermeture automatique (temporisés) PN10", "CŒUR DE MÉTIER Presto — prioritaire", "Tests durée fermeture, volume délivré, résistance pressions, endurance mécanique"],
    ["EN 817", "Mitigeurs mécaniques PN10", "Mitigeurs thermostatiques", "Mélange eau F/C, limites anti-brûlure, PN10"],
    ["EN 1112", "Douches sanitaires PN10", "Douches ERP (sport, santé, prison)", "Tests pression dynamique, débit, anti-retour"],
    ["EN 1113", "Flexibles de douche", "Accessoires douches ERP", "Résistance pression, durabilité"],
    ["prSIA 500:2025", "Constructions sans obstacles (PMR)", "Obligatoire ERP neufs et rénovés", "Hauteur 90-130 cm, levier ou IR préférés, espace fauteuil roulant 1,50 m"],
])
doc.add_paragraph()

add_heading(doc, "8.4 Réglementation eau potable OPBD", 2)
add_bullet(doc, "OPBD (Ordonnance sur les denrées alimentaires — eau potable) : matériaux en contact eau potable très strictement réglementés.")
add_bullet(doc, "Interdiction du plomb depuis 1904 — standard mondial le plus ancien. Laiton sans plomb obligatoire : alliages CC246E (CuSi4Zn9MnP) et équivalents certifiés SVGW.")
add_bullet(doc, "Matières plastiques : doivent figurer sur listes positives SVGW. Non-conformité = retrait du marché.")
add_bullet(doc, "Implication Presto : vérifier conformité OPBD pour chaque référence avant commercialisation en Suisse. (blv.admin.ch / svgw.ch)")

add_heading(doc, "8.5 Norme SIA 500 (accessibilité PMR)", 2)
add_bullet(doc, "Révision en cours : prSIA 500:2025-07 (consultation publique). Obligatoire dans tout ERP neuf ou rénové (droit cantonal bâtiment harmonisé). (sia.ch)")
add_bullet(doc, "Hauteur robinetterie / commandes : 90 à 130 cm du sol. Type préféré : robinetterie à levier ou infrarouge.")
add_bullet(doc, "Espace lavabo : approche frontale fauteuil roulant requise. Zone de rotation : minimum 1,50 m de diamètre.")

add_heading(doc, "8.6 Suissetec — rôle prescripteur", 2)
add_bullet(doc, "3 600 entreprises membres, 24 sections régionales. Diffuse les directives SVGW aux installateurs et prescripteurs.")
add_bullet(doc, "Formation 'Connaissances techniques Eau potable' pour installateurs. Référencement Presto auprès de Suissetec = levier de prescription indirecte à activer.")

add_heading(doc, "8.7 Comparaison avec normes françaises", 2)
add_bullet(doc, "Marquage CE/EN : valable en Suisse pour les normes EN harmonisées. Ne remplace pas SVGW pour eau potable.")
add_bullet(doc, "NF/ACS (France) vs SVGW (Suisse) : normes proches mais procédures distinctes. Aucune équivalence automatique. L'ACS peut accélérer l'obtention de la SVGW.")
add_bullet(doc, "Droits de douane : 0% sur HS 8481 depuis janvier 2024. Pas de barrière tarifaire pour produits français importés. (kmu.admin.ch)")
doc.add_paragraph()

# ─── PARTIE 9 ────────────────────────────────────────────────────────────────
add_heading(doc, "PARTIE 9 — POINTS À REVÉRIFIER", 1)
add_para(doc, "Données incertaines ou manquantes à revalider par recherche complémentaire ou terrain :", italic=True)
doc.add_paragraph()

bold_bullet(doc, "Certification SVGW des produits Presto", "Les commerciaux terrain ont-ils des produits certifiés SVGW ? Si non, c'est la priorité #1 avant tout développement commercial significatif. Contact : info@svgw.ch | +41 44 288 33 33.")
bold_bullet(doc, "Taille de marché robinetterie collective CH", "Aucune donnée directe. À obtenir via Meier Tobler (demande confidentielle), SVGW, Suissetec. Alternative : commander l'étude BRG Suisse si elle existe (europe@brgbuildingsolutions.com).")
bold_bullet(doc, "Ajustement structurel +12%", "Hypothèse à confirmer par interviews terrain (Meier Tobler, installateurs). Variables utilisées : marché premium, culture entretien, prix CHF. Tester aussi avec PIB/hab en PPA.")
bold_bullet(doc, "Données scolaires actualisées", "Le parc scolaire (~11 700 établissements) est daté de 2017/18. Actualiser via OFS : bfs.admin.ch/fr/home/statistiques/education-science/institutions-eleves-etudiants.html")
bold_bullet(doc, "Dépenses sport % PIB Suisse", "Non disponibles via Eurostat (CH hors UE). Chercher via OFS (statistiques sport) ou OFSPO (Office fédéral du sport, osp.admin.ch).")
bold_bullet(doc, "Aéroports ZRH (Zurich) et BSL (Bâle-Mulhouse)", "Projets d'expansion/rénovation non identifiés. Vérifier sur flughafen-zuerich.ch et euroairport.com.")
bold_bullet(doc, "CHUV Lausanne", "Projets de rénovation hospitalière en cours non chiffrés. Vérifier sur chuv.ch ou simap.ch.")
bold_bullet(doc, "Nombre campings / HPA Suisse", "Non trouvé. Chercher via OFS tourisme ou TCS (Touring Club Suisse, tcs.ch).")
bold_bullet(doc, "Parts de marché concurrents (Grohe, Hansgrohe, Delabie) en CH", "Non disponibles. À estimer via entretiens Meier Tobler ou installateurs.")
bold_bullet(doc, "Prix catalogue robinetterie collective CHF", "Aucun exemple trouvé. À obtenir via Meier Tobler ou catalogues Delabie/KWC (kwc.com). Indicateur de positionnement tarifaire essentiel.")
bold_bullet(doc, "Appels d'offres simap.ch", "Recherche directe sur simap.ch (CPV 44411000) requise. Mots-clés : Sanitär, Armaturen, Unterputz, robinetterie, temporisé. Archive pré-juillet 2024 séparée.")
bold_bullet(doc, "Presto en Suisse", "Coordonnées exactes des commerciaux terrain, clients actuels, produits déjà référencés, distributeurs partenaires CH.")
doc.add_paragraph()

# Sauvegarde
path_etude = os.path.join(OUTPUT_DIR, "MAB_Suisse_Etude.docx")
doc.save(path_etude)
print(f"✓ Étude v1 sauvegardée : {path_etude}")


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT ANNEXES
# ════════════════════════════════════════════════════════════════════════════
ann = Document()
set_margins(ann)

t2 = ann.add_heading("MAB SUISSE — ANNEXES & SOURCES v1", 0)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2 = ann.add_paragraph("Sources complètes, données brutes et compléments — Les Robinets Presto")
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.runs[0].font.size = Pt(11)
sub2.runs[0].font.name = FONT
ann.add_paragraph()

add_heading(ann, "ANNEXE 1 — LISTE DES SOURCES UTILISÉES", 1)
ann.add_paragraph("Toutes les sources consultées pour MAB_Suisse_Etude.docx v1 :").runs[0].font.italic = True
ann.add_paragraph()

sources = [
    ("Worldometers — PIB Suisse 2024", "https://www.worldometers.info/fr/pib/suisse-pib/", "Français", "Juin 2026"),
    ("Trading Economics — PIB/hab Suisse 2024", "https://tradingeconomics.com/switzerland/gdp-per-capita", "Anglais", "Juin 2026"),
    ("OFS / bfs.admin.ch — Statistiques Suisse", "https://www.bfs.admin.ch/bfs/fr/home.html", "Français/Allemand", "Juin 2026"),
    ("Eminence.ch — Swiss Construction Insights 2025", "https://eminence.ch/en/swiss-construction-insights-2025/", "Anglais", "2025"),
    ("SSE / Baumeister.swiss — Construction 23 Mds CHF 2024", "https://baumeister.swiss/fr/construction-23-milliards-de-francs-en-2024/", "Français", "2024"),
    ("SSE — Prévisions construction 2025-2026 (+2%)", "https://www.presseportal.ch/fr/pm/100051907/100938632", "Français", "2024"),
    ("Direction Générale du Trésor — Échanges France-Suisse 2026", "https://www.tresor.economie.gouv.fr/Articles/2026/04/10/les-echanges-commerciaux-bilateraux-france-suisse", "Français", "2026"),
    ("Direction du Trésor — Présence française en Suisse", "https://www.tresor.economie.gouv.fr/Articles/2019/02/21/presence-et-investissements-francais-en-suisse", "Français", "2019/2021"),
    ("Delabie — Acquisition KWC Professional finalisée (communiqué officiel)", "https://www.delabie.com/our-group/group-news/group-news/the-delabie-group-finalises-the-acquisition-of-kwc-professional", "Anglais", "2025"),
    ("Architecture Hospitalière — Delabie-KWC dimension européenne", "https://www.architecture-hospitaliere.fr/blog/2025/09/25/delabie-une-nouvelle-dimension-europeenne-avec-lintegration-de-kwc-professional/", "Français", "2025"),
    ("Geberit — L'entreprise", "https://www.geberit.ch/a-propos-de-nous/l-entreprise-geberit/", "Français", "Juin 2026"),
    ("Meier Tobler SA — Résultats 2024", "https://domotech-magazine.ch/meier-tobler-resiste-dans-un-environnement-de-marche-difficile", "Français", "2024"),
    ("SVGW — Certification produits (ZertW)", "https://www.svgw.ch/fr/certification/certification-des-produits/", "Français", "Juin 2026"),
    ("SVGW — Directive W3 installations eau potable", "https://www.svgw.ch/fr/eau/reglementation/faq/w3-installations-deau-potable/", "Français", "Juin 2026"),
    ("OFS — Hôpitaux 2022", "https://www.bfs.admin.ch/bfs/fr/home/statistiques/sante/systeme-sante/hopitaux.html", "Français", "2022"),
    ("OFS — EMS / SOMED 2023", "https://www.bfs.admin.ch/bfs/fr/home/statistiques/sante/systeme-sante/etablissements-medico-sociaux.html", "Français", "2023"),
    ("Senesuisse — Faits et chiffres EMS 2023", "https://www.senesuisse.ch/fr/news/actualites/1184-faits-et-chiffres-les-plus-recents-des-ems-suisses", "Français", "2023"),
    ("HUG Genève — Construire l'hôpital de demain", "https://www.hug.ch/construire-hopital-demain", "Français", "2024"),
    ("OFS — Statistiques hôtellerie 2024 (record nuitées)", "https://www.bfs.admin.ch/bfs/en/home.assetdetail.34307641.html", "Anglais", "2024"),
    ("news.admin.ch — Programme Bâtiments 528 M CHF 2024", "https://www.news.admin.ch/fr/newnsb/IKoj6VY8s85qnTeXJrLxc", "Français", "2024"),
    ("Prison Insider — Suisse 2024", "https://www.prison-insider.com/en/countryprofile/suisse-2024-6707ab9ed4369", "Anglais", "2024"),
    ("Genève Aéroport — CAP2030 (neuf questions)", "https://newsroom.gva.ch/le-projet-cap2030-en-neuf-questions/", "Français", "2023"),
    ("CFF / SBB — Accessibilité transports publics", "https://news.sbb.ch/fr/medias/article/120058/accessibilite-des-transports-publics-les-cff-poursuivent-leur-effort", "Français", "2024"),
    ("Union des villes suisses — Statistiques 2024", "https://uniondesvilles.ch/fr/detail/statistiques-des-villes-suisses-2024", "Français", "2024"),
    ("CBRE / immoday.ch — Marché bureaux CH 2024", "https://immoday.ch", "Français", "2024"),
    ("EDA Suisse — Installations sportives", "https://www.aboutswitzerland.eda.admin.ch", "Français", "Juin 2026"),
    ("Watson.ch / LFM — Fitness Suisse 2024", "https://www.watson.ch", "Français/Allemand", "2024"),
    ("Suissetec — Site officiel", "https://www.suissetec.ch", "Français/Allemand", "Juin 2026"),
    ("kmu.admin.ch — Suppression droits de douane biens industriels 2024", "https://www.kmu.admin.ch", "Français", "2024"),
    ("Global Innovation Index 2024 — Suisse rang 2", "https://www.globalinnovationindex.org", "Anglais", "2024"),
    ("Analyse Marché Sanitaire Lieux publics France — source interne Presto", "MAB-core/sources-internes/Analyse Marché Sanitaire Lieux public France.pdf", "Français", "Décembre 2024"),
    ("EMAE — Extrapolation, notes de recherche — source interne", "MAB-core/sources-internes/EMAE - Extrapolation, notes de recherche.docx", "Français", "2025"),
    ("MAB_Suisse_PREP.md — Corpus pré-étude PREP", "MAB-core/pays/suisse/MAB_Suisse_PREP.md", "Français", "Juin 2026"),
]

for i, (title, url, lang, date) in enumerate(sources, 1):
    p = ann.add_paragraph(style="List Number")
    r = p.add_run(f"[{i}] {title}")
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = FONT
    p2 = ann.add_paragraph(f"    URL/Chemin : {url}")
    p2.runs[0].font.size = Pt(8)
    p2.runs[0].font.name = FONT
    p2.runs[0].font.color.rgb = RGBColor(0x00, 0x56, 0xA2)
    p3 = ann.add_paragraph(f"    Langue : {lang} | Consulté/daté : {date}")
    p3.runs[0].font.size = Pt(8)
    p3.runs[0].font.italic = True
    p3.runs[0].font.name = FONT

ann.add_paragraph()

add_heading(ann, "ANNEXE 2 — CALCUL DÉTAILLÉ EXTRAPOLATION MARCHÉ ROBINETTERIE SUISSE", 1)
ann.add_paragraph("Source base France : Analyse de Marché France — Équipements Sanitaires Automatiques pour Lieux Publics, Presto, décembre 2024. Marché total France 2024 : 485-550 M€ HT.").runs[0].font.italic = True
ann.add_paragraph()
make_table(ann, [
    ["Paramètre", "Valeur", "Source"],
    ["PIB/hab France 2025 (constante MAB)", "48 982 USD", "Worldometers (constante CLAUDE.md)"],
    ["Population France 2025 (constante MAB)", "69,1 millions", "Worldometers (constante CLAUDE.md)"],
    ["PIB/hab Suisse 2024", "89 783 USD", "Trading Economics / Worldometers"],
    ["Population Suisse 2024", "8,9 millions", "OFS / Worldometers"],
    ["Ratio PIB/hab (CH/FR)", "89 783 / 48 982 = 1,833", "Calcul MAB"],
    ["Ratio population (CH/FR)", "8,9 / 69,1 = 0,1288", "Calcul MAB"],
    ["Coefficient brut X", "1,833 × 0,1288 = 0,236", "Calcul MAB"],
    ["Ajustement structurel", "+12% (marché premium, PIB/hab très élevé, culture entretien)", "Hypothèse MAB"],
    ["Coefficient ajusté", "0,236 × 1,12 = 0,265", "Calcul MAB"],
    ["Taux de change", "1 CHF = 1,05 EUR (moyenne 2024)", "CDTF / Banque de France"],
])
ann.add_paragraph()
make_table(ann, [
    ["Segment France (base ÷2)", "Base France (M€)", "× 0,265", "Estimation CH (M€)", "En M CHF (×1,05)"],
    ["Robinetterie de collectivités", "100–125 M€", "× 0,265", "26,5–33,1 M€", "~28–35 M CHF"],
    ["Chasses d'eau & WC collectifs", "90–110 M€", "× 0,265", "23,9–29,2 M€", "~25–31 M CHF"],
    ["Douches & équipements connexes", "52–65 M€", "× 0,265", "13,8–17,2 M€", "~14–18 M CHF"],
    ["TOTAL", "242–300 M€", "× 0,265", "64,1–79,4 M€", "~67–83 M CHF"],
])
ann.add_paragraph()
ann.add_paragraph("Justification ajustement structurel +12% : (1) PIB/hab suisse 1,8× la France → pouvoir d'achat et prix acceptés bien supérieurs ; (2) Marché premium avec exigences qualité élevées sur ERP publics ; (3) Culture d'entretien et de renouvellement rigoureuse ; (4) Prix en CHF supérieurs aux prix EUR à qualité équivalente. Hypothèse à confirmer via entretiens terrain (Meier Tobler, installateurs Suissetec).").runs[0].font.size = Pt(9)
ann.add_paragraph()

add_heading(ann, "ANNEXE 3 — DONNÉES CONSTRUCTION SUISSE 2024-2026", 1)
make_table(ann, [
    ["Indicateur", "2024", "2025 (prév.)", "2026 (prév.)", "Source"],
    ["Investissement total construction (Mds CHF)", "68,9 (+1,8%)", "~70,4 (+2,1%)", "~71,7 (+1,9%)", "Eminence.ch / SSE"],
    ["CA entrepreneurs construction (Mds CHF)", "23,4", "23,9 (+2,1%)", "24,4 (+1,9%)", "SSE / Baumeister.swiss"],
    ["Part rénovation vs neuf", "57% réno / 43% neuf", "~57%", "~58%", "Eminence.ch 2024"],
    ["Résidentiel neuf (croissance)", "+4,8% (2025)", "+4,8%", "stable", "EMAE / Eminence.ch"],
    ["Non-résidentiel", "Stable — pas de croissance", "Stable", "Stable", "EMAE 2025"],
    ["Génie civil", "Stable (FORTA/CFF)", "Stable", "Stable", "EMAE 2025"],
    ["Carnets de commandes", "7,4 mois (record 6 ans)", "—", "—", "SIA Q4 2024"],
    ["Emplois secteur", "352 000 (8,1% emploi national)", "—", "—", "Eminence.ch"],
    ["Programme Bâtiments versé", "528 M CHF (dont 247 M fédéral)", "~528 M CHF/an", "~528 M CHF/an", "news.admin.ch"],
])
ann.add_paragraph()

add_heading(ann, "ANNEXE 4 — PIPELINE INVESTISSEMENTS ERP SUISSE 2024-2035", 1)
make_table(ann, [
    ["Segment", "Porteur", "Montant", "Calendrier", "Statut", "Source"],
    ["HUG Genève — rénovation hospitalière (total)", "Cantons GE + Confédération", ">1,5 Mds CHF HUG+CHUV", "2024-2031", "En cours", "hug.ch"],
    ["HUG — investissements 2024", "HUG", "58,5 M CHF", "2024", "Réalisé", "hug.ch"],
    ["Programme accessibilité CFF (PMR)", "CFF / Confédération", ">2,5 Mds CHF", "2024-2028", "En cours", "news.sbb.ch"],
    ["Plan ferroviaire CFF 2035", "CFF / Confédération", "8,5 Mds CHF", "2024-2035", "Engagé", "CFF"],
    ["CAP2030 Aéroport Genève", "Genève Aéroport", "560 M CHF", "2025-2032", "En cours", "newsroom.gva.ch"],
    ["Programme Bâtiments (rénovation énergétique)", "Confédération + Cantons", "528 M CHF/an (récurrent)", "2024-2030+", "Continu", "news.admin.ch"],
    ["FORTA (routes nationales)", "Confédération", "~3 Mds CHF/an ; 11,6 Mds à 2030", "Continu", "Continu", "DETEC"],
    ["Budget FRI 2025-2028 (éducation)", "Confédération", "29,2 Mds CHF (coupes 460 M/an)", "2025-2028", "Adopté (avec coupes)", "Conseil fédéral"],
    ["Prisons suisses (saturation 97%)", "Cantons", "Non chiffré", "À planifier", "En discussion", "Prison Insider 2024"],
], font_size=8)
ann.add_paragraph()

add_heading(ann, "ANNEXE 5 — DONNÉES ERP SUISSE PAR SEGMENT", 1)
make_table(ann, [
    ["Segment", "Données clés", "Source", "Niveau de confiance"],
    ["Éducation", "~11 700 établissements primaire+secondaire (données 2017/18) ; ETH 23 500 étudiants ; EPFL 12 000 ; 37 HES ; 549 595 élèves primaire 2023", "OFS 2017/18 et 2023 / ETH / EPFL", "Basse (données 2017/18 à actualiser)"],
    ["Santé / Hôpitaux", "278 établissements (264 H + 14 maternités) ; 37 925-38 000 lits ; 4,4-4,6 lits/1000 hab", "OFS 2022-2023 / OCDE", "Haute"],
    ["EMS", "1 465 établissements ; 100 540 places ; 170 211 résidents ; coûts 11,65 Mds CHF", "OFS SOMED 2023 / Senesuisse", "Très haute"],
    ["Tertiaire / Bureaux", "995 500 m² disponibles (5 villes) ; vacance 5,0%", "CBRE / immoday.ch 2024", "Haute"],
    ["CHR", "~5 000 hôtels ; 42,8 M nuitées 2024 ; 22 M nuitées étrangères ; recettes 20,33 Mds EUR", "OFS 2024 / Fédération suisse du tourisme", "Très haute"],
    ["Sport & Loisirs", "32 000 installations sportives ; 980-1 000 centres fitness ; 1,37 M membres fitness ; CA CHF 1,3 Md", "EDA / LFM / Watson.ch 2024", "Bonne"],
    ["Pénitentiaire", "90 établissements ; 7 373 places ; 7 119 détenus (2026) ; taux occupation 97% ; 78/100 000 hab", "Prison Insider / RTS 2024", "Haute"],
    ["Transports", "764 gares CFF ; 17,8 M passagers GVA 2024 ; investis. CFF PMR >2,5 Mds CHF ; CAP2030 560 M CHF", "CFF / GVA newsroom / news.sbb.ch 2024", "Très haute"],
    ["Lieux de culte", "1 599 paroisses catholiques ; ~200 mosquées ; ~40 synagogues ; 36,8% sans religion 2024", "OFS 2022 / ciao.ch / Kirchenstatistik", "Bonne"],
    ["HPA campings", "NON DISPONIBLE", "—", "Non disponible"],
    ["Culturel", "NON DISPONIBLE (nombre cinémas, théâtres, musées)", "—", "Non disponible"],
])
ann.add_paragraph()

add_heading(ann, "ANNEXE 6 — DISTRIBUTEURS ROBINETTERIE SUISSE", 1)
make_table(ann, [
    ["Distributeur", "Type", "Zone", "CA / Taille", "Remarques"],
    ["Meier Tobler SA", "Grossiste sanitaire-chauffage #1 CH", "National (47 points de vente)", "CA CHF 496 M (2024, -9,1% YoY)", "Canal d'entrée prioritaire Presto ; 10 000 clients installateurs ; partenaire Geberit"],
    ["Debrunner Acifer", "Grossiste acier + technique bâtiment", "National", ">160 000 articles ; e-shop B2B", "Robinetterie dans catalogue ; réseau national"],
    ["Bringhen Group / Saneo SA", "Grossiste sanitaire Suisse romande", "Suisse romande (Bulle)", "PME régionale", "Partenaire Geberit ; KWC, Laufen, Vola, Neoperl"],
    ["simap.ch", "Plateforme marchés publics fédérale", "National", "—", "Obligatoire depuis juillet 2024 ; CPV 44411000 = robinetterie"],
    ["Suissetec (installateurs)", "Réseau prescripteurs-installateurs", "National (24 sections)", "3 600 entreprises membres", "Canal prescripteur indirect ; formation SVGW"],
])
ann.add_paragraph()

add_heading(ann, "ANNEXE 7 — NORMES ET CERTIFICATIONS DÉTAILLÉES ROBINETTERIE SUISSE", 1)
make_table(ann, [
    ["Norme / Certification", "Objet", "Statut CH", "Implications techniques Presto", "Organisme"],
    ["SVGW ZertW", "Certification eau potable suisse", "OBLIGATOIRE de facto", "Hygiène, hydraulique, mécanique, physique ; 5 ans renouvelable", "SVGW (Zurich / Lausanne)"],
    ["Directive W3 SVGW", "Installations eau potable en bâtiment", "Obligatoire (référence)", "Vitesse 1-2 m/s ; pertes de charge 100 kPa standard", "SVGW"],
    ["EN 816 (SN EN 816)", "Robinets fermeture automatique PN10", "Applicable (adopté SNV)", "CŒUR GAMME Presto — durée fermeture, volume, endurance, acoustique", "SNV / CEN"],
    ["EN 200 (SN EN 200)", "Robinetterie sanitaire générale PN10", "Applicable (adopté SNV)", "Base robinetterie ERP — pression, débit, étanchéité, endurance", "SNV / CEN"],
    ["EN 817 (SN EN 817)", "Mitigeurs mécaniques PN10", "Applicable (adopté SNV)", "Mitigeurs thermostatiques — mélange F/C, anti-brûlure", "SNV / CEN"],
    ["EN 1112 (SN EN 1112)", "Douches sanitaires PN10", "Applicable (adopté SNV)", "Douches ERP — pression dynamique, débit, anti-retour", "SNV / CEN"],
    ["prSIA 500:2025", "Accessibilité PMR / constructions sans obstacles", "Obligatoire ERP neufs", "Hauteur 90-130 cm ; espace fauteuil roulant 1,50 m", "SIA"],
    ["OPBD", "Ordonnance eau potable — matériaux", "Obligatoire", "Sans plomb depuis 1904 ; laiton CC246E ; listes positives plastiques", "OSAV / blv.admin.ch"],
    ["NF/ACS (France)", "Certification eau potable française", "Non reconnue directement", "Peut être présenté comme support dossier SVGW", "AFNOR / CRECEP"],
    ["Marquage CE (EU 305/2011)", "Mise sur marché UE", "Insuffisant seul pour eau potable CH", "Prérequis mais ne remplace pas SVGW", "Commission UE"],
], font_size=8)
ann.add_paragraph()

add_heading(ann, "ANNEXE 8 — MATRICE DE CONFIANCE DONNÉES CLÉS", 1)
make_table(ann, [
    ["Donnée", "Valeur", "Source", "Niveau confiance", "À revérifier"],
    ["PIB/hab CH 2024", "89 783 USD", "Trading Economics / Worldometers", "Très haute", "Non"],
    ["Population CH 2024", "8,9 M (9,007 M en 2026)", "OFS / Worldometers", "Très haute", "Non"],
    ["Taux change 2024", "1 CHF = 1,05 EUR", "CDTF / Banque de France", "Très haute", "Non"],
    ["Investissement construction CH 2024", "CHF 68,9 Mds", "Eminence.ch", "Bonne", "Croiser OFS"],
    ["CA entrepreneurs construction 2024", "CHF 23,4 Mds", "SSE / Baumeister.swiss", "Très haute", "Non"],
    ["Split neuf/réno", "57% réno / 43% neuf", "Eminence.ch", "Bonne", "Croiser OFS"],
    ["Nombre hôpitaux CH", "278 (264 + 14 maternités)", "OFS 2022", "Haute", "Actualiser"],
    ["Lits hospitaliers", "~37 925-38 000", "OFS 2023", "Haute", "Non"],
    ["Nombre EMS", "1 465 / 100 540 places", "OFS SOMED 2023", "Très haute", "Non"],
    ["Nuitées hôtelières 2024", "42,8 M (record)", "OFS 2024", "Très haute", "Non"],
    ["Installations sportives", "32 000", "EDA", "Bonne", "Non"],
    ["Détenus CH 2024", "6 881 (janv. 2024)", "Prison Insider", "Très haute", "Non"],
    ["Gares CFF", "764", "CFF", "Très haute", "Non"],
    ["CAP2030 Genève", "560 M CHF 2025-2032", "GVA newsroom", "Très haute", "Non"],
    ["Acquisition Delabie-KWC", "Finalisée Q3 2025", "delabie.com", "Très haute", "Non"],
    ["Geberit CA mondial 2024", "CHF 3,08 Mds", "geberit.ch / zonebourse", "Très haute", "Non"],
    ["Meier Tobler CA 2024", "CHF 496 M", "domotech-magazine.ch", "Très haute", "Non"],
    ["SVGW obligatoire (principe)", "Oui, 5 ans renouvelable", "svgw.ch", "Très haute", "Coûts/délais"],
    ["Programme Bâtiments 2024", "528 M CHF versés", "news.admin.ch", "Très haute", "Non"],
    ["Coefficient extrapolation brut", "0,236", "Calcul MAB", "Bonne", "Non"],
    ["Estimation rob. collective CH (adj.)", "26,5–33 M€", "Extrapolation × 0,265", "Moyenne", "Confirmer terrain"],
    ["Nombre campings CH", "NON DISPONIBLE", "—", "—", "OFS tourisme / TCS"],
    ["Taille marché rob. CH (directe)", "NON DISPONIBLE", "—", "—", "Meier Tobler / terrain"],
    ["Étude BRG Suisse", "NON DISPONIBLE", "—", "—", "BRG Building Solutions"],
], font_size=8)
ann.add_paragraph()

# Sauvegarde annexes
path_ann = os.path.join(OUTPUT_DIR, "MAB_Suisse_Annexes.docx")
ann.save(path_ann)
print(f"✓ Annexes v1 sauvegardées : {path_ann}")
print()
print("═" * 60)
print("MAB SUISSE v1 — Génération terminée")
print(f"  Étude   → {path_etude}")
print(f"  Annexes → {path_ann}")
print("═" * 60)
